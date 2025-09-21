#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
图片删除功能测试脚本
测试DOCX文档中的图片删除功能
"""

import os
import sys
import logging
from pathlib import Path

# 添加项目根目录到Python路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from core.enhanced_docx_processor import EnhancedDocxProcessor

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_test_document_with_images():
    """创建一个包含多张图片的测试文档"""
    print("📝 创建包含多张图片的测试文档...")
    
    processor = EnhancedDocxProcessor()
    
    # 创建新文档
    doc_path = "../docs/test_image_deletion.docx"
    processor.create_document(doc_path)
    
    # 添加标题
    processor.add_heading("图片删除功能测试", 1)
    processor.add_paragraph("这个文档将测试图片删除功能")
    
    # 插入第一张图片
    print("📸 插入第一张图片...")
    if os.path.exists("../images/test_image.jpg"):
        result = processor.add_image("../images/test_image.jpg", width="2in", height="1.5in")
        print(f"第一张图片插入结果: {result}")
    else:
        print("⚠️ test_image.jpg 不存在")
    
    # 添加一些文字
    processor.add_paragraph("这是第一张图片下方的文字")
    
    # 插入第二张图片
    print("📸 插入第二张图片...")
    if os.path.exists("../images/test_image.png"):
        result = processor.add_image("../images/test_image.png", width="2.5in", height="2in", alignment="center")
        print(f"第二张图片插入结果: {result}")
    else:
        print("⚠️ test_image.png 不存在")
    
    # 添加更多文字
    processor.add_paragraph("这是第二张图片下方的文字")
    
    # 插入第三张图片
    print("📸 插入第三张图片...")
    if os.path.exists("../images/test_image.bmp"):
        result = processor.add_image("../images/test_image.bmp", width="3in", height="2.5in")
        print(f"第三张图片插入结果: {result}")
    else:
        print("⚠️ test_image.bmp 不存在")
    
    # 添加结尾文字
    processor.add_paragraph("这是文档的结尾")
    
    # 保存文档
    processor.save_document()
    print(f"✅ 测试文档已创建: {doc_path}")
    
    return doc_path

def test_image_deletion_scenarios():
    """测试各种图片删除场景"""
    print("\n🗑️ 开始测试图片删除功能...")
    
    processor = EnhancedDocxProcessor()
    
    # 打开测试文档
    doc_path = "../docs/test_image_deletion.docx"
    if not os.path.exists(doc_path):
        print(f"❌ 测试文档不存在: {doc_path}")
        return
    
    processor.open_document(doc_path)
    
    # 检查初始图片数量
    print("\n🔍 检查初始图片数量...")
    result = processor.list_images()
    print(f"初始图片列表: {result}")
    
    # 测试场景1: 删除第一张图片
    print("\n📸 测试场景1: 删除第一张图片 (索引0)")
    result = processor.delete_image(0)
    print(f"删除结果: {result}")
    
    # 检查删除后的图片数量
    result = processor.list_images()
    print(f"删除第一张图片后的列表: {result}")
    
    # 保存删除第一张图片后的文档
    processor.save_as_document("../docs/test_image_deletion_after_first_delete.docx")
    print("✅ 删除第一张图片后的文档已保存")
    
    # 测试场景2: 删除最后一张图片
    print("\n📸 测试场景2: 删除最后一张图片 (索引1)")
    result = processor.delete_image(1)
    print(f"删除结果: {result}")
    
    # 检查删除后的图片数量
    result = processor.list_images()
    print(f"删除最后一张图片后的列表: {result}")
    
    # 保存删除最后一张图片后的文档
    processor.save_as_document("../docs/test_image_deletion_after_last_delete.docx")
    print("✅ 删除最后一张图片后的文档已保存")
    
    # 测试场景3: 删除剩余的图片
    print("\n📸 测试场景3: 删除剩余的图片 (索引0)")
    result = processor.delete_image(0)
    print(f"删除结果: {result}")
    
    # 检查删除后的图片数量
    result = processor.list_images()
    print(f"删除所有图片后的列表: {result}")
    
    # 保存删除所有图片后的文档
    processor.save_as_document("../docs/test_image_deletion_all_deleted.docx")
    print("✅ 删除所有图片后的文档已保存")
    
    # 测试场景4: 尝试删除不存在的图片
    print("\n📸 测试场景4: 尝试删除不存在的图片 (索引0)")
    result = processor.delete_image(0)
    print(f"删除结果: {result}")
    
    # 测试场景5: 尝试删除负索引
    print("\n📸 测试场景5: 尝试删除负索引 (-1)")
    result = processor.delete_image(-1)
    print(f"删除结果: {result}")
    
    # 测试场景6: 尝试删除超出范围的索引
    print("\n📸 测试场景6: 尝试删除超出范围的索引 (999)")
    result = processor.delete_image(999)
    print(f"删除结果: {result}")

def verify_document_integrity():
    """验证文档完整性"""
    print("\n🔬 验证文档完整性...")
    
    try:
        from docx import Document
        
        # 检查原始文档
        original_doc = Document("../docs/test_image_deletion.docx")
        print(f"原始文档段落数: {len(original_doc.paragraphs)}")
        
        # 检查删除第一张图片后的文档
        first_delete_doc = Document("../docs/test_image_deletion_after_first_delete.docx")
        print(f"删除第一张图片后段落数: {len(first_delete_doc.paragraphs)}")
        
        # 检查删除最后一张图片后的文档
        last_delete_doc = Document("../docs/test_image_deletion_after_last_delete.docx")
        print(f"删除最后一张图片后段落数: {len(last_delete_doc.paragraphs)}")
        
        # 检查删除所有图片后的文档
        all_deleted_doc = Document("../docs/test_image_deletion_all_deleted.docx")
        print(f"删除所有图片后段落数: {len(all_deleted_doc.paragraphs)}")
        
        # 验证文字内容是否保持完整
        print("\n📋 验证文字内容完整性...")
        
        # 检查原始文档的文字内容
        original_text = [p.text for p in original_doc.paragraphs if p.text.strip()]
        print(f"原始文档文字段落: {len(original_text)}")
        for i, text in enumerate(original_text):
            print(f"  段落 {i}: {text[:50]}...")
        
        # 检查删除所有图片后的文档文字内容
        final_text = [p.text for p in all_deleted_doc.paragraphs if p.text.strip()]
        print(f"删除所有图片后文字段落: {len(final_text)}")
        for i, text in enumerate(final_text):
            print(f"  段落 {i}: {text[:50]}...")
        
        # 验证文字内容是否一致
        if original_text == final_text:
            print("✅ 文字内容保持完整，删除图片操作成功")
        else:
            print("⚠️ 文字内容有变化，需要检查删除逻辑")
            
    except Exception as e:
        print(f"❌ 验证文档完整性失败: {e}")

def test_edge_cases():
    """测试边界情况"""
    print("\n🧪 测试边界情况...")
    
    processor = EnhancedDocxProcessor()
    
    # 创建一个只有图片的文档
    print("\n📝 创建只有图片的文档...")
    doc_path = "../docs/test_image_deletion_edge_cases.docx"
    processor.create_document(doc_path)
    
    # 插入图片
    if os.path.exists("../images/test_image.jpg"):
        processor.add_image("../images/test_image.jpg", width="2in", height="1.5in")
        print("✅ 插入测试图片")
    
    # 保存文档
    processor.save_document()
    
    # 打开文档
    processor.open_document(doc_path)
    
    # 检查图片
    result = processor.list_images()
    print(f"图片列表: {result}")
    
    # 删除唯一的图片
    print("\n🗑️ 删除唯一的图片...")
    result = processor.delete_image(0)
    print(f"删除结果: {result}")
    
    # 检查删除后的状态
    result = processor.list_images()
    print(f"删除后的图片列表: {result}")
    
    # 保存结果
    processor.save_as_document("../docs/test_image_deletion_edge_cases_result.docx")
    print("✅ 边界情况测试文档已保存")

def main():
    """主函数"""
    print("🚀 图片删除功能测试开始")
    print("=" * 60)
    
    # 创建测试文档
    doc_path = create_test_document_with_images()
    
    # 测试图片删除功能
    test_image_deletion_scenarios()
    
    # 验证文档完整性
    verify_document_integrity()
    
    # 测试边界情况
    test_edge_cases()
    
    print("\n" + "=" * 60)
    print("🎯 图片删除功能测试完成！")
    print("\n📋 生成的文件:")
    print("- ../docs/test_image_deletion.docx (原始测试文档)")
    print("- ../docs/test_image_deletion_after_first_delete.docx (删除第一张图片后)")
    print("- ../docs/test_image_deletion_after_last_delete.docx (删除最后一张图片后)")
    print("- ../docs/test_image_deletion_all_deleted.docx (删除所有图片后)")
    print("- ../docs/test_image_deletion_edge_cases.docx (边界情况测试)")
    print("- ../docs/test_image_deletion_edge_cases_result.docx (边界情况测试结果)")
    
    print("\n💡 测试要点:")
    print("1. 图片删除后，文档结构应保持完整")
    print("2. 文字内容不应被意外删除")
    print("3. 空段落应被自动清理")
    print("4. 索引验证应正确工作")
    print("5. 边界情况应被正确处理")

if __name__ == "__main__":
    main()
