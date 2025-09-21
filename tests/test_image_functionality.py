#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
图片功能测试脚本
测试DOCX文档中的图片插入、显示和操作功能
"""

import os
import sys
import logging
from pathlib import Path

# 添加项目根目录到Python路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from core.enhanced_docx_processor import EnhancedDocxProcessor
from core.state_manager import StateManager

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def test_image_functionality():
    """测试图片功能"""
    print("🔍 开始测试图片功能...")
    
    # 创建处理器实例
    processor = EnhancedDocxProcessor()
    
    # 创建测试文档
    print("\n1️⃣ 创建测试文档...")
    result = processor.create_document("test_image_document.docx")
    print(f"结果: {result}")
    
    # 添加一些内容
    print("\n2️⃣ 添加标题和段落...")
    processor.add_heading("图片功能测试", 1)
    processor.add_paragraph("这是一个测试图片插入功能的文档")
    
    # 检查是否有测试图片
    test_image_path = "../images/test_image.jpg"
    if not os.path.exists(test_image_path):
        # 创建一个简单的测试图片（如果不存在）
        print(f"\n⚠️ 测试图片 {test_image_path} 不存在，尝试创建...")
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            # 创建一个简单的测试图片
            img = Image.new('RGB', (400, 300), color='lightblue')
            draw = ImageDraw.Draw(img)
            
            # 添加文字
            try:
                font = ImageFont.truetype("arial.ttf", 24)
            except:
                font = ImageFont.load_default()
            
            draw.text((50, 50), "测试图片", fill='black', font=font)
            draw.text((50, 100), "DOCX MCP", fill='red', font=font)
            draw.text((50, 150), "图片插入测试", fill='green', font=font)
            
            img.save(test_image_path, "JPEG")
            print(f"✅ 成功创建测试图片: {test_image_path}")
            
        except ImportError:
            print("❌ 无法创建测试图片，PIL库未安装")
            print("请安装: pip install Pillow")
            return
        except Exception as e:
            print(f"❌ 创建测试图片失败: {e}")
            return
    
    # 测试图片插入
    print(f"\n3️⃣ 测试插入图片: {test_image_path}")
    result = processor.add_image(test_image_path, width="3in", height="2in", alignment="center")
    print(f"插入结果: {result}")
    
    # 检查文档中的图片
    print("\n4️⃣ 检查文档中的图片...")
    result = processor.list_images()
    print(f"图片列表: {result}")
    
    # 保存文档
    print("\n5️⃣ 保存文档...")
    result = processor.save_document()
    print(f"保存结果: {result}")
    
    # 另存为测试
    test_doc_path = "../docs/test_image_document.docx"
    result = processor.save_as_document(test_doc_path)
    print(f"另存为结果: {result}")
    
    # 验证文件是否存在
    if os.path.exists(test_doc_path):
        file_size = os.path.getsize(test_doc_path)
        print(f"✅ 文档已保存: {test_doc_path} (大小: {file_size} 字节)")
    else:
        print(f"❌ 文档保存失败: {test_doc_path}")
    
    # 测试图片大小调整
    print("\n6️⃣ 测试图片大小调整...")
    result = processor.resize_image(0, width="4in", height="3in")
    print(f"调整结果: {result}")
    
    # 再次保存
    processor.save_as_document("../docs/test_image_document_resized.docx")
    
    print("\n🎯 测试完成！")
    print(f"请检查以下文件:")
    print(f"- {test_doc_path}")
    print(f"- ../docs/test_image_document_resized.docx")
    print(f"- {test_image_path}")

def diagnose_image_issues():
    """诊断图片功能问题"""
    print("\n🔧 图片功能问题诊断...")
    
    # 检查python-docx版本
    try:
        import docx
        print(f"✅ python-docx版本: {docx.__version__}")
    except ImportError:
        print("❌ python-docx未安装")
        return
    
    # 检查PIL/Pillow
    try:
        from PIL import Image
        print(f"✅ Pillow版本: {Image.__version__}")
    except ImportError:
        print("❌ Pillow未安装，图片处理可能受限")
    
    # 检查lxml
    try:
        import lxml
        print(f"✅ lxml版本: {lxml.__version__}")
    except ImportError:
        print("❌ lxml未安装，XML操作可能受限")
    
    # 检查支持的图片格式
    print("\n📋 支持的图片格式:")
    supported_formats = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']
    for fmt in supported_formats:
        print(f"  - {fmt}")
    
    print("\n💡 常见问题解决方案:")
    print("1. 图片文件不存在或路径错误")
    print("2. 图片格式不支持")
    print("3. 图片文件损坏")
    print("4. 权限问题")
    print("5. python-docx版本兼容性问题")

def test_basic_image_operations():
    """测试基本图片操作"""
    print("\n🧪 测试基本图片操作...")
    
    processor = EnhancedDocxProcessor()
    
    # 创建文档
    processor.create_document("basic_image_test.docx")
    processor.add_heading("基本图片操作测试", 1)
    
    # 测试不同格式的图片插入
    test_images = [
        ("../images/test_image.jpg", "JPEG格式"),
        ("../images/test_image.png", "PNG格式"),
        ("../images/test_image.bmp", "BMP格式")
    ]
    
    for img_path, description in test_images:
        if os.path.exists(img_path):
            print(f"\n插入 {description}: {img_path}")
            result = processor.add_image(img_path, width="2in", height="1.5in")
            print(f"结果: {result}")
        else:
            print(f"\n⚠️ {description}文件不存在: {img_path}")
    
    # 保存文档
    processor.save_as_document("../docs/basic_image_test.docx")
    print("\n✅ 基本图片操作测试完成")

if __name__ == "__main__":
    print("🚀 DOCX图片功能测试开始")
    print("=" * 50)
    
    # 诊断问题
    diagnose_image_issues()
    
    # 测试基本功能
    test_basic_image_operations()
    
    # 测试完整功能
    test_image_functionality()
    
    print("\n" + "=" * 50)
    print("🎉 所有测试完成！")
