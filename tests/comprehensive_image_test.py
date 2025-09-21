#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
全面图片功能测试脚本
深入测试DOCX文档中的图片功能，分析可能的问题
"""

import os
import sys
import logging
from pathlib import Path

# 添加项目根目录到Python路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from core.enhanced_docx_processor import EnhancedDocxProcessor

# 设置日志
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_test_images():
    """创建多种格式的测试图片"""
    print("🎨 创建测试图片...")
    
    try:
        from PIL import Image, ImageDraw, ImageFont
        
        # 创建JPEG测试图片
        img_jpg = Image.new('RGB', (400, 300), color='lightblue')
        draw = ImageDraw.Draw(img_jpg)
        draw.text((50, 50), "JPEG测试图片", fill='black')
        draw.text((50, 100), "DOCX MCP", fill='red')
        img_jpg.save("../images/test_image.jpg", "JPEG")
        print("✅ 创建JPEG测试图片成功")
        
        # 创建PNG测试图片
        img_png = Image.new('RGBA', (400, 300), color='lightgreen')
        draw = ImageDraw.Draw(img_png)
        draw.text((50, 50), "PNG测试图片", fill='black')
        draw.text((50, 100), "透明背景", fill='blue')
        img_png.save("../images/test_image.png", "PNG")
        print("✅ 创建PNG测试图片成功")
        
        # 创建BMP测试图片
        img_bmp = Image.new('RGB', (400, 300), color='lightyellow')
        draw = ImageDraw.Draw(img_bmp)
        draw.text((50, 50), "BMP测试图片", fill='black')
        draw.text((50, 100), "无压缩格式", fill='purple')
        img_bmp.save("../images/test_image.bmp", "BMP")
        print("✅ 创建BMP测试图片成功")
        
    except ImportError:
        print("❌ PIL库未安装，无法创建测试图片")
        return False
    except Exception as e:
        print(f"❌ 创建测试图片失败: {e}")
        return False
    
    return True

def test_image_insertion_with_details():
    """详细测试图片插入功能"""
    print("\n🔍 详细测试图片插入功能...")
    
    processor = EnhancedDocxProcessor()
    
    # 创建新文档
    doc_path = "../docs/detailed_image_test.docx"
    processor.create_document(doc_path)
    
    # 添加标题
    processor.add_heading("详细图片功能测试", 1)
    processor.add_paragraph("这个文档将测试各种图片插入场景")
    
    # 测试1: 基本图片插入
    print("\n📸 测试1: 基本图片插入")
    if os.path.exists("../images/test_image.jpg"):
        result = processor.add_image("../images/test_image.jpg")
        print(f"JPEG插入结果: {result}")
    else:
        print("⚠️ ../images/test_image.jpg 不存在")
    
    # 测试2: 指定尺寸的图片插入
    print("\n📏 测试2: 指定尺寸的图片插入")
    if os.path.exists("../images/test_image.png"):
        result = processor.add_image("../images/test_image.png", width="2.5in", height="2in", alignment="center")
        print(f"PNG插入结果: {result}")
    else:
        print("⚠️ ../images/test_image.png 不存在")
    
    # 测试3: 在指定位置插入图片
    print("\n📍 测试3: 在指定位置插入图片")
    if os.path.exists("../images/test_image.bmp"):
        result = processor.add_image("../images/test_image.bmp", width="3in", height="2.5in", paragraph_index=2)
        print(f"BMP插入结果: {result}")
    else:
        print("⚠️ ../images/test_image.bmp 不存在")
    
    # 检查文档中的图片
    print("\n🔍 检查文档中的图片...")
    result = processor.list_images()
    print(f"图片列表: {result}")
    
    # 保存文档
    processor.save_document()
    print(f"✅ 文档已保存: {doc_path}")
    
    return doc_path

def analyze_document_structure(doc_path):
    """分析文档结构，检查图片是否正确插入"""
    print(f"\n🔬 分析文档结构: {doc_path}")
    
    try:
        from docx import Document
        doc = Document(doc_path)
        
        print(f"文档段落数: {len(doc.paragraphs)}")
        print(f"文档表格数: {len(doc.tables)}")
        
        # 检查每个段落
        for i, para in enumerate(doc.paragraphs):
            print(f"\n段落 {i}: '{para.text[:50]}...'")
            print(f"  运行数: {len(para.runs)}")
            
            for j, run in enumerate(para.runs):
                print(f"    运行 {j}: '{run.text[:30]}...'")
                
                # 检查是否包含图片
                if run._element.xpath('.//a:blip'):
                    print(f"      ✅ 包含图片")
                    
                    # 获取图片信息
                    try:
                        drawing = run._element.xpath('.//w:drawing')[0]
                        extent = drawing.xpath('.//wp:extent')[0]
                        width_emu = int(extent.get('cx'))
                        height_emu = int(extent.get('cy'))
                        width_inches = width_emu / 914400
                        height_inches = height_emu / 914400
                        print(f"      图片尺寸: {width_inches:.2f}\" x {height_inches:.2f}\"")
                    except Exception as e:
                        print(f"      无法获取图片尺寸: {e}")
                else:
                    print(f"      ❌ 不包含图片")
        
        # 检查文档的XML结构
        print(f"\n🔍 检查XML结构...")
        root = doc._element
        images = root.xpath('.//a:blip')
        print(f"XML中找到的图片数量: {len(images)}")
        
        if images:
            for i, img in enumerate(images):
                print(f"  图片 {i}: {img}")
        
    except Exception as e:
        print(f"❌ 分析文档失败: {e}")

def test_image_operations():
    """测试图片操作功能"""
    print("\n🔄 测试图片操作功能...")
    
    processor = EnhancedDocxProcessor()
    
    # 打开刚才创建的文档
    doc_path = "../docs/detailed_image_test.docx"
    if not os.path.exists(doc_path):
        print(f"❌ 文档不存在: {doc_path}")
        return
    
    processor.open_document(doc_path)
    
    # 测试图片大小调整
    print("\n📏 测试图片大小调整...")
    result = processor.resize_image(0, width="4in", height="3in")
    print(f"调整结果: {result}")
    
    # 测试图片位置设置
    print("\n📍 测试图片位置设置...")
    result = processor.set_image_position(0, position_type="inline", horizontal_position="center")
    print(f"位置设置结果: {result}")
    
    # 再次检查图片
    result = processor.list_images()
    print(f"调整后的图片列表: {result}")
    
    # 保存修改后的文档
    processor.save_as_document("../docs/detailed_image_test_modified.docx")
    print("✅ 修改后的文档已保存")

def troubleshoot_common_issues():
    """排查常见问题"""
    print("\n🔧 排查常见问题...")
    
    # 检查文件权限
    print("\n1️⃣ 检查文件权限...")
    test_files = ["../images/test_image.jpg", "../images/test_image.png", "../images/test_image.bmp"]
    for file in test_files:
        if os.path.exists(file):
            try:
                with open(file, 'rb') as f:
                    f.read(1)
                print(f"✅ {file}: 可读")
            except Exception as e:
                print(f"❌ {file}: 不可读 - {e}")
        else:
            print(f"⚠️ {file}: 不存在")
    
    # 检查python-docx版本兼容性
    print("\n2️⃣ 检查python-docx版本兼容性...")
    try:
        import docx
        version = docx.__version__
        print(f"当前版本: {version}")
        
        if version.startswith('0.8'):
            print("⚠️ 版本0.8.x可能存在兼容性问题")
        elif version.startswith('1.0'):
            print("✅ 版本1.0.x应该兼容")
        else:
            print(f"⚠️ 未知版本: {version}")
    except Exception as e:
        print(f"❌ 无法获取版本信息: {e}")
    
    # 检查依赖库
    print("\n3️⃣ 检查依赖库...")
    dependencies = ['lxml', 'PIL', 'docx']
    for dep in dependencies:
        try:
            if dep == 'PIL':
                from PIL import Image
                print(f"✅ {dep}: 已安装")
            elif dep == 'docx':
                import docx
                print(f"✅ {dep}: 已安装")
            else:
                __import__(dep)
                print(f"✅ {dep}: 已安装")
        except ImportError:
            print(f"❌ {dep}: 未安装")

def main():
    """主函数"""
    print("🚀 全面图片功能测试开始")
    print("=" * 60)
    
    # 创建测试图片
    if not create_test_images():
        print("❌ 无法创建测试图片，测试终止")
        return
    
    # 测试图片插入
    doc_path = test_image_insertion_with_details()
    
    # 分析文档结构
    analyze_document_structure(doc_path)
    
    # 测试图片操作
    test_image_operations()
    
    # 排查问题
    troubleshoot_common_issues()
    
    print("\n" + "=" * 60)
    print("🎯 测试完成！")
    print("\n📋 生成的文件:")
    print("- ../docs/detailed_image_test.docx (原始测试文档)")
    print("- ../docs/detailed_image_test_modified.docx (修改后的文档)")
    print("- ../images/test_image.jpg, ../images/test_image.png, ../images/test_image.bmp (测试图片)")
    
    print("\n💡 如果看不到图片，可能的原因:")
    print("1. 图片文件路径错误或不存在")
    print("2. 图片格式不支持")
    print("3. python-docx版本兼容性问题")
    print("4. 文档保存时出现问题")
    print("5. 查看文档的软件不支持该图片格式")

if __name__ == "__main__":
    main()
