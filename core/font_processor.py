#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
字体处理模块
负责文档中文本的字体样式、大小、颜色等格式处理
"""

import logging
from typing import Optional, List, Union, Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

class FontProcessor:
    """字体处理器，提供文本格式化的所有操作"""
    
    @staticmethod
    def set_paragraph_font(
        document: Document,
        paragraph_index: int,
        font_name: Optional[str] = None,
        font_size: Optional[int] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        underline: Optional[bool] = None,
        color: Optional[str] = None,
        alignment: Optional[str] = None
    ) -> str:
        """
        设置段落字体样式
        
        Parameters:
        - document: docx文档对象
        - paragraph_index: 段落索引
        - font_name: 字体名称
        - font_size: 字体大小（磅）
        - bold: 是否加粗
        - italic: 是否斜体
        - underline: 是否下划线
        - color: 字体颜色（十六进制，如#FF0000）
        - alignment: 对齐方式（left/center/right/justify）
        
        Returns:
        - 操作结果消息
        """
        try:
            if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
                return f"段落索引超出范围: {paragraph_index}，文档共有 {len(document.paragraphs)} 个段落"
            
            paragraph = document.paragraphs[paragraph_index]
            
            # 设置段落对齐
            if alignment:
                alignment_map = {
                    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                }
                if alignment in alignment_map:
                    paragraph.alignment = alignment_map[alignment]
            
            # 如果段落没有runs，创建一个
            if not paragraph.runs:
                run = paragraph.add_run(paragraph.text)
                paragraph.clear()
            
            # 对所有runs应用格式
            for run in paragraph.runs:
                FontProcessor._apply_run_formatting(
                    run, font_name, font_size, bold, italic, underline, color
                )
            
            return f"段落 {paragraph_index} 字体样式已更新"
            
        except Exception as e:
            logger.error(f"设置段落字体失败: {e}")
            return f"设置段落字体失败: {str(e)}"
    
    @staticmethod
    def set_text_range_font(
        document: Document,
        paragraph_index: int,
        start_pos: int,
        end_pos: int,
        font_name: Optional[str] = None,
        font_size: Optional[int] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        underline: Optional[bool] = None,
        color: Optional[str] = None
    ) -> str:
        """
        设置文本范围的字体样式
        
        Parameters:
        - document: docx文档对象
        - paragraph_index: 段落索引
        - start_pos: 开始位置
        - end_pos: 结束位置
        - 其他参数同set_paragraph_font
        
        Returns:
        - 操作结果消息
        """
        try:
            if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
                return f"段落索引超出范围: {paragraph_index}"
            
            paragraph = document.paragraphs[paragraph_index]
            text = paragraph.text
            
            if start_pos < 0 or start_pos >= len(text):
                return f"开始位置超出范围: {start_pos}"
            
            if end_pos <= start_pos or end_pos > len(text):
                return f"结束位置无效: {end_pos}"
            
            # 分割文本并重新构建runs
            before_text = text[:start_pos]
            target_text = text[start_pos:end_pos]
            after_text = text[end_pos:]
            
            # 保存原始格式
            original_format = FontProcessor._get_run_format(paragraph.runs[0]) if paragraph.runs else {}
            
            # 清除所有runs并重新添加
            paragraph.clear()
            
            # 添加前置文本
            if before_text:
                run = paragraph.add_run(before_text)
                FontProcessor._apply_format_dict(run, original_format)
            
            # 添加目标文本（应用新格式）
            if target_text:
                run = paragraph.add_run(target_text)
                FontProcessor._apply_format_dict(run, original_format)
                FontProcessor._apply_run_formatting(
                    run, font_name, font_size, bold, italic, underline, color
                )
            
            # 添加后置文本
            if after_text:
                run = paragraph.add_run(after_text)
                FontProcessor._apply_format_dict(run, original_format)
            
            return f"段落 {paragraph_index} 位置 {start_pos}-{end_pos} 字体样式已更新"
            
        except Exception as e:
            logger.error(f"设置文本范围字体失败: {e}")
            return f"设置文本范围字体失败: {str(e)}"
    
    @staticmethod
    def set_heading_style(
        document: Document,
        paragraph_index: int,
        level: int,
        font_name: Optional[str] = None,
        font_size: Optional[int] = None,
        color: Optional[str] = None
    ) -> str:
        """
        设置标题样式
        
        Parameters:
        - document: docx文档对象
        - paragraph_index: 段落索引
        - level: 标题级别（1-9）
        - font_name: 字体名称
        - font_size: 字体大小
        - color: 字体颜色
        
        Returns:
        - 操作结果消息
        """
        try:
            if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
                return f"段落索引超出范围: {paragraph_index}"
            
            if level < 1 or level > 9:
                return f"标题级别必须在1-9之间: {level}"
            
            paragraph = document.paragraphs[paragraph_index]
            
            # 设置标题样式
            paragraph.style = document.styles[f'Heading {level}']
            
            # 应用自定义格式
            if paragraph.runs:
                for run in paragraph.runs:
                    FontProcessor._apply_run_formatting(
                        run, font_name, font_size, None, None, None, color
                    )
            
            return f"段落 {paragraph_index} 已设置为 {level} 级标题"
            
        except Exception as e:
            logger.error(f"设置标题样式失败: {e}")
            return f"设置标题样式失败: {str(e)}"
    
    @staticmethod
    def apply_text_style(
        document: Document,
        style_name: str,
        paragraph_indices: List[int]
    ) -> str:
        """
        批量应用文本样式
        
        Parameters:
        - document: docx文档对象
        - style_name: 样式名称
        - paragraph_indices: 段落索引列表
        
        Returns:
        - 操作结果消息
        """
        try:
            # 检查样式是否存在
            available_styles = [style.name for style in document.styles]
            if style_name not in available_styles:
                return f"样式不存在: {style_name}，可用样式: {', '.join(available_styles[:10])}"
            
            applied_count = 0
            errors = []
            
            for idx in paragraph_indices:
                try:
                    if 0 <= idx < len(document.paragraphs):
                        document.paragraphs[idx].style = document.styles[style_name]
                        applied_count += 1
                    else:
                        errors.append(f"索引 {idx} 超出范围")
                except Exception as e:
                    errors.append(f"索引 {idx}: {str(e)}")
            
            result = f"成功应用样式到 {applied_count} 个段落"
            if errors:
                result += f"，错误: {'; '.join(errors)}"
            
            return result
            
        except Exception as e:
            logger.error(f"批量应用样式失败: {e}")
            return f"批量应用样式失败: {str(e)}"
    
    @staticmethod
    def create_custom_style(
        document: Document,
        style_name: str,
        font_name: str = "微软雅黑",
        font_size: int = 12,
        bold: bool = False,
        italic: bool = False,
        color: str = "#000000",
        alignment: str = "left"
    ) -> str:
        """
        创建自定义样式
        
        Parameters:
        - document: docx文档对象
        - style_name: 样式名称
        - font_name: 字体名称
        - font_size: 字体大小
        - bold: 是否加粗
        - italic: 是否斜体
        - color: 字体颜色
        - alignment: 对齐方式
        
        Returns:
        - 操作结果消息
        """
        try:
            # 检查样式是否已存在
            existing_styles = [style.name for style in document.styles]
            if style_name in existing_styles:
                return f"样式已存在: {style_name}"
            
            # 创建新样式
            from docx.enum.style import WD_STYLE_TYPE
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            
            # 设置字体
            font = style.font
            font.name = font_name
            font.size = Pt(font_size)
            font.bold = bold
            font.italic = italic
            
            # 设置颜色
            if color and color.startswith('#'):
                try:
                    r = int(color[1:3], 16)
                    g = int(color[3:5], 16)
                    b = int(color[5:7], 16)
                    font.color.rgb = RGBColor(r, g, b)
                except ValueError:
                    pass
            
            # 设置对齐
            alignment_map = {
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            if alignment in alignment_map:
                style.paragraph_format.alignment = alignment_map[alignment]
            
            # 设置中文字体
            style.font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            
            return f"自定义样式 '{style_name}' 创建成功"
            
        except Exception as e:
            logger.error(f"创建自定义样式失败: {e}")
            return f"创建自定义样式失败: {str(e)}"
    
    @staticmethod
    def get_font_info(document: Document, paragraph_index: int) -> str:
        """
        获取段落字体信息
        
        Parameters:
        - document: docx文档对象
        - paragraph_index: 段落索引
        
        Returns:
        - 字体信息字符串
        """
        try:
            if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
                return f"段落索引超出范围: {paragraph_index}"
            
            paragraph = document.paragraphs[paragraph_index]
            
            if not paragraph.runs:
                return f"段落 {paragraph_index} 没有文本内容"
            
            # 获取第一个run的格式信息
            run = paragraph.runs[0]
            font = run.font
            
            info = [
                f"段落 {paragraph_index} 字体信息:",
                f"  字体名称: {font.name or '默认'}",
                f"  字体大小: {font.size.pt if font.size else '默认'}磅",
                f"  加粗: {'是' if font.bold else '否'}",
                f"  斜体: {'是' if font.italic else '否'}",
                f"  下划线: {'是' if font.underline else '否'}",
                f"  颜色: {font.color.rgb if font.color.rgb else '默认'}",
                f"  样式: {paragraph.style.name}",
                f"  对齐: {paragraph.alignment}"
            ]
            
            return "\n".join(info)
            
        except Exception as e:
            logger.error(f"获取字体信息失败: {e}")
            return f"获取字体信息失败: {str(e)}"
    
    @staticmethod
    def _apply_run_formatting(
        run,
        font_name: Optional[str],
        font_size: Optional[int],
        bold: Optional[bool],
        italic: Optional[bool],
        underline: Optional[bool],
        color: Optional[str]
    ):
        """应用run格式"""
        font = run.font
        
        if font_name:
            font.name = font_name
            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        if font_size:
            font.size = Pt(font_size)
        
        if bold is not None:
            font.bold = bold
        
        if italic is not None:
            font.italic = italic
        
        if underline is not None:
            font.underline = underline
        
        if color and color.startswith('#') and len(color) == 7:
            try:
                r = int(color[1:3], 16)
                g = int(color[3:5], 16)
                b = int(color[5:7], 16)
                font.color.rgb = RGBColor(r, g, b)
            except ValueError:
                pass
    
    @staticmethod
    def _get_run_format(run) -> Dict[str, Any]:
        """获取run的格式信息"""
        font = run.font
        return {
            "font_name": font.name,
            "font_size": font.size.pt if font.size else None,
            "bold": font.bold,
            "italic": font.italic,
            "underline": font.underline,
            "color": font.color.rgb if font.color.rgb else None
        }
    
    @staticmethod
    def _apply_format_dict(run, format_dict: Dict[str, Any]):
        """应用格式字典到run"""
        FontProcessor._apply_run_formatting(
            run,
            format_dict.get("font_name"),
            format_dict.get("font_size"),
            format_dict.get("bold"),
            format_dict.get("italic"),
            format_dict.get("underline"),
            format_dict.get("color")
        )
