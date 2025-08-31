#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
增强版状态管理器
提供完整的操作历史记录、状态持久化和智能恢复功能
"""

import os
import json
import pickle
import logging
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass, field, asdict
from datetime import datetime, timedelta
from pathlib import Path
from enum import Enum
import hashlib

logger = logging.getLogger(__name__)

# ==================== 枚举定义 ====================

class OperationType(Enum):
    """操作类型"""
    CREATE_DOCUMENT = "create_document"
    OPEN_DOCUMENT = "open_document"
    SAVE_DOCUMENT = "save_document"
    ADD_PARAGRAPH = "add_paragraph"
    ADD_HEADING = "add_heading"
    ADD_TABLE = "add_table"
    ADD_IMAGE = "add_image"
    FORMAT_TEXT = "format_text"
    APPLY_TEMPLATE = "apply_template"
    WORKFLOW_EXECUTION = "workflow_execution"
    BATCH_OPERATION = "batch_operation"

class OperationStatus(Enum):
    """操作状态"""
    PENDING = "pending"
    RUNNING = "running"
    COMPLETED = "completed"
    FAILED = "failed"
    CANCELLED = "cancelled"

# ==================== 数据结构定义 ====================

@dataclass
class OperationRecord:
    """操作记录"""
    operation_id: str
    operation_type: OperationType
    timestamp: datetime
    status: OperationStatus
    parameters: Dict[str, Any]
    result: Optional[Any] = None
    error_message: Optional[str] = None
    execution_time: float = 0.0
    context: Dict[str, Any] = field(default_factory=dict)
    dependencies: List[str] = field(default_factory=list)
    rollback_data: Optional[Dict[str, Any]] = None

@dataclass
class DocumentState:
    """文档状态"""
    document_id: str
    file_path: str
    last_modified: datetime
    content_hash: str
    structure_info: Dict[str, Any]
    metadata: Dict[str, Any] = field(default_factory=dict)
    operation_history: List[str] = field(default_factory=list)  # 操作ID列表

@dataclass
class SessionState:
    """会话状态"""
    session_id: str
    start_time: datetime
    last_activity: datetime
    current_document: Optional[str] = None
    active_operations: List[str] = field(default_factory=list)
    completed_operations: List[str] = field(default_factory=list)
    failed_operations: List[str] = field(default_factory=list)
    context: Dict[str, Any] = field(default_factory=dict)

@dataclass
class StateSnapshot:
    """状态快照"""
    snapshot_id: str
    timestamp: datetime
    session_state: SessionState
    document_states: Dict[str, DocumentState]
    operation_records: Dict[str, OperationRecord]
    checksum: str

# ==================== 增强版状态管理器 ====================

class EnhancedStateManager:
    """增强版状态管理器 - 提供完整的操作历史和状态持久化"""
    
    def __init__(self, state_dir: str = "state_data"):
        self.state_dir = Path(state_dir)
        self.state_dir.mkdir(exist_ok=True)
        
        # 状态存储
        self.current_session: Optional[SessionState] = None
        self.document_states: Dict[str, DocumentState] = {}
        self.operation_records: Dict[str, OperationRecord] = {}
        self.state_snapshots: Dict[str, StateSnapshot] = {}
        
        # 配置
        self.max_history_size = 1000  # 最大历史记录数
        self.snapshot_interval = 300  # 快照间隔（秒）
        self.auto_save_interval = 60  # 自动保存间隔（秒）
        
        # 加载之前的状态
        self._load_persisted_state()
        
        # 创建新会话
        self._create_new_session()
        
        logger.info("增强版状态管理器初始化完成")
    
    def _load_persisted_state(self):
        """加载持久化的状态"""
        try:
            # 加载会话状态
            session_file = self.state_dir / "current_session.json"
            if session_file.exists():
                with open(session_file, 'r', encoding='utf-8') as f:
                    session_data = json.load(f)
                    self.current_session = self._deserialize_session(session_data)
            
            # 加载文档状态
            docs_file = self.state_dir / "document_states.json"
            if docs_file.exists():
                with open(docs_file, 'r', encoding='utf-8') as f:
                    docs_data = json.load(f)
                    for doc_id, doc_data in docs_data.items():
                        self.document_states[doc_id] = self._deserialize_document_state(doc_data)
            
            # 加载操作记录
            ops_file = self.state_dir / "operation_records.json"
            if ops_file.exists():
                with open(ops_file, 'r', encoding='utf-8') as f:
                    ops_data = json.load(f)
                    for op_id, op_data in ops_data.items():
                        self.operation_records[op_id] = self._deserialize_operation_record(op_data)
            
            # 加载状态快照
            snapshots_dir = self.state_dir / "snapshots"
            if snapshots_dir.exists():
                for snapshot_file in snapshots_dir.glob("*.json"):
                    with open(snapshot_file, 'r', encoding='utf-8') as f:
                        snapshot_data = json.load(f)
                        snapshot = self._deserialize_state_snapshot(snapshot_data)
                        self.state_snapshots[snapshot.snapshot_id] = snapshot
            
            logger.info(f"加载了 {len(self.document_states)} 个文档状态和 {len(self.operation_records)} 个操作记录")
            
        except Exception as e:
            logger.error(f"加载持久化状态失败: {e}")
    
    def _create_new_session(self):
        """创建新会话"""
        if not self.current_session:
            session_id = f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            self.current_session = SessionState(
                session_id=session_id,
                start_time=datetime.now(),
                last_activity=datetime.now()
            )
            logger.info(f"创建新会话: {session_id}")
    
    def record_operation(self, operation_type: OperationType, parameters: Dict[str, Any], 
                        context: Dict[str, Any] = None) -> str:
        """记录操作"""
        operation_id = f"{operation_type.value}_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}"
        
        operation = OperationRecord(
            operation_id=operation_id,
            operation_type=operation_type,
            timestamp=datetime.now(),
            status=OperationStatus.PENDING,
            parameters=parameters,
            context=context or {}
        )
        
        self.operation_records[operation_id] = operation
        self.current_session.active_operations.append(operation_id)
        self.current_session.last_activity = datetime.now()
        
        logger.info(f"记录操作: {operation_id} - {operation_type.value}")
        return operation_id
    
    def update_operation_status(self, operation_id: str, status: OperationStatus, 
                               result: Any = None, error_message: str = None, 
                               execution_time: float = 0.0):
        """更新操作状态"""
        if operation_id not in self.operation_records:
            logger.warning(f"操作记录不存在: {operation_id}")
            return
        
        operation = self.operation_records[operation_id]
        operation.status = status
        operation.result = result
        operation.error_message = error_message
        operation.execution_time = execution_time
        
        # 更新会话状态
        if status == OperationStatus.COMPLETED:
            if operation_id in self.current_session.active_operations:
                self.current_session.active_operations.remove(operation_id)
            self.current_session.completed_operations.append(operation_id)
        elif status == OperationStatus.FAILED:
            if operation_id in self.current_session.active_operations:
                self.current_session.active_operations.remove(operation_id)
            self.current_session.failed_operations.append(operation_id)
        
        self.current_session.last_activity = datetime.now()
        
        logger.info(f"更新操作状态: {operation_id} - {status.value}")
    
    def set_current_document(self, file_path: str, document_id: str = None):
        """设置当前文档"""
        if not document_id:
            document_id = self._generate_document_id(file_path)
        
        # 创建或更新文档状态
        if document_id not in self.document_states:
            self.document_states[document_id] = DocumentState(
                document_id=document_id,
                file_path=file_path,
                last_modified=datetime.now(),
                content_hash=self._calculate_file_hash(file_path),
                structure_info=self._analyze_document_structure(file_path)
            )
        else:
            # 更新现有文档状态
            doc_state = self.document_states[document_id]
            doc_state.file_path = file_path
            doc_state.last_modified = datetime.now()
            doc_state.content_hash = self._calculate_file_hash(file_path)
            doc_state.structure_info = self._analyze_document_structure(file_path)
        
        self.current_session.current_document = document_id
        self.current_session.last_activity = datetime.now()
        
        logger.info(f"设置当前文档: {document_id} - {file_path}")
    
    def add_operation_to_document(self, document_id: str, operation_id: str):
        """将操作添加到文档历史"""
        if document_id in self.document_states:
            self.document_states[document_id].operation_history.append(operation_id)
            # 限制历史记录大小
            if len(self.document_states[document_id].operation_history) > self.max_history_size:
                self.document_states[document_id].operation_history = \
                    self.document_states[document_id].operation_history[-self.max_history_size:]
    
    def get_operation_history(self, document_id: str = None, operation_type: OperationType = None, 
                             limit: int = 50) -> List[OperationRecord]:
        """获取操作历史"""
        operations = []
        
        if document_id:
            # 获取特定文档的操作历史
            if document_id in self.document_states:
                operation_ids = self.document_states[document_id].operation_history
                for op_id in operation_ids:
                    if op_id in self.operation_records:
                        operations.append(self.operation_records[op_id])
        else:
            # 获取所有操作
            operations = list(self.operation_records.values())
        
        # 按操作类型过滤
        if operation_type:
            operations = [op for op in operations if op.operation_type == operation_type]
        
        # 按时间排序（最新的在前）
        operations.sort(key=lambda x: x.timestamp, reverse=True)
        
        # 限制返回数量
        return operations[:limit]
    
    def get_document_state(self, document_id: str) -> Optional[DocumentState]:
        """获取文档状态"""
        return self.document_states.get(document_id)
    
    def get_current_session_info(self) -> Dict[str, Any]:
        """获取当前会话信息"""
        if not self.current_session:
            return {}
        
        return {
            "session_id": self.current_session.session_id,
            "start_time": self.current_session.start_time.isoformat(),
            "last_activity": self.current_session.last_activity.isoformat(),
            "current_document": self.current_session.current_document,
            "active_operations_count": len(self.current_session.active_operations),
            "completed_operations_count": len(self.current_session.completed_operations),
            "failed_operations_count": len(self.current_session.failed_operations),
            "total_documents": len(self.document_states),
            "total_operations": len(self.operation_records)
        }
    
    def create_state_snapshot(self, description: str = "") -> str:
        """创建状态快照"""
        snapshot_id = f"snapshot_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        snapshot = StateSnapshot(
            snapshot_id=snapshot_id,
            timestamp=datetime.now(),
            session_state=self.current_session,
            document_states=self.document_states.copy(),
            operation_records=self.operation_records.copy(),
            checksum=""
        )
        
        # 计算校验和
        snapshot.checksum = self._calculate_snapshot_checksum(snapshot)
        
        self.state_snapshots[snapshot_id] = snapshot
        
        # 保存快照到文件
        self._save_snapshot_to_file(snapshot)
        
        logger.info(f"创建状态快照: {snapshot_id}")
        return snapshot_id
    
    def restore_from_snapshot(self, snapshot_id: str) -> bool:
        """从快照恢复状态"""
        if snapshot_id not in self.state_snapshots:
            logger.error(f"快照不存在: {snapshot_id}")
            return False
        
        try:
            snapshot = self.state_snapshots[snapshot_id]
            
            # 验证校验和
            if not self._verify_snapshot_checksum(snapshot):
                logger.error(f"快照校验和验证失败: {snapshot_id}")
                return False
            
            # 恢复状态
            self.current_session = snapshot.session_state
            self.document_states = snapshot.document_states.copy()
            self.operation_records = snapshot.operation_records.copy()
            
            # 保存恢复的状态
            self._persist_state()
            
            logger.info(f"从快照恢复状态成功: {snapshot_id}")
            return True
            
        except Exception as e:
            logger.error(f"从快照恢复状态失败: {e}")
            return False
    
    def get_available_snapshots(self) -> List[Dict[str, Any]]:
        """获取可用的快照列表"""
        snapshots = []
        for snapshot in self.state_snapshots.values():
            snapshots.append({
                "snapshot_id": snapshot.snapshot_id,
                "timestamp": snapshot.timestamp.isoformat(),
                "checksum": snapshot.checksum,
                "documents_count": len(snapshot.document_states),
                "operations_count": len(snapshot.operation_records)
            })
        
        # 按时间排序（最新的在前）
        snapshots.sort(key=lambda x: x["timestamp"], reverse=True)
        return snapshots
    
    def cleanup_old_data(self, days_to_keep: int = 30):
        """清理旧数据"""
        cutoff_date = datetime.now() - timedelta(days=days_to_keep)
        
        # 清理旧的操作记录
        old_operations = [
            op_id for op_id, op in self.operation_records.items()
            if op.timestamp < cutoff_date
        ]
        for op_id in old_operations:
            del self.operation_records[op_id]
        
        # 清理旧的快照
        old_snapshots = [
            snap_id for snap_id, snap in self.state_snapshots.items()
            if snap.timestamp < cutoff_date
        ]
        for snap_id in old_snapshots:
            del self.state_snapshots[snap_id]
            # 删除快照文件
            snapshot_file = self.state_dir / "snapshots" / f"{snap_id}.json"
            if snapshot_file.exists():
                snapshot_file.unlink()
        
        logger.info(f"清理了 {len(old_operations)} 个旧操作记录和 {len(old_snapshots)} 个旧快照")
    
    def _generate_document_id(self, file_path: str) -> str:
        """生成文档ID"""
        return hashlib.md5(file_path.encode()).hexdigest()[:16]
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """计算文件哈希值"""
        if not os.path.exists(file_path):
            return ""
        
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                return hashlib.md5(content).hexdigest()
        except Exception:
            return ""
    
    def _analyze_document_structure(self, file_path: str) -> Dict[str, Any]:
        """分析文档结构"""
        if not os.path.exists(file_path):
            return {}
        
        try:
            from docx import Document
            doc = Document(file_path)
            
            structure = {
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
                "images_count": len(doc.inline_shapes),
                "sections_count": len(doc.sections),
                "last_modified": os.path.getmtime(file_path)
            }
            
            return structure
        except Exception as e:
            logger.error(f"分析文档结构失败: {e}")
            return {}
    
    def _calculate_snapshot_checksum(self, snapshot: StateSnapshot) -> str:
        """计算快照校验和"""
        data = {
            "session_state": asdict(snapshot.session_state),
            "document_states": {k: asdict(v) for k, v in snapshot.document_states.items()},
            "operation_records": {k: asdict(v) for k, v in snapshot.operation_records.items()}
        }
        
        data_str = json.dumps(data, sort_keys=True, default=str)
        return hashlib.md5(data_str.encode()).hexdigest()
    
    def _verify_snapshot_checksum(self, snapshot: StateSnapshot) -> bool:
        """验证快照校验和"""
        calculated_checksum = self._calculate_snapshot_checksum(snapshot)
        return calculated_checksum == snapshot.checksum
    
    def _save_snapshot_to_file(self, snapshot: StateSnapshot):
        """保存快照到文件"""
        snapshots_dir = self.state_dir / "snapshots"
        snapshots_dir.mkdir(exist_ok=True)
        
        snapshot_file = snapshots_dir / f"{snapshot.snapshot_id}.json"
        
        snapshot_data = {
            "snapshot_id": snapshot.snapshot_id,
            "timestamp": snapshot.timestamp.isoformat(),
            "session_state": self._serialize_session(snapshot.session_state),
            "document_states": {k: self._serialize_document_state(v) for k, v in snapshot.document_states.items()},
            "operation_records": {k: self._serialize_operation_record(v) for k, v in snapshot.operation_records.items()},
            "checksum": snapshot.checksum
        }
        
        with open(snapshot_file, 'w', encoding='utf-8') as f:
            json.dump(snapshot_data, f, ensure_ascii=False, indent=2)
    
    def _serialize_session(self, session: SessionState) -> Dict[str, Any]:
        """序列化会话状态"""
        return {
            "session_id": session.session_id,
            "start_time": session.start_time.isoformat(),
            "last_activity": session.last_activity.isoformat(),
            "current_document": session.current_document,
            "active_operations": session.active_operations,
            "completed_operations": session.completed_operations,
            "failed_operations": session.failed_operations,
            "context": session.context
        }
    
    def _deserialize_session(self, data: Dict[str, Any]) -> SessionState:
        """反序列化会话状态"""
        return SessionState(
            session_id=data["session_id"],
            start_time=datetime.fromisoformat(data["start_time"]),
            last_activity=datetime.fromisoformat(data["last_activity"]),
            current_document=data.get("current_document"),
            active_operations=data.get("active_operations", []),
            completed_operations=data.get("completed_operations", []),
            failed_operations=data.get("failed_operations", []),
            context=data.get("context", {})
        )
    
    def _serialize_document_state(self, doc_state: DocumentState) -> Dict[str, Any]:
        """序列化文档状态"""
        return {
            "document_id": doc_state.document_id,
            "file_path": doc_state.file_path,
            "last_modified": doc_state.last_modified.isoformat(),
            "content_hash": doc_state.content_hash,
            "structure_info": doc_state.structure_info,
            "metadata": doc_state.metadata,
            "operation_history": doc_state.operation_history
        }
    
    def _deserialize_document_state(self, data: Dict[str, Any]) -> DocumentState:
        """反序列化文档状态"""
        return DocumentState(
            document_id=data["document_id"],
            file_path=data["file_path"],
            last_modified=datetime.fromisoformat(data["last_modified"]),
            content_hash=data["content_hash"],
            structure_info=data.get("structure_info", {}),
            metadata=data.get("metadata", {}),
            operation_history=data.get("operation_history", [])
        )
    
    def _serialize_operation_record(self, op_record: OperationRecord) -> Dict[str, Any]:
        """序列化操作记录"""
        return {
            "operation_id": op_record.operation_id,
            "operation_type": op_record.operation_type.value,
            "timestamp": op_record.timestamp.isoformat(),
            "status": op_record.status.value,
            "parameters": op_record.parameters,
            "result": op_record.result,
            "error_message": op_record.error_message,
            "execution_time": op_record.execution_time,
            "context": op_record.context,
            "dependencies": op_record.dependencies,
            "rollback_data": op_record.rollback_data
        }
    
    def _deserialize_operation_record(self, data: Dict[str, Any]) -> OperationRecord:
        """反序列化操作记录"""
        return OperationRecord(
            operation_id=data["operation_id"],
            operation_type=OperationType(data["operation_type"]),
            timestamp=datetime.fromisoformat(data["timestamp"]),
            status=OperationStatus(data["status"]),
            parameters=data.get("parameters", {}),
            result=data.get("result"),
            error_message=data.get("error_message"),
            execution_time=data.get("execution_time", 0.0),
            context=data.get("context", {}),
            dependencies=data.get("dependencies", []),
            rollback_data=data.get("rollback_data")
        )
    
    def _deserialize_state_snapshot(self, data: Dict[str, Any]) -> StateSnapshot:
        """反序列化状态快照"""
        return StateSnapshot(
            snapshot_id=data["snapshot_id"],
            timestamp=datetime.fromisoformat(data["timestamp"]),
            session_state=self._deserialize_session(data["session_state"]),
            document_states={k: self._deserialize_document_state(v) for k, v in data["document_states"].items()},
            operation_records={k: self._deserialize_operation_record(v) for k, v in data["operation_records"].items()},
            checksum=data["checksum"]
        )
    
    def _persist_state(self):
        """持久化状态"""
        try:
            # 保存会话状态
            session_file = self.state_dir / "current_session.json"
            with open(session_file, 'w', encoding='utf-8') as f:
                json.dump(self._serialize_session(self.current_session), f, ensure_ascii=False, indent=2)
            
            # 保存文档状态
            docs_file = self.state_dir / "document_states.json"
            docs_data = {k: self._serialize_document_state(v) for k, v in self.document_states.items()}
            with open(docs_file, 'w', encoding='utf-8') as f:
                json.dump(docs_data, f, ensure_ascii=False, indent=2)
            
            # 保存操作记录
            ops_file = self.state_dir / "operation_records.json"
            ops_data = {k: self._serialize_operation_record(v) for k, v in self.operation_records.items()}
            with open(ops_file, 'w', encoding='utf-8') as f:
                json.dump(ops_data, f, ensure_ascii=False, indent=2)
            
            logger.debug("状态持久化完成")
            
        except Exception as e:
            logger.error(f"状态持久化失败: {e}")
    
    def __del__(self):
        """析构函数，确保状态被保存"""
        try:
            self._persist_state()
        except Exception:
            pass
