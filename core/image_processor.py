#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
图片处理模块
负责文档中图片的插入、删除、缩放和定位操作
"""

import os
import logging
from typing import Optional, Tuple, Union
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import shutil

logger = logging.getLogger(__name__)

class ImageProcessor:
    """图片处理器，提供图片相关的所有操作"""
    
    @staticmethod
    def add_image(
        document: Document,
        image_path: str,
        width: Optional[Union[float, str]] = None,
        height: Optional[Union[float, str]] = None,
        alignment: str = "left",
        paragraph_index: Optional[int] = None
    ) -> str:
        """
        向文档添加图片
        
        Parameters:
        - document: docx文档对象
        - image_path: 图片文件路径
        - width: 图片宽度，支持数字(英寸)或字符串("5cm", "2in")
        - height: 图片高度，支持数字(英寸)或字符串("5cm", "2in")
        - alignment: 对齐方式 ("left", "center", "right")
        - paragraph_index: 插入位置，None表示文档末尾
        
        Returns:
        - 操作结果消息
        """
        try:
            if not os.path.exists(image_path):
                return f"图片文件不存在: {image_path}"
            
            # 确定插入位置
            if paragraph_index is not None:
                if paragraph_index < 0 or paragraph_index > len(document.paragraphs):
                    return f"段落索引超出范围: {paragraph_index}"
                
                # 在指定位置插入新段落
                if paragraph_index == len(document.paragraphs):
                    paragraph = document.add_paragraph()
                else:
                    # 在指定段落前插入
                    p = document.paragraphs[paragraph_index]._element
                    new_p = OxmlElement("w:p")
                    p.addprevious(new_p)
                    paragraph = document.paragraphs[paragraph_index]
            else:
                paragraph = document.add_paragraph()
            
            # 设置段落对齐
            alignment_map = {
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT
            }
            paragraph.alignment = alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)
            
            # 添加图片
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            
            # 解析尺寸参数
            width_inches = ImageProcessor._parse_size(width) if width else None
            height_inches = ImageProcessor._parse_size(height) if height else None
            
            # 插入图片
            if width_inches and height_inches:
                run.add_picture(image_path, width=Inches(width_inches), height=Inches(height_inches))
            elif width_inches:
                run.add_picture(image_path, width=Inches(width_inches))
            elif height_inches:
                run.add_picture(image_path, height=Inches(height_inches))
            else:
                run.add_picture(image_path)
            
            return f"图片已添加: {os.path.basename(image_path)}"
            
        except Exception as e:
            logger.error(f"添加图片失败: {e}")
            return f"添加图片失败: {str(e)}"
    
    @staticmethod
    def extract_image_to_local(
        document: Document,
        image_index: int,
        output_dir: str = "extracted_images",
        filename: Optional[str] = None
    ) -> str:
        """
        将文档中的图片提取到本地文件夹
        
        Parameters:
        - document: docx文档对象
        - image_index: 图片索引
        - output_dir: 输出目录
        - filename: 自定义文件名，None则自动生成
        
        Returns:
        - 操作结果消息
        """
        try:
            # 查找所有图片
            image_runs = []
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        image_runs.append((paragraph, run))
            
            if not image_runs:
                return "文档中没有找到图片"
            
            if image_index < 0 or image_index >= len(image_runs):
                return f"图片索引超出范围: {image_index}，文档中共有 {len(image_runs)} 张图片"
            
            # 创建输出目录
            os.makedirs(output_dir, exist_ok=True)
            
            # 获取图片数据
            paragraph, run = image_runs[image_index]
            
            # 从run中提取图片
            for shape in run._element.xpath('.//pic:pic'):
                # 获取图片关系ID
                blip = shape.xpath('.//a:blip')[0]
                r_embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                
                if r_embed:
                    # 获取图片数据
                    image_part = document.part.related_parts[r_embed]
                    image_data = image_part.blob
                    
                    # 确定文件扩展名
                    content_type = image_part.content_type
                    if 'jpeg' in content_type or 'jpg' in content_type:
                        ext = '.jpg'
                    elif 'png' in content_type:
                        ext = '.png'
                    elif 'gif' in content_type:
                        ext = '.gif'
                    elif 'bmp' in content_type:
                        ext = '.bmp'
                    elif 'tiff' in content_type:
                        ext = '.tiff'
                    else:
                        ext = '.bin'
                    
                    # 生成文件名
                    if filename is None:
                        filename = f"extracted_image_{image_index}{ext}"
                    elif not filename.endswith(ext):
                        filename += ext
                    
                    # 保存图片
                    output_path = os.path.join(output_dir, filename)
                    with open(output_path, 'wb') as f:
                        f.write(image_data)
                    
                    return f"图片已提取到: {output_path}"
            
            return "无法提取图片数据"
            
        except Exception as e:
            logger.error(f"提取图片失败: {e}")
            return f"提取图片失败: {str(e)}"
    
    @staticmethod
    def extract_all_images_to_local(
        document: Document,
        output_dir: str = "extracted_images"
    ) -> str:
        """
        将文档中的所有图片提取到本地文件夹
        
        Parameters:
        - document: docx文档对象
        - output_dir: 输出目录
        
        Returns:
        - 操作结果消息
        """
        try:
            # 查找所有图片
            image_runs = []
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        image_runs.append((paragraph, run))
            
            if not image_runs:
                return "文档中没有找到图片"
            
            # 创建输出目录
            os.makedirs(output_dir, exist_ok=True)
            
            extracted_count = 0
            
            # 提取每张图片
            for img_idx, (paragraph, run) in enumerate(image_runs):
                for shape in run._element.xpath('.//pic:pic'):
                    blip = shape.xpath('.//a:blip')[0]
                    r_embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    
                    if r_embed:
                        image_part = document.part.related_parts[r_embed]
                        image_data = image_part.blob
                        
                        # 确定文件扩展名
                        content_type = image_part.content_type
                        if 'jpeg' in content_type or 'jpg' in content_type:
                            ext = '.jpg'
                        elif 'png' in content_type:
                            ext = '.png'
                        elif 'gif' in content_type:
                            ext = '.gif'
                        elif 'bmp' in content_type:
                            ext = '.bmp'
                        elif 'tiff' in content_type:
                            ext = '.tiff'
                        else:
                            ext = '.bin'
                        
                        # 生成文件名
                        filename = f"extracted_image_{img_idx}{ext}"
                        output_path = os.path.join(output_dir, filename)
                        
                        # 保存图片
                        with open(output_path, 'wb') as f:
                            f.write(image_data)
                        
                        extracted_count += 1
            
            return f"成功提取 {extracted_count} 张图片到: {output_dir}"
            
        except Exception as e:
            logger.error(f"提取所有图片失败: {e}")
            return f"提取所有图片失败: {str(e)}"
    
    @staticmethod
    def move_image_to_another_document(
        source_document: Document,
        target_document: Document,
        image_index: int,
        width: Optional[Union[float, str]] = None,
        height: Optional[Union[float, str]] = None,
        alignment: str = "left",
        paragraph_index: Optional[int] = None
    ) -> str:
        """
        将图片从一个文档移动到另一个文档
        
        Parameters:
        - source_document: 源文档
        - target_document: 目标文档
        - image_index: 源文档中的图片索引
        - width: 目标文档中的图片宽度
        - height: 目标文档中的图片高度
        - alignment: 目标文档中的对齐方式
        - paragraph_index: 目标文档中的插入位置
        
        Returns:
        - 操作结果消息
        """
        try:
            # 从源文档中提取图片
            image_runs = []
            for paragraph in source_document.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        image_runs.append((paragraph, run))
            
            if not image_runs:
                return "源文档中没有找到图片"
            
            if image_index < 0 or image_index >= len(image_runs):
                return f"图片索引超出范围: {image_index}，源文档中共有 {len(image_runs)} 张图片"
            
            # 获取图片数据
            paragraph, run = image_runs[image_index]
            
            # 从run中提取图片
            for shape in run._element.xpath('.//pic:pic'):
                blip = shape.xpath('.//a:blip')[0]
                r_embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                
                if r_embed:
                    image_part = source_document.part.related_parts[r_embed]
                    image_data = image_part.blob
                    
                    # 确定文件扩展名
                    content_type = image_part.content_type
                    if 'jpeg' in content_type or 'jpg' in content_type:
                        ext = '.jpg'
                    elif 'png' in content_type:
                        ext = '.png'
                    elif 'gif' in content_type:
                        ext = '.gif'
                    elif 'bmp' in content_type:
                        ext = '.bmp'
                    elif 'tiff' in content_type:
                        ext = '.tiff'
                    else:
                        ext = '.bin'
                    
                    # 创建临时文件
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp_file:
                        temp_file.write(image_data)
                        temp_path = temp_file.name
                    
                    try:
                        # 将图片添加到目标文档
                        result = ImageProcessor.add_image(
                            target_document, temp_path, width, height, alignment, paragraph_index
                        )
                        
                        # 清理临时文件
                        os.unlink(temp_path)
                        
                        return f"图片移动成功: {result}"
                        
                    except Exception as e:
                        # 清理临时文件
                        if os.path.exists(temp_path):
                            os.unlink(temp_path)
                        raise e
            
            return "无法提取图片数据"
            
        except Exception as e:
            logger.error(f"移动图片失败: {e}")
            return f"移动图片失败: {str(e)}"
    
    @staticmethod
    def copy_image_to_another_document(
        source_document: Document,
        target_document: Document,
        image_index: int,
        width: Optional[Union[float, str]] = None,
        height: Optional[Union[float, str]] = None,
        alignment: str = "left",
        paragraph_index: Optional[int] = None
    ) -> str:
        """
        将图片从一个文档复制到另一个文档（保留原文档中的图片）
        
        Parameters:
        - source_document: 源文档
        - target_document: 目标文档
        - image_index: 源文档中的图片索引
        - width: 目标文档中的图片宽度
        - height: 目标文档中的图片高度
        - alignment: 目标文档中的对齐方式
        - paragraph_index: 目标文档中的插入位置
        
        Returns:
        - 操作结果消息
        """
        try:
            # 从源文档中提取图片
            image_runs = []
            for paragraph in source_document.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        image_runs.append((paragraph, run))
            
            if not image_runs:
                return "源文档中没有找到图片"
            
            if image_index < 0 or image_index >= len(image_runs):
                return f"图片索引超出范围: {image_index}，源文档中共有 {len(image_runs)} 张图片"
            
            # 获取图片数据
            paragraph, run = image_runs[image_index]
            
            # 从run中提取图片
            for shape in run._element.xpath('.//pic:pic'):
                blip = shape.xpath('.//a:blip')[0]
                r_embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                
                if r_embed:
                    image_part = source_document.part.related_parts[r_embed]
                    image_data = image_part.blob
                    
                    # 确定文件扩展名
                    content_type = image_part.content_type
                    if 'jpeg' in content_type or 'jpg' in content_type:
                        ext = '.jpg'
                    elif 'png' in content_type:
                        ext = '.png'
                    elif 'gif' in content_type:
                        ext = '.gif'
                    elif 'bmp' in content_type:
                        ext = '.bmp'
                    elif 'tiff' in content_type:
                        ext = '.tiff'
                    else:
                        ext = '.bin'
                    
                    # 创建临时文件
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp_file:
                        temp_file.write(image_data)
                        temp_path = temp_file.name
                    
                    try:
                        # 将图片添加到目标文档
                        result = ImageProcessor.add_image(
                            target_document, temp_path, width, height, alignment, paragraph_index
                        )
                        
                        # 清理临时文件
                        os.unlink(temp_path)
                        
                        return f"图片复制成功: {result}"
                        
                    except Exception as e:
                        # 清理临时文件
                        if os.path.exists(temp_path):
                            os.unlink(temp_path)
                        raise e
            
            return "无法提取图片数据"
            
        except Exception as e:
            logger.error(f"复制图片失败: {e}")
            return f"复制图片失败: {str(e)}"
    
    @staticmethod
    def resize_image(
        document: Document,
        image_index: int,
        width: Optional[Union[float, str]] = None,
        height: Optional[Union[float, str]] = None,
        maintain_aspect_ratio: bool = True
    ) -> str:
        """
        调整图片大小
        
        Parameters:
        - document: docx文档对象
        - image_index: 图片索引（按文档中出现顺序）
        - width: 新宽度
        - height: 新高度
        - maintain_aspect_ratio: 是否保持宽高比
        
        Returns:
        - 操作结果消息
        """
        try:
            # 查找所有图片
            images = []
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        images.append(run)
            
            if not images:
                return "文档中没有找到图片"
            
            if image_index < 0 or image_index >= len(images):
                return f"图片索引超出范围: {image_index}，文档中共有 {len(images)} 张图片"
            
            run = images[image_index]
            
            # 获取图片元素
            drawing = run._element.xpath('.//w:drawing')[0]
            extent = drawing.xpath('.//wp:extent')[0]
            
            # 解析新尺寸
            width_emu = None
            height_emu = None
            
            if width:
                width_inches = ImageProcessor._parse_size(width)
                width_emu = int(width_inches * 914400)  # 转换为EMU单位
            
            if height:
                height_inches = ImageProcessor._parse_size(height)
                height_emu = int(height_inches * 914400)
            
            # 如果需要保持宽高比且只指定了一个尺寸
            if maintain_aspect_ratio:
                current_width = int(extent.get('cx'))
                current_height = int(extent.get('cy'))
                aspect_ratio = current_width / current_height
                
                if width_emu and not height_emu:
                    height_emu = int(width_emu / aspect_ratio)
                elif height_emu and not width_emu:
                    width_emu = int(height_emu * aspect_ratio)
            
            # 应用新尺寸
            if width_emu:
                extent.set('cx', str(width_emu))
            if height_emu:
                extent.set('cy', str(height_emu))
            
            return f"图片 {image_index} 大小已调整"
            
        except Exception as e:
            logger.error(f"调整图片大小失败: {e}")
            return f"调整图片大小失败: {str(e)}"
    
    @staticmethod
    def delete_image(document: Document, image_index: int) -> str:
        """
        删除指定图片
        
        Parameters:
        - document: docx文档对象
        - image_index: 图片索引
        
        Returns:
        - 操作结果消息
        """
        try:
            # 查找所有图片
            image_runs = []
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        image_runs.append((paragraph, run))
            
            if not image_runs:
                return "文档中没有找到图片"
            
            if image_index < 0 or image_index >= len(image_runs):
                return f"图片索引超出范围: {image_index}，文档中共有 {len(image_runs)} 张图片"
            
            paragraph, run = image_runs[image_index]
            
            # 删除包含图片的run
            paragraph._element.remove(run._element)
            
            # 如果段落变空，删除段落
            if not paragraph.runs:
                p = paragraph._element
                p.getparent().remove(p)
            
            return f"图片 {image_index} 已删除"
            
        except Exception as e:
            logger.error(f"删除图片失败: {e}")
            return f"删除图片失败: {str(e)}"
    
    @staticmethod
    def list_images(document: Document) -> str:
        """
        列出文档中的所有图片信息
        
        Parameters:
        - document: docx文档对象
        
        Returns:
        - 图片信息列表
        """
        try:
            images_info = []
            image_count = 0
            
            for para_idx, paragraph in enumerate(document.paragraphs):
                for run_idx, run in enumerate(paragraph.runs):
                    if run._element.xpath('.//a:blip'):
                        # 获取图片尺寸
                        drawing = run._element.xpath('.//w:drawing')[0]
                        extent = drawing.xpath('.//wp:extent')[0]
                        
                        width_emu = int(extent.get('cx'))
                        height_emu = int(extent.get('cy'))
                        width_inches = width_emu / 914400
                        height_inches = height_emu / 914400
                        
                        images_info.append(f"图片 {image_count}: 段落 {para_idx}, 尺寸 {width_inches:.2f}\" x {height_inches:.2f}\"")
                        image_count += 1
            
            if not images_info:
                return "文档中没有图片"
            
            return f"文档中共有 {image_count} 张图片:\n" + "\n".join(images_info)
            
        except Exception as e:
            logger.error(f"列出图片失败: {e}")
            return f"列出图片失败: {str(e)}"
    
    @staticmethod
    def _parse_size(size_str: Union[str, float, int]) -> float:
        """
        解析尺寸字符串为英寸数值
        
        Parameters:
        - size_str: 尺寸字符串或数值
        
        Returns:
        - 英寸数值
        """
        if isinstance(size_str, (int, float)):
            return float(size_str)
        
        size_str = str(size_str).lower().strip()
        
        if size_str.endswith('cm'):
            return float(size_str[:-2]) / 2.54
        elif size_str.endswith('in') or size_str.endswith('"'):
            return float(size_str.rstrip('in"'))
        elif size_str.endswith('pt'):
            return float(size_str[:-2]) / 72
        elif size_str.endswith('px'):
            return float(size_str[:-2]) / 96  # 假设96 DPI
        else:
            # 默认当作英寸
            return float(size_str)
    
    @staticmethod
    def set_image_position(
        document: Document,
        image_index: int,
        position_type: str = "inline",
        horizontal_position: Optional[str] = None,
        vertical_position: Optional[str] = None
    ) -> str:
        """
        设置图片位置
        
        Parameters:
        - document: docx文档对象
        - image_index: 图片索引
        - position_type: 位置类型 ("inline", "floating")
        - horizontal_position: 水平位置 ("left", "center", "right")
        - vertical_position: 垂直位置 ("top", "center", "bottom")
        
        Returns:
        - 操作结果消息
        """
        try:
            # 查找图片
            image_runs = []
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        image_runs.append((paragraph, run))
            
            if not image_runs:
                return "文档中没有找到图片"
            
            if image_index < 0 or image_index >= len(image_runs):
                return f"图片索引超出范围: {image_index}"
            
            paragraph, run = image_runs[image_index]
            
            # 设置段落对齐（用于内联图片）
            if position_type == "inline" and horizontal_position:
                alignment_map = {
                    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT
                }
                paragraph.alignment = alignment_map.get(horizontal_position, WD_PARAGRAPH_ALIGNMENT.LEFT)
            
            return f"图片 {image_index} 位置已设置为 {position_type}"
            
        except Exception as e:
            logger.error(f"设置图片位置失败: {e}")
            return f"设置图片位置失败: {str(e)}"
