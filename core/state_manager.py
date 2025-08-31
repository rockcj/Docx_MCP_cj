#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
状态管理器模块
负责处理文档会话状态的持久化和恢复
"""

import os
import tempfile
import logging
from typing import Optional, Dict, Any
from docx import Document

logger = logging.getLogger(__name__)

class StateManager:
    """文档状态管理器，支持会话持久化"""
    
    def __init__(self):
        self.state_file = os.path.join(tempfile.gettempdir(), "docx_mcp_current_doc.txt")
        self.current_document: Optional[Document] = None
        self.current_file_path: Optional[str] = None
        self.documents: Dict[str, Document] = {}
        
        # 尝试加载之前的状态
        self._load_current_document()
    
    def _load_current_document(self) -> bool:
        """从状态文件加载当前文档"""
        if not os.path.exists(self.state_file):
            return False
        
        try:
            with open(self.state_file, 'r', encoding='utf-8') as f:
                file_path = f.read().strip()
            
            if file_path and os.path.exists(file_path):
                try:
                    self.current_file_path = file_path
                    self.current_document = Document(file_path)
                    self.documents[file_path] = self.current_document
                    logger.info(f"Successfully loaded document from {file_path}")
                    return True
                except Exception as e:
                    logger.error(f"Failed to load document at {file_path}: {e}")
                    self._remove_invalid_state_file()
            else:
                self._remove_invalid_state_file()
                
        except Exception as e:
            logger.error(f"Failed to load current document: {e}")
            self._remove_invalid_state_file()
        
        return False
    
    def _remove_invalid_state_file(self):
        """删除无效的状态文件"""
        try:
            if os.path.exists(self.state_file):
                os.remove(self.state_file)
                logger.info("Removed invalid state file")
        except Exception as e:
            logger.error(f"Failed to remove state file: {e}")
    
    def _save_current_document_state(self) -> bool:
        """保存当前文档路径到状态文件"""
        if not self.current_file_path:
            return False
        
        try:
            with open(self.state_file, 'w', encoding='utf-8') as f:
                f.write(self.current_file_path)
            return True
        except Exception as e:
            logger.error(f"Failed to save current document path: {e}")
            return False
    
    def set_current_document(self, file_path: str, document: Document):
        """设置当前文档"""
        self.current_file_path = file_path
        self.current_document = document
        self.documents[file_path] = document
        self._save_current_document_state()
    
    def get_current_document(self) -> Optional[Document]:
        """获取当前文档"""
        return self.current_document
    
    def get_current_file_path(self) -> Optional[str]:
        """获取当前文档路径"""
        return self.current_file_path
    
    def save_current_document(self):
        """保存当前文档"""
        if self.current_document and self.current_file_path:
            try:
                self.current_document.save(self.current_file_path)
                self._save_current_document_state()
                logger.info(f"Document saved to {self.current_file_path}")
            except Exception as e:
                logger.error(f"Failed to save current document: {e}")
                raise
    
    def clear_state(self):
        """清除状态"""
        self.current_document = None
        self.current_file_path = None
        self.documents.clear()
        self._remove_invalid_state_file()
    
    def has_current_document(self) -> bool:
        """检查是否有当前文档"""
        return self.current_document is not None
    
    def get_document_info(self) -> Dict[str, Any]:
        """获取当前文档信息"""
        if not self.current_document:
            return {"error": "No document is open"}
        
        doc = self.current_document
        return {
            "file_path": self.current_file_path,
            "sections_count": len(doc.sections),
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
            "has_changes": True  # 简化实现，实际可以跟踪变更状态
        }
