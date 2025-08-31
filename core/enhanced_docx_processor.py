#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
增强版DOCX处理器
整合了状态管理、图片处理、字体处理等功能的完整文档处理系统
"""

import os
import tempfile
import logging
from typing import List, Dict, Any, Optional, Union
from io import BytesIO
from contextlib import contextmanager

from docx import Document
from docx.table import Table
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from .models import DocumentPatch, TableOperation, TableOperationType, CellAlignment, BorderStyle
from .state_manager import StateManager
from .state_machine import IntelligentStateManager
from .template_engine import DocumentTemplateEngine, TemplateRenderer
from .ai_interface import AISmartInterface
from .image_processor import ImageProcessor
from .font_processor import FontProcessor
from .oss_processor import OSSProcessor

logger = logging.getLogger(__name__)

class EnhancedDocxProcessor:
    """
    增强版DOCX处理器
    整合了原有功能和新增的状态管理、图片处理、字体处理功能
    """
    
    def __init__(self):
        """初始化处理器"""
        # 原有组件
        self.state_manager = StateManager()
        self.image_processor = ImageProcessor()
        self.font_processor = FontProcessor()
        self.oss_processor = OSSProcessor()
        
        # 新增智能组件
        self.intelligent_state_manager = IntelligentStateManager()
        self.template_engine = DocumentTemplateEngine()
        self.template_renderer = TemplateRenderer(self.template_engine)
        self.ai_interface = AISmartInterface(self.intelligent_state_manager, self.template_engine)
        
        # 配置日志
        logging.basicConfig(
            level=logging.INFO,
            format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
            handlers=[
                logging.FileHandler(os.path.join(tempfile.gettempdir(), "enhanced_docx_processor.log")),
                logging.StreamHandler()
            ]
        )
    
    # ==================== 文档生命周期管理 ====================
    
    def create_document(self, file_path: str) -> str:
        """创建新文档"""
        try:
            document = Document()
            self.state_manager.set_current_document(file_path, document)
            document.save(file_path)
            logger.info(f"Document created: {file_path}")
            return f"文档创建成功: {file_path}"
        except Exception as e:
            logger.error(f"Failed to create document: {e}")
            return f"创建文档失败: {str(e)}"
    
    def open_document(self, file_path: str) -> str:
        """打开现有文档"""
        try:
            if not os.path.exists(file_path):
                return f"文件不存在: {file_path}"
            
            document = Document(file_path)
            self.state_manager.set_current_document(file_path, document)
            logger.info(f"Document opened: {file_path}")
            return f"文档打开成功: {file_path}"
        except Exception as e:
            logger.error(f"Failed to open document: {e}")
            return f"打开文档失败: {str(e)}"
    
    def save_document(self) -> str:
        """保存当前文档"""
        try:
            if not self.state_manager.has_current_document():
                return "没有打开的文档"
            
            self.state_manager.save_current_document()
            return f"文档保存成功: {self.state_manager.get_current_file_path()}"
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            return f"保存文档失败: {str(e)}"
    
    def save_as_document(self, new_file_path: str) -> str:
        """另存为新文件"""
        try:
            if not self.state_manager.has_current_document():
                return "没有打开的文档"
            
            document = self.state_manager.get_current_document()
            document.save(new_file_path)
            self.state_manager.set_current_document(new_file_path, document)
            return f"文档另存为: {new_file_path}"
        except Exception as e:
            logger.error(f"Failed to save document as: {e}")
            return f"另存文档失败: {str(e)}"
    
    def close_document(self) -> str:
        """关闭当前文档"""
        try:
            if not self.state_manager.has_current_document():
                return "没有打开的文档"
            
            file_path = self.state_manager.get_current_file_path()
            self.state_manager.clear_state()
            return f"文档已关闭: {file_path}"
        except Exception as e:
            logger.error(f"Failed to close document: {e}")
            return f"关闭文档失败: {str(e)}"
    
    def get_document_info(self) -> str:
        """获取文档信息"""
        try:
            info = self.state_manager.get_document_info()
            if "error" in info:
                return info["error"]
            
            result = []
            result.append(f"文档路径: {info['file_path']}")
            result.append(f"节数: {info['sections_count']}")
            result.append(f"段落数: {info['paragraphs_count']}")
            result.append(f"表格数: {info['tables_count']}")
            
            return "\n".join(result)
        except Exception as e:
            logger.error(f"Failed to get document info: {e}")
            return f"获取文档信息失败: {str(e)}"
    
    # ==================== 内容编辑功能 ====================
    
    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        font_size: Optional[int] = None,
        font_name: Optional[str] = None,
        color: Optional[str] = None,
        alignment: Optional[str] = None,
        style: Optional[str] = None
    ) -> str:
        """添加段落"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            # 添加段落
            paragraph = document.add_paragraph(text)
            
            # 设置样式
            if style:
                try:
                    paragraph.style = document.styles[style]
                except KeyError:
                    pass  # 样式不存在时忽略
            
            # 应用格式
            if paragraph.runs:
                run = paragraph.runs[0]
                self.font_processor._apply_run_formatting(
                    run, font_name, font_size, bold, italic, underline, color
                )
            
            # 设置对齐
            if alignment:
                alignment_map = {
                    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                }
                if alignment in alignment_map:
                    paragraph.alignment = alignment_map[alignment]
            
            return "段落添加成功"
        except Exception as e:
            logger.error(f"Failed to add paragraph: {e}")
            return f"添加段落失败: {str(e)}"
    
    def add_heading(self, text: str, level: int) -> str:
        """添加标题"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            if level < 1 or level > 9:
                return f"标题级别必须在1-9之间: {level}"
            
            document.add_heading(text, level=level)
            return f"添加了 {level} 级标题"
        except Exception as e:
            logger.error(f"Failed to add heading: {e}")
            return f"添加标题失败: {str(e)}"
    
    def add_page_break(self) -> str:
        """添加分页符"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            document.add_page_break()
            return "分页符添加成功"
        except Exception as e:
            logger.error(f"Failed to add page break: {e}")
            return f"添加分页符失败: {str(e)}"
    
    def delete_paragraph(self, paragraph_index: int) -> str:
        """删除段落"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
                return f"段落索引超出范围: {paragraph_index}"
            
            paragraph = document.paragraphs[paragraph_index]
            p = paragraph._element
            p.getparent().remove(p)
            
            return f"段落 {paragraph_index} 删除成功"
        except Exception as e:
            logger.error(f"Failed to delete paragraph: {e}")
            return f"删除段落失败: {str(e)}"
    
    # ==================== 表格操作功能 ====================
    
    def add_table(self, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
        """添加表格"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            table = document.add_table(rows=rows, cols=cols, style="Table Grid")
            
            # 填充数据
            if data:
                for i, row_data in enumerate(data):
                    if i < rows:
                        row = table.rows[i]
                        for j, cell_text in enumerate(row_data):
                            if j < cols:
                                row.cells[j].text = str(cell_text)
            
            return f"添加了 {rows}x{cols} 表格"
        except Exception as e:
            logger.error(f"Failed to add table: {e}")
            return f"添加表格失败: {str(e)}"
    
    def add_table_row(self, table_index: int, data: Optional[List[str]] = None, row_index: Optional[int] = None) -> str:
        """
        添加表格行
        :param table_index: 表格索引
        :param data: 行数据列表
        :param row_index: 插入位置，None表示在末尾添加，0表示在开头插入，1表示在第1行和第2行之间插入
        """
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            if table_index < 0 or table_index >= len(document.tables):
                return f"表格索引超出范围: {table_index}"
            
            table = document.tables[table_index]
            
            # 如果没有指定插入位置，在末尾添加
            if row_index is None:
                new_row = table.add_row()
                if data:
                    for i, cell_text in enumerate(data):
                        if i < len(new_row.cells):
                            new_row.cells[i].text = str(cell_text)
                return f"表格 {table_index} 在末尾添加行成功"
            
            # 验证插入位置
            if row_index < 0 or row_index > len(table.rows):
                return f"行索引超出范围: {row_index}，当前表格有 {len(table.rows)} 行"
            
            # 在指定位置插入行
            from docx.oxml.shared import OxmlElement
            
            # 创建新行元素
            new_tr = OxmlElement('w:tr')
            
            # 复制第一行的样式作为模板（如果表格不为空）
            template_row = None
            if len(table.rows) > 0:
                template_row = table.rows[0]._tr
            
            # 为新行创建单元格
            for col_idx in range(len(table.columns)):
                new_tc = OxmlElement('w:tc')
                
                # 创建基本的单元格属性
                tc_pr = OxmlElement('w:tcPr')
                tc_w = OxmlElement('w:tcW')
                tc_w.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', "2000")
                tc_w.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', "dxa")
                tc_pr.append(tc_w)
                new_tc.append(tc_pr)
                
                # 创建段落元素
                p_element = OxmlElement('w:p')
                new_tc.append(p_element)
                new_tr.append(new_tc)
            
            # 在指定位置插入新行
            table_element = table._tbl
            if row_index >= len(table.rows):
                # 在末尾添加
                table_element.append(new_tr)
            else:
                # 在指定位置插入
                target_row = table.rows[row_index]._tr
                table_element.insert(table_element.index(target_row), new_tr)
            
            # 刷新表格以获取新添加的行
            table = document.tables[table_index]  # 重新获取表格引用
            
            # 填充数据
            if data:
                inserted_row = table.rows[row_index]
                for i, cell_text in enumerate(data):
                    if i < len(inserted_row.cells):
                        inserted_row.cells[i].text = str(cell_text)
            
            return f"表格 {table_index} 在位置 {row_index} 插入行成功"
            
        except Exception as e:
            logger.error(f"Failed to add table row: {e}")
            return f"添加表格行失败: {str(e)}"
    
    def delete_table_row(self, table_index: int, row_index: int) -> str:
        """删除表格行"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            if table_index < 0 or table_index >= len(document.tables):
                return f"表格索引超出范围: {table_index}"
            
            table = document.tables[table_index]
            
            if row_index < 0 or row_index >= len(table.rows):
                return f"行索引超出范围: {row_index}"
            
            row = table.rows[row_index]._tr
            row.getparent().remove(row)
            
            return f"表格 {table_index} 行 {row_index} 删除成功"
        except Exception as e:
            logger.error(f"Failed to delete table row: {e}")
            return f"删除表格行失败: {str(e)}"
    
    def add_table_column(self, table_index: int, column_index: Optional[int] = None, data: Optional[List[str]] = None) -> str:
        """添加表格列"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            if table_index < 0 or table_index >= len(document.tables):
                return f"表格索引超出范围: {table_index}"
            
            table = document.tables[table_index]
            
            # 确定插入位置
            if column_index is None:
                column_index = len(table.columns)
            elif column_index < 0 or column_index > len(table.columns):
                return f"列索引超出范围: {column_index}"
            
            # 为每一行添加新单元格
            for row_idx, row in enumerate(table.rows):
                # 创建新单元格
                new_cell_element = OxmlElement('w:tc')
                
                # 如果有模板单元格，复制其属性
                if len(row.cells) > 0:
                    template_cell = row.cells[0]._tc
                    for child in template_cell:
                        if child.tag.endswith('tcPr'):
                            new_cell_element.append(child)
                            break
                
                # 添加空段落
                p_element = OxmlElement('w:p')
                new_cell_element.append(p_element)
                
                # 插入到指定位置
                if column_index >= len(row.cells):
                    row._tr.append(new_cell_element)
                else:
                    row._tr.insert(column_index, new_cell_element)
                
                # 填充数据
                if data and row_idx < len(data):
                    # 需要重新获取单元格引用
                    table = document.tables[table_index]  # 重新获取表格
                    new_cell = table.rows[row_idx].cells[column_index]
                    new_cell.text = str(data[row_idx])
            
            return f"表格 {table_index} 添加列成功"
        except Exception as e:
            logger.error(f"Failed to add table column: {e}")
            return f"添加表格列失败: {str(e)}"
    
    def delete_table_column(self, table_index: int, column_index: int) -> str:
        """删除表格列"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            if table_index < 0 or table_index >= len(document.tables):
                return f"表格索引超出范围: {table_index}"
            
            table = document.tables[table_index]
            
            if column_index < 0 or column_index >= len(table.columns):
                return f"列索引超出范围: {column_index}"
            
            # 从每一行删除指定列的单元格
            for row in table.rows:
                if column_index < len(row.cells):
                    cell_element = row.cells[column_index]._tc
                    cell_element.getparent().remove(cell_element)
            
            return f"表格 {table_index} 列 {column_index} 删除成功"
        except Exception as e:
            logger.error(f"Failed to delete table column: {e}")
            return f"删除表格列失败: {str(e)}"
    
    def edit_table_cell(self, table_index: int, row_index: int, col_index: int, text: str) -> str:
        """编辑表格单元格"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            if table_index < 0 or table_index >= len(document.tables):
                return f"表格索引超出范围: {table_index}"
            
            table = document.tables[table_index]
            
            if row_index < 0 or row_index >= len(table.rows):
                return f"行索引超出范围: {row_index}"
            
            if col_index < 0 or col_index >= len(table.columns):
                return f"列索引超出范围: {col_index}"
            
            table.cell(row_index, col_index).text = text
            
            return f"单元格 ({row_index}, {col_index}) 编辑成功"
        except Exception as e:
            logger.error(f"Failed to edit table cell: {e}")
            return f"编辑单元格失败: {str(e)}"
    
    def merge_table_cells(
        self,
        table_index: int,
        start_row: int,
        start_col: int,
        end_row: int,
        end_col: int
    ) -> str:
        """合并表格单元格"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            if table_index < 0 or table_index >= len(document.tables):
                return f"表格索引超出范围: {table_index}"
            
            table = document.tables[table_index]
            
            # 验证索引
            if not (0 <= start_row <= end_row < len(table.rows)):
                return "行索引无效"
            
            if not (0 <= start_col <= end_col < len(table.columns)):
                return "列索引无效"
            
            start_cell = table.cell(start_row, start_col)
            end_cell = table.cell(end_row, end_col)
            start_cell.merge(end_cell)
            
            return f"单元格合并成功: ({start_row},{start_col}) 到 ({end_row},{end_col})"
        except Exception as e:
            logger.error(f"Failed to merge cells: {e}")
            return f"合并单元格失败: {str(e)}"
    
    # ==================== 搜索和替换功能 ====================
    
    def search_text(self, keyword: str) -> str:
        """搜索文本"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            results = []
            
            # 在段落中搜索
            for i, paragraph in enumerate(document.paragraphs):
                if keyword in paragraph.text:
                    results.append(f"段落 {i}: {paragraph.text[:100]}...")
            
            # 在表格中搜索
            for t_idx, table in enumerate(document.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if keyword in cell.text:
                            results.append(f"表格 {t_idx} 单元格 ({r_idx},{c_idx}): {cell.text[:50]}...")
            
            if not results:
                return f"未找到关键词: {keyword}"
            
            return f"找到 {len(results)} 个匹配项:\n" + "\n".join(results)
        except Exception as e:
            logger.error(f"Failed to search text: {e}")
            return f"搜索失败: {str(e)}"
    
    def find_and_replace(self, find_text: str, replace_text: str) -> str:
        """查找并替换文本"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            replace_count = 0
            
            # 在段落中替换
            for paragraph in document.paragraphs:
                if find_text in paragraph.text:
                    original_count = paragraph.text.count(find_text)
                    paragraph.text = paragraph.text.replace(find_text, replace_text)
                    replace_count += original_count
            
            # 在表格中替换
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if find_text in paragraph.text:
                                original_count = paragraph.text.count(find_text)
                                paragraph.text = paragraph.text.replace(find_text, replace_text)
                                replace_count += original_count
            
            return f"替换完成，共替换 {replace_count} 处"
        except Exception as e:
            logger.error(f"Failed to find and replace: {e}")
            return f"查找替换失败: {str(e)}"
    
    # ==================== 图片处理功能代理 ====================
    
    def add_image(self, image_path: str, width: Optional[str] = None, height: Optional[str] = None, 
                  alignment: str = "left", paragraph_index: Optional[int] = None) -> str:
        """添加图片"""
        document = self.state_manager.get_current_document()
        if not document:
            return "没有打开的文档"
        
        return self.image_processor.add_image(
            document, image_path, width, height, alignment, paragraph_index
        )
    
    def resize_image(self, image_index: int, width: Optional[str] = None, 
                     height: Optional[str] = None, maintain_aspect_ratio: bool = True) -> str:
        """调整图片大小"""
        document = self.state_manager.get_current_document()
        if not document:
            return "没有打开的文档"
        
        return self.image_processor.resize_image(
            document, image_index, width, height, maintain_aspect_ratio
        )
    
    def delete_image(self, image_index: int) -> str:
        """删除图片"""
        document = self.state_manager.get_current_document()
        if not document:
            return "没有打开的文档"
        
        return self.image_processor.delete_image(document, image_index)
    
    def list_images(self) -> str:
        """列出所有图片"""
        document = self.state_manager.get_current_document()
        if not document:
            return "没有打开的文档"
        
        return self.image_processor.list_images(document)
    
    def extract_image_to_local(self, image_index: int, output_dir: str = "extracted_images", 
                              filename: Optional[str] = None) -> str:
        """提取图片到本地"""
        document = self.state_manager.get_current_document()
        if not document:
            return "没有打开的文档"
        
        return self.image_processor.extract_image_to_local(document, image_index, output_dir, filename)
    
    def extract_all_images_to_local(self, output_dir: str = "extracted_images") -> str:
        """提取所有图片到本地"""
        document = self.state_manager.get_current_document()
        if not document:
            return "没有打开的文档"
        
        return self.image_processor.extract_all_images_to_local(document, output_dir)
    
    def move_image_to_another_document(self, target_doc_path: str, image_index: int,
                                     width: Optional[str] = None, height: Optional[str] = None,
                                     alignment: str = "left", paragraph_index: Optional[int] = None) -> str:
        """移动图片到另一个文档"""
        source_document = self.state_manager.get_current_document()
        if not source_document:
            return "没有打开的源文档"
        
        try:
            from docx import Document
            target_document = Document(target_doc_path)
            
            result = self.image_processor.move_image_to_another_document(
                source_document, target_document, image_index, width, height, alignment, paragraph_index
            )
            
            # 保存目标文档
            target_document.save(target_doc_path)
            
            return result
            
        except Exception as e:
            return f"移动图片失败: {str(e)}"
    
    def copy_image_to_another_document(self, target_doc_path: str, image_index: int,
                                     width: Optional[str] = None, height: Optional[str] = None,
                                     alignment: str = "left", paragraph_index: Optional[int] = None) -> str:
        """复制图片到另一个文档"""
        source_document = self.state_manager.get_current_document()
        if not source_document:
            return "没有打开的源文档"
        
        try:
            from docx import Document
            target_document = Document(target_doc_path)
            
            result = self.image_processor.copy_image_to_another_document(
                source_document, target_document, image_index, width, height, alignment, paragraph_index
            )
            
            # 保存目标文档
            target_document.save(target_doc_path)
            
            return result
            
        except Exception as e:
            return f"复制图片失败: {str(e)}"
    
    # ==================== 字体处理功能代理 ====================
    
    def set_paragraph_font(self, paragraph_index: int, font_name: Optional[str] = None,
                          font_size: Optional[int] = None, bold: Optional[bool] = None,
                          italic: Optional[bool] = None, underline: Optional[bool] = None,
                          color: Optional[str] = None, alignment: Optional[str] = None) -> str:
        """设置段落字体"""
        document = self.state_manager.get_current_document()
        if not document:
            return "没有打开的文档"
        
        return self.font_processor.set_paragraph_font(
            document, paragraph_index, font_name, font_size, bold, italic, underline, color, alignment
        )
    
    def set_text_range_font(self, paragraph_index: int, start_pos: int, end_pos: int,
                           font_name: Optional[str] = None, font_size: Optional[int] = None,
                           bold: Optional[bool] = None, italic: Optional[bool] = None,
                           underline: Optional[bool] = None, color: Optional[str] = None) -> str:
        """设置文本范围字体"""
        document = self.state_manager.get_current_document()
        if not document:
            return "没有打开的文档"
        
        return self.font_processor.set_text_range_font(
            document, paragraph_index, start_pos, end_pos, font_name, font_size, bold, italic, underline, color
        )
    
    def get_font_info(self, paragraph_index: int) -> str:
        """获取字体信息"""
        document = self.state_manager.get_current_document()
        if not document:
            return "没有打开的文档"
        
        return self.font_processor.get_font_info(document, paragraph_index)
    
    # ==================== 页面设置功能 ====================
    
    def set_page_margins(self, top: Optional[float] = None, bottom: Optional[float] = None,
                        left: Optional[float] = None, right: Optional[float] = None) -> str:
        """设置页边距"""
        try:
            document = self.state_manager.get_current_document()
            if not document:
                return "没有打开的文档"
            
            section = document.sections[0]
            
            if top is not None:
                section.top_margin = Cm(top)
            if bottom is not None:
                section.bottom_margin = Cm(bottom)
            if left is not None:
                section.left_margin = Cm(left)
            if right is not None:
                section.right_margin = Cm(right)
            
            return "页边距设置成功"
        except Exception as e:
            logger.error(f"Failed to set margins: {e}")
            return f"设置页边距失败: {str(e)}"
    
    # ==================== 上下文管理器 ====================
    
    @contextmanager
    def document_context(self, file_path: str, create_if_not_exists: bool = False):
        """文档上下文管理器，确保文档正确打开和关闭"""
        original_doc = self.state_manager.get_current_document()
        original_path = self.state_manager.get_current_file_path()
        
        try:
            if create_if_not_exists and not os.path.exists(file_path):
                self.create_document(file_path)
            else:
                self.open_document(file_path)
            
            yield self
        finally:
            # 恢复原始状态
            if original_doc and original_path:
                self.state_manager.set_current_document(original_path, original_doc)
            else:
                self.state_manager.clear_state()
    
    # ==================== 静态方法保持兼容性 ====================
    
    @staticmethod
    def extract_structure_with_ids(file_stream: BytesIO) -> Dict[str, Any]:
        """保持与原有代码的兼容性"""
        from .docx_processor import DocxProcessor
        return DocxProcessor.extract_structure_with_ids(file_stream)
    
    @staticmethod
    def apply_patches(original_stream: BytesIO, modified_stream: BytesIO, patches: List[DocumentPatch]):
        """保持与原有代码的兼容性"""
        from .docx_processor import DocxProcessor
        return DocxProcessor.apply_patches(original_stream, modified_stream, patches)
    
    # ==================== OSS云存储功能 ====================
    
    def upload_current_document_to_oss(self, custom_filename: str = None) -> Dict[str, Any]:
        """上传当前文档到OSS"""
        try:
            if not self.state_manager.has_current_document():
                return {"error": "没有打开的文档"}
            
            file_path = self.state_manager.get_current_file_path()
            if not file_path:
                return {"error": "当前文档未保存，请先保存文档"}
            
            # 先保存文档
            self.save_document()
            
            return self.oss_processor.upload_file_to_oss(file_path, custom_filename)
        except Exception as e:
            logger.error(f"Failed to upload current document to OSS: {e}")
            return {"error": f"上传当前文档失败: {str(e)}"}
    
    def upload_file_to_oss(self, file_path: str, custom_filename: str = None) -> Dict[str, Any]:
        """上传指定文件到OSS"""
        try:
            return self.oss_processor.upload_file_to_oss(file_path, custom_filename)
        except Exception as e:
            logger.error(f"Failed to upload file to OSS: {e}")
            return {"error": f"上传文件失败: {str(e)}"}
    
    def download_file_from_oss(self, filename: str, local_path: str = None) -> Dict[str, Any]:
        """从OSS下载文件"""
        try:
            return self.oss_processor.download_file_from_oss(filename, local_path)
        except Exception as e:
            logger.error(f"Failed to download file from OSS: {e}")
            return {"error": f"从OSS下载文件失败: {str(e)}"}
    
    def download_file_from_url(self, url: str, local_path: str = None) -> Dict[str, Any]:
        """从网络URL下载文件"""
        try:
            return self.oss_processor.download_file_from_url(url, local_path)
        except Exception as e:
            logger.error(f"Failed to download file from URL: {e}")
            return {"error": f"从URL下载文件失败: {str(e)}"}
    
    def open_document_from_url(self, url: str) -> str:
        """从网络URL下载并打开文档"""
        try:
            # 下载文件
            result = self.download_file_from_url(url)
            if "error" in result:
                return result["error"]
            
            # 打开下载的文件
            return self.open_document(result["local_path"])
        except Exception as e:
            logger.error(f"Failed to open document from URL: {e}")
            return f"从URL打开文档失败: {str(e)}"
    
    def list_oss_files(self, prefix: str = "", max_keys: int = 100) -> Dict[str, Any]:
        """列出OSS中的文件"""
        try:
            return self.oss_processor.list_oss_files(prefix, max_keys)
        except Exception as e:
            logger.error(f"Failed to list OSS files: {e}")
            return {"error": f"列出OSS文件失败: {str(e)}"}
    
    def delete_oss_file(self, filename: str) -> Dict[str, Any]:
        """删除OSS中的文件"""
        try:
            return self.oss_processor.delete_oss_file(filename)
        except Exception as e:
            logger.error(f"Failed to delete OSS file: {e}")
            return {"error": f"删除OSS文件失败: {str(e)}"}
    
    def process_document_with_oss_upload(self, modifications: List[Dict[str, Any]]) -> Dict[str, Any]:
        """处理文档并上传到OSS"""
        try:
            if not self.state_manager.has_current_document():
                return {"error": "没有打开的文档"}
            
            # 应用修改（这里简化处理，实际可以根据需要实现更复杂的修改逻辑）
            for modification in modifications:
                if modification.get("type") == "add_paragraph":
                    self.add_paragraph(modification.get("text", ""))
                elif modification.get("type") == "add_heading":
                    self.add_heading(modification.get("text", ""), modification.get("level", 1))
            
            # 保存文档
            self.save_document()
            
            # 上传到OSS
            upload_result = self.upload_current_document_to_oss()
            
            return upload_result
        except Exception as e:
            logger.error(f"Failed to process document with OSS upload: {e}")
            return {"error": f"处理文档并上传失败: {str(e)}"}
    
    # ==================== AI智能接口方法 ====================
    
    def smart_execute(self, user_intent: str, context: Dict[str, Any] = None) -> str:
        """智能执行 - AI友好的主要接口"""
        try:
            result = self.ai_interface.smart_execute(user_intent, context)
            
            if result.success:
                # 如果建议了具体操作，执行状态转换
                if result.data and "suggested_operation" in result.data:
                    operation = result.data["suggested_operation"]
                    self.intelligent_state_manager.execute_state_transition(operation)
                
                return result.message + "\n" + "\n".join([f"- {s}" for s in result.suggestions])
            else:
                return f"执行失败: {result.message}\n建议: " + "\n".join([f"- {s}" for s in result.suggestions])
                
        except Exception as e:
            logger.error(f"Smart execute failed: {e}")
            return f"智能执行失败: {str(e)}"
    
    def suggest_template(self, user_intent: str) -> str:
        """建议模板"""
        try:
            suggestions = self.template_engine.suggest_templates_by_intent(user_intent)
            
            if not suggestions:
                return "未找到合适的模板，请尝试更具体的描述"
            
            result_lines = ["找到以下模板建议:"]
            for i, suggestion in enumerate(suggestions[:3], 1):
                result_lines.append(f"{i}. {suggestion.template_name}")
                result_lines.append(f"   匹配度: {suggestion.match_score:.1%}")
                result_lines.append(f"   原因: {suggestion.reason}")
                result_lines.append(f"   预览: {suggestion.preview}")
                result_lines.append("")
            
            return "\n".join(result_lines)
            
        except Exception as e:
            logger.error(f"Suggest template failed: {e}")
            return f"模板建议失败: {str(e)}"
    
    def fill_template(self, template_id: str, variables: Dict[str, Any]) -> str:
        """填写模板数据"""
        try:
            result = self.ai_interface.fill_template(template_id, variables)
            
            if result.success:
                return f"模板填写成功: {template_id}\n" + "\n".join([f"- {s}" for s in result.suggestions])
            else:
                return f"模板填写失败: {result.message}\n" + "\n".join([f"- {s}" for s in result.suggestions])
                
        except Exception as e:
            logger.error(f"Fill template failed: {e}")
            return f"模板填写失败: {str(e)}"
    
    def render_template_document(self, template_id: str, variables: Dict[str, Any], output_path: str) -> str:
        """从模板渲染文档"""
        try:
            # 验证模板数据
            validation_result = self.template_engine.validate_template_data(template_id, variables)
            if not validation_result.is_valid:
                return f"模板数据验证失败: {', '.join(validation_result.error_messages)}"
            
            # 渲染模板结构
            document_structure = self.template_renderer.render_template_to_document(template_id, variables)
            
            # 创建新文档
            document = Document()
            
            # 应用页面设置
            if "page_settings" in document_structure:
                page_settings = document_structure["page_settings"]
                if "margins" in page_settings:
                    margins = page_settings["margins"]
                    section = document.sections[0]
                    section.top_margin = Cm(margins.get("top", 2.54))
                    section.bottom_margin = Cm(margins.get("bottom", 2.54))
                    section.left_margin = Cm(margins.get("left", 3.18))
                    section.right_margin = Cm(margins.get("right", 3.18))
            
            # 渲染段落
            for section in document_structure["sections"]:
                if section["section_type"] == "paragraph":
                    self._render_paragraph_from_template(document, section)
                elif section["section_type"] == "heading":
                    self._render_heading_from_template(document, section)
                elif section["section_type"] == "table":
                    self._render_table_from_template(document, section)
            
            # 保存文档
            document.save(output_path)
            
            # 更新状态
            self.state_manager.set_current_document(output_path, document)
            self.intelligent_state_manager.execute_state_transition("create_document")
            
            return f"模板文档创建成功: {output_path}"
            
        except Exception as e:
            logger.error(f"Render template document failed: {e}")
            return f"模板文档渲染失败: {str(e)}"
    
    def _render_paragraph_from_template(self, document: Document, section: Dict[str, Any]):
        """从模板渲染段落"""
        content = section["content"]
        paragraph = document.add_paragraph(content["text"])
        
        # 设置对齐
        if "alignment" in content:
            alignment_map = {
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            if content["alignment"] in alignment_map:
                paragraph.alignment = alignment_map[content["alignment"]]
        
        # 设置字体
        if paragraph.runs:
            run = paragraph.runs[0]
            if "font_size" in content:
                run.font.size = Pt(content["font_size"])
            if "bold" in content and content["bold"]:
                run.bold = True
    
    def _render_heading_from_template(self, document: Document, section: Dict[str, Any]):
        """从模板渲染标题"""
        content = section["content"]
        level = content.get("level", 1)
        document.add_heading(content["text"], level=level)
    
    def _render_table_from_template(self, document: Document, section: Dict[str, Any]):
        """从模板渲染表格"""
        content = section["content"]
        rows = content.get("rows", 3)
        cols = content.get("columns", 3)
        
        table = document.add_table(rows=rows, cols=cols, style=content.get("style", "Table Grid"))
        
        # 填充表头数据
        if "data" in content and content["data"]:
            for row_data in content["data"]:
                if row_data["row"] < len(table.rows):
                    row = table.rows[row_data["row"]]
                    for i, cell_data in enumerate(row_data["cells"]):
                        if i < len(row.cells):
                            cell = row.cells[i]
                            cell.text = cell_data["text"]
                            
                            # 设置单元格格式
                            if "bold" in cell_data and cell_data["bold"]:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
    
    def get_available_templates(self, category: str = None) -> str:
        """获取可用模板列表"""
        try:
            templates = self.template_engine.get_available_templates(category)
            
            if not templates:
                return "未找到可用模板"
            
            result_lines = ["可用模板列表:"]
            for template in templates:
                result_lines.append(f"- {template.metadata.name} ({template.metadata.id})")
                result_lines.append(f"  分类: {template.metadata.category}")
                result_lines.append(f"  描述: {template.metadata.description}")
                result_lines.append("")
            
            return "\n".join(result_lines)
            
        except Exception as e:
            logger.error(f"Get available templates failed: {e}")
            return f"获取模板列表失败: {str(e)}"
    
    def get_current_state_info(self) -> str:
        """获取当前状态信息"""
        try:
            state_summary = self.intelligent_state_manager.get_current_state_summary()
            stats = self.intelligent_state_manager.get_operation_statistics()
            
            result_lines = ["当前状态信息:"]
            result_lines.append(f"全局状态: {state_summary['global']}")
            result_lines.append(f"内容编辑状态: {state_summary['content']}")
            result_lines.append(f"表格操作状态: {state_summary['table']}")
            result_lines.append(f"图片处理状态: {state_summary['image']}")
            result_lines.append(f"字体处理状态: {state_summary['font']}")
            result_lines.append("")
            result_lines.append("操作统计:")
            result_lines.append(f"总操作数: {stats['total_operations']}")
            result_lines.append(f"成功操作数: {stats['successful_operations']}")
            result_lines.append(f"成功率: {stats['success_rate']:.1%}")
            
            return "\n".join(result_lines)
            
        except Exception as e:
            logger.error(f"Get current state info failed: {e}")
            return f"获取状态信息失败: {str(e)}"
