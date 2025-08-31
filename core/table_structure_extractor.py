#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
表格结构提取器 - 支持A3和A4格式文档的表格结构提取和缓存管理
"""

import json
import logging
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, asdict
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.oxml import OxmlElement
import threading
import time

logger = logging.getLogger(__name__)

@dataclass
class CellInfo:
    """单元格信息"""
    row_index: int
    col_index: int
    text: str
    is_merged: bool = False
    merge_span: Tuple[int, int] = (1, 1)  # (row_span, col_span)
    cell_type: str = "normal"  # normal, header, data, empty
    style_info: Dict[str, Any] = None
    size_info: Dict[str, Any] = None  # 大小信息
    position_info: Dict[str, Any] = None  # 位置信息
    format_info: Dict[str, Any] = None  # 格式信息

@dataclass
class TableStructure:
    """表格结构信息"""
    table_index: int
    rows: int
    columns: int
    cells: List[List[CellInfo]]
    table_type: str = "unknown"  # evaluation, signature, award, internship, etc.
    page_format: str = "A4"  # A3, A4
    has_merged_cells: bool = False
    header_rows: int = 0
    data_rows: int = 0
    extracted_at: float = 0.0

class TableStructureCache:
    """表格结构缓存管理器"""
    
    def __init__(self, max_cache_size: int = 10, cache_timeout: int = 3600):
        """
        初始化缓存管理器
        
        Args:
            max_cache_size: 最大缓存数量
            cache_timeout: 缓存超时时间（秒）
        """
        self._cache: Dict[str, TableStructure] = {}
        self._cache_timestamps: Dict[str, float] = {}
        self._max_cache_size = max_cache_size
        self._cache_timeout = cache_timeout
        self._lock = threading.RLock()
        
    def _generate_cache_key(self, file_path: str, table_index: int) -> str:
        """生成缓存键"""
        file_hash = hash(str(Path(file_path).resolve()))
        return f"{file_hash}_{table_index}"
    
    def _cleanup_expired_cache(self):
        """清理过期的缓存"""
        current_time = time.time()
        expired_keys = []
        
        for key, timestamp in self._cache_timestamps.items():
            if current_time - timestamp > self._cache_timeout:
                expired_keys.append(key)
        
        for key in expired_keys:
            self._cache.pop(key, None)
            self._cache_timestamps.pop(key, None)
            logger.debug(f"清理过期缓存: {key}")
    
    def _evict_oldest_cache(self):
        """驱逐最旧的缓存"""
        if not self._cache_timestamps:
            return
            
        oldest_key = min(self._cache_timestamps.keys(), 
                        key=lambda k: self._cache_timestamps[k])
        self._cache.pop(oldest_key, None)
        self._cache_timestamps.pop(oldest_key, None)
        logger.debug(f"驱逐最旧缓存: {oldest_key}")
    
    def get(self, file_path: str, table_index: int) -> Optional[TableStructure]:
        """获取缓存的表格结构"""
        with self._lock:
            self._cleanup_expired_cache()
            
            cache_key = self._generate_cache_key(file_path, table_index)
            if cache_key in self._cache:
                # 更新访问时间
                self._cache_timestamps[cache_key] = time.time()
                logger.debug(f"从缓存获取表格结构: {cache_key}")
                return self._cache[cache_key]
            
            return None
    
    def put(self, file_path: str, table_index: int, structure: TableStructure):
        """存储表格结构到缓存"""
        with self._lock:
            self._cleanup_expired_cache()
            
            # 如果缓存已满，驱逐最旧的
            if len(self._cache) >= self._max_cache_size:
                self._evict_oldest_cache()
            
            cache_key = self._generate_cache_key(file_path, table_index)
            self._cache[cache_key] = structure
            self._cache_timestamps[cache_key] = time.time()
            logger.debug(f"存储表格结构到缓存: {cache_key}")
    
    def clear(self):
        """清空所有缓存"""
        with self._lock:
            self._cache.clear()
            self._cache_timestamps.clear()
            logger.info("清空所有表格结构缓存")
    
    def get_cache_info(self) -> Dict[str, Any]:
        """获取缓存信息"""
        with self._lock:
            return {
                "cache_size": len(self._cache),
                "max_cache_size": self._max_cache_size,
                "cache_timeout": self._cache_timeout,
                "cached_files": list(self._cache.keys())
            }

class TableStructureExtractor:
    """表格结构提取器"""
    
    def __init__(self):
        self.cache = TableStructureCache()
        self._page_formats = {
            (7772400, 10058400): "A4",  # 标准A4尺寸
            (14171295, 9829165): "A3",  # 标准A3尺寸
        }
        
    def _detect_page_format(self, doc: Document) -> str:
        """检测页面格式"""
        if not doc.sections:
            return "A4"
            
        section = doc.sections[0]
        page_width = section.page_width
        page_height = section.page_height
        
        # 检查标准尺寸
        for (width, height), format_name in self._page_formats.items():
            if abs(page_width - width) < 1000 and abs(page_height - height) < 1000:
                return format_name
        
        # 根据尺寸比例判断
        ratio = page_width / page_height if page_height > 0 else 1
        if ratio > 1.4:  # A3通常更宽
            return "A3"
        else:
            return "A4"
    
    def _detect_table_type(self, table: Table, page_format: str) -> str:
        """检测表格类型"""
        if not table.rows or not table.columns:
            return "unknown"
        
        # 获取表格前几行的文本内容
        sample_text = []
        for row in table.rows[:3]:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                sample_text.extend(row_text)
        
        sample_text_str = " ".join(sample_text).lower()
        
        # 根据关键词判断表格类型
        if any(keyword in sample_text_str for keyword in ["评价", "德育", "智育", "文体", "综合素质"]):
            return "evaluation"
        elif any(keyword in sample_text_str for keyword in ["签名", "盖章", "日期", "学院"]):
            return "signature"
        elif any(keyword in sample_text_str for keyword in ["获奖", "时间", "项目", "等级"]):
            return "award"
        elif any(keyword in sample_text_str for keyword in ["实习", "鉴定", "单位", "指导"]):
            return "internship"
        elif any(keyword in sample_text_str for keyword in ["姓名", "学号", "班级", "专业"]):
            return "student_info"
        else:
            return "general"
    
    def _extract_cell_info(self, table: Table, row_idx: int, col_idx: int) -> CellInfo:
        """提取单元格信息"""
        try:
            cell = table.cell(row_idx, col_idx)
            cell_text = cell.text.strip()
            
            # 检测合并单元格
            is_merged = False
            merge_span = (1, 1)
            row_span = 1
            col_span = 1
            
            # 检查单元格是否被合并
            if hasattr(cell, '_tc'):
                tc = cell._tc
                if hasattr(tc, 'grid_span'):
                    col_span = tc.grid_span
                    if col_span > 1:
                        is_merged = True
                
                # 检查行合并
                if hasattr(tc, 'v_merge'):
                    v_merge = tc.v_merge
                    if v_merge is not None:
                        is_merged = True
                        # 尝试计算行合并数量
                        row_span = self._calculate_row_span(table, row_idx, col_idx)
                
                merge_span = (row_span, col_span)
            
            # 判断单元格类型
            cell_type = "normal"
            if row_idx == 0:  # 第一行通常是标题
                cell_type = "header"
            elif not cell_text:
                cell_type = "empty"
            else:
                cell_type = "data"
            
            # 提取样式信息
            style_info = self._extract_style_info(cell)
            
            # 提取大小信息
            size_info = self._extract_size_info(cell, table, row_idx, col_idx)
            
            # 提取位置信息
            position_info = self._extract_position_info(table, row_idx, col_idx)
            
            # 提取格式信息
            format_info = self._extract_format_info(cell)
            
            return CellInfo(
                row_index=row_idx,
                col_index=col_idx,
                text=cell_text,
                is_merged=is_merged,
                merge_span=merge_span,
                cell_type=cell_type,
                style_info=style_info,
                size_info=size_info,
                position_info=position_info,
                format_info=format_info
            )
            
        except Exception as e:
            logger.warning(f"提取单元格信息失败 ({row_idx}, {col_idx}): {e}")
            return CellInfo(
                row_index=row_idx,
                col_index=col_idx,
                text="",
                cell_type="error"
            )
    
    def _analyze_table_structure(self, table: Table) -> Tuple[int, int, bool]:
        """分析表格结构"""
        rows = len(table.rows)
        columns = len(table.columns)
        has_merged_cells = False
        
        # 检查是否有合并单元格
        for row in table.rows:
            for cell in row.cells:
                if hasattr(cell, '_tc'):
                    tc = cell._tc
                    if hasattr(tc, 'grid_span') and tc.grid_span > 1:
                        has_merged_cells = True
                        break
                if has_merged_cells:
                    break
        
        return rows, columns, has_merged_cells
    
    def _calculate_row_span(self, table: Table, row_idx: int, col_idx: int) -> int:
        """计算行合并数量"""
        try:
            # 这是一个简化的实现，实际的行合并检测比较复杂
            # 这里返回默认值1，实际应用中可以使用更复杂的算法
            return 1
        except Exception:
            return 1
    
    def _extract_style_info(self, cell) -> Dict[str, Any]:
        """提取单元格样式信息"""
        try:
            style_info = {
                "has_border": True,  # 默认有边框
                "alignment": "left",  # 默认左对齐
                "background_color": None,
                "border_style": "solid"
            }
            
            # 尝试提取更多样式信息
            if hasattr(cell, '_tc'):
                tc = cell._tc
                
                # 提取对齐方式
                try:
                    for p in tc.iterchildren():
                        if hasattr(p, 'pPr'):
                            p_pr = p.pPr
                            if hasattr(p_pr, 'jc'):
                                alignment = str(p_pr.jc.val)
                                style_info["alignment"] = alignment
                except:
                    pass
            
            return style_info
        except Exception:
            return {"has_border": True, "alignment": "left"}
    
    def _extract_size_info(self, cell, table: Table, row_idx: int, col_idx: int) -> Dict[str, Any]:
        """提取单元格大小信息"""
        try:
            size_info = {
                "width": None,
                "height": None,
                "width_units": "twips",  # Word中的默认单位
                "height_units": "twips"
            }
            
            # 尝试从表格列获取宽度信息
            try:
                if col_idx < len(table.columns):
                    col = table.columns[col_idx]
                    if hasattr(col, '_gridCol') and hasattr(col._gridCol, 'w'):
                        width = col._gridCol.w
                        size_info["width"] = int(width)
            except:
                pass
            
            # 尝试从表格行获取高度信息
            try:
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    if hasattr(row, '_tr') and hasattr(row._tr, 'trPr'):
                        tr_pr = row._tr.trPr
                        if hasattr(tr_pr, 'trHeight'):
                            height = tr_pr.trHeight
                            if hasattr(height, 'val'):
                                size_info["height"] = int(height.val)
            except:
                pass
            
            return size_info
        except Exception:
            return {"width": None, "height": None}
    
    def _extract_position_info(self, table: Table, row_idx: int, col_idx: int) -> Dict[str, Any]:
        """提取单元格位置信息"""
        try:
            position_info = {
                "row_position": row_idx,
                "col_position": col_idx,
                "is_first_row": row_idx == 0,
                "is_last_row": row_idx == len(table.rows) - 1,
                "is_first_col": col_idx == 0,
                "is_last_col": col_idx == len(table.columns) - 1,
                "total_rows": len(table.rows),
                "total_cols": len(table.columns)
            }
            return position_info
        except Exception:
            return {
                "row_position": row_idx,
                "col_position": col_idx,
                "total_rows": 0,
                "total_cols": 0
            }
    
    def _extract_format_info(self, cell) -> Dict[str, Any]:
        """提取单元格格式信息"""
        try:
            format_info = {
                "font_name": None,
                "font_size": None,
                "font_bold": False,
                "font_italic": False,
                "font_underline": False,
                "text_color": None,
                "background_color": None,
                "paragraph_count": 0,
                "text_length": len(cell.text)
            }
            
            # 提取段落信息
            try:
                paragraphs = cell.paragraphs
                format_info["paragraph_count"] = len(paragraphs)
                
                # 从第一个段落提取格式信息
                if paragraphs:
                    para = paragraphs[0]
                    if para.runs:
                        run = para.runs[0]
                        if run.font:
                            font = run.font
                            if font.name:
                                format_info["font_name"] = font.name
                            if font.size:
                                format_info["font_size"] = font.size.pt
                            format_info["font_bold"] = font.bold is True
                            format_info["font_italic"] = font.italic is True
                            format_info["font_underline"] = font.underline is True
                            if font.color and font.color.rgb:
                                format_info["text_color"] = str(font.color.rgb)
            except:
                pass
            
            return format_info
        except Exception:
            return {
                "font_name": None,
                "font_size": None,
                "text_length": len(cell.text) if cell.text else 0,
                "paragraph_count": 0
            }
    
    def extract_table_structure(self, file_path: str, table_index: int) -> Optional[TableStructure]:
        """
        提取指定表格的结构
        
        Args:
            file_path: 文档路径
            table_index: 表格索引
            
        Returns:
            表格结构信息，如果提取失败返回None
        """
        try:
            # 检查缓存
            cached_structure = self.cache.get(file_path, table_index)
            if cached_structure:
                logger.info(f"从缓存获取表格 {table_index} 结构")
                return cached_structure
            
            # 打开文档
            doc = Document(file_path)
            
            # 检查表格索引是否有效
            if table_index < 0 or table_index >= len(doc.tables):
                logger.error(f"表格索引超出范围: {table_index}")
                return None
            
            table = doc.tables[table_index]
            
            # 检测页面格式
            page_format = self._detect_page_format(doc)
            
            # 检测表格类型
            table_type = self._detect_table_type(table, page_format)
            
            # 分析表格结构
            rows, columns, has_merged_cells = self._analyze_table_structure(table)
            
            # 提取所有单元格信息
            cells = []
            for row_idx in range(rows):
                row_cells = []
                for col_idx in range(columns):
                    cell_info = self._extract_cell_info(table, row_idx, col_idx)
                    row_cells.append(cell_info)
                cells.append(row_cells)
            
            # 计算标题行和数据行数量
            header_rows = 0
            data_rows = 0
            
            for row_idx, row_cells in enumerate(cells):
                has_header_content = any(
                    cell.cell_type == "header" or 
                    any(keyword in cell.text.lower() for keyword in ["项目", "名称", "时间", "等级", "分值"])
                    for cell in row_cells
                )
                
                if has_header_content and row_idx < 2:  # 前两行可能是标题
                    header_rows += 1
                elif any(cell.cell_type == "data" for cell in row_cells):
                    data_rows += 1
            
            # 创建表格结构对象
            structure = TableStructure(
                table_index=table_index,
                rows=rows,
                columns=columns,
                cells=cells,
                table_type=table_type,
                page_format=page_format,
                has_merged_cells=has_merged_cells,
                header_rows=header_rows,
                data_rows=data_rows,
                extracted_at=time.time()
            )
            
            # 存储到缓存
            self.cache.put(file_path, table_index, structure)
            
            logger.info(f"成功提取表格 {table_index} 结构: {rows}行 x {columns}列, 类型: {table_type}, 格式: {page_format}")
            return structure
            
        except Exception as e:
            logger.error(f"提取表格结构失败: {e}")
            return None
    
    def extract_all_tables(self, file_path: str) -> List[TableStructure]:
        """
        提取文档中所有表格的结构
        
        Args:
            file_path: 文档路径
            
        Returns:
            所有表格结构信息列表
        """
        try:
            doc = Document(file_path)
            table_count = len(doc.tables)
            
            if table_count == 0:
                logger.warning(f"文档中没有表格: {file_path}")
                return []
            
            structures = []
            for i in range(table_count):
                structure = self.extract_table_structure(file_path, i)
                if structure:
                    structures.append(structure)
            
            logger.info(f"成功提取 {len(structures)} 个表格结构")
            return structures
            
        except Exception as e:
            logger.error(f"提取所有表格结构失败: {e}")
            return []
    
    def get_table_summary(self, file_path: str) -> Dict[str, Any]:
        """
        获取文档表格摘要信息
        
        Args:
            file_path: 文档路径
            
        Returns:
            表格摘要信息
        """
        try:
            structures = self.extract_all_tables(file_path)
            
            summary = {
                "file_path": file_path,
                "total_tables": len(structures),
                "page_format": structures[0].page_format if structures else "unknown",
                "tables": []
            }
            
            for structure in structures:
                table_info = {
                    "index": structure.table_index,
                    "rows": structure.rows,
                    "columns": structure.columns,
                    "type": structure.table_type,
                    "has_merged_cells": structure.has_merged_cells,
                    "header_rows": structure.header_rows,
                    "data_rows": structure.data_rows
                }
                summary["tables"].append(table_info)
            
            return summary
            
        except Exception as e:
            logger.error(f"获取表格摘要失败: {e}")
            return {"error": str(e)}
    
    def clear_cache(self):
        """清空缓存"""
        self.cache.clear()
    
    def get_cache_info(self) -> Dict[str, Any]:
        """获取缓存信息"""
        return self.cache.get_cache_info()

# 全局实例
table_extractor = TableStructureExtractor()
