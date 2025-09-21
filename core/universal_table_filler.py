#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用表格填充器 - 基于智能分析的表格填充执行器
支持任意格式的Word表格智能填充
"""

import json
import logging
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path
from docx import Document
from .intelligent_table_analyzer import IntelligentTableAnalyzer

logger = logging.getLogger(__name__)

class UniversalTableFiller:
    """通用表格填充器"""
    
    def __init__(self):
        self.analyzer = IntelligentTableAnalyzer()
    
    def analyze_and_get_coordinates(self, file_path: str) -> str:
        """
        分析文档并返回坐标信息（重点功能）
        
        Args:
            file_path: 文档路径
            
        Returns:
            坐标映射信息，供AI使用
        """
        try:
            # 分析文档结构
            logger.info(f"分析文档结构: {file_path}")
            analysis_result = self.analyzer.analyze_document(file_path)
            
            if 'error' in analysis_result:
                return f"文档分析失败: {analysis_result['error']}"
            
            # 创建坐标映射
            logger.info("创建坐标映射信息")
            coordinate_mapping = self.analyzer.create_coordinate_mapping(analysis_result)
            
            return json.dumps(coordinate_mapping, ensure_ascii=False, indent=2)
            
        except Exception as e:
            logger.error(f"坐标分析失败: {e}")
            return f"分析失败: {str(e)}"
    
    def get_document_analysis(self, file_path: str) -> str:
        """
        获取文档分析结果（专注于坐标信息）
        
        Args:
            file_path: 文档路径
            
        Returns:
            JSON格式的坐标分析结果
        """
        try:
            # 直接调用坐标分析方法
            return self.analyze_and_get_coordinates(file_path)
            
        except Exception as e:
            logger.error(f"文档分析失败: {e}")
            return f"分析失败: {str(e)}"
    
    def fill_with_coordinates(self, file_path: str, 
                            coordinate_data: Dict[str, Tuple[int, int, int]]) -> str:
        """
        使用坐标数据直接填充表格
        
        Args:
            file_path: 文档路径
            coordinate_data: 坐标数据，格式为 {数据内容: (表格索引, 行索引, 列索引)}
            
        Returns:
            填充结果
        """
        try:
            result = self._execute_fill_plan(file_path, coordinate_data)
            return self._format_coordinate_result(coordinate_data, result)
            
        except Exception as e:
            logger.error(f"坐标填充失败: {e}")
            return f"填充失败: {str(e)}"
    
    def _execute_fill_plan(self, file_path: str, 
                          fill_plan: Dict[str, Tuple[int, int, int]]) -> Dict[str, Any]:
        """
        执行填充计划
        
        Args:
            file_path: 文档路径
            fill_plan: 填充计划
            
        Returns:
            执行结果
        """
        try:
            # 打开文档
            doc = Document(file_path)
            
            filled_count = 0
            filled_positions = []
            failed_fills = []
            
            # 执行每个填充操作
            #（表，行，列）
            for data_value, (table_idx, row_idx, col_idx) in fill_plan.items():
                try:
                    # 验证表格索引
                    if table_idx >= len(doc.tables):
                        failed_fills.append({
                            'data': data_value,
                            'position': (table_idx, row_idx, col_idx),
                            'error': f'表格索引超出范围: {table_idx}'
                        })
                        continue
                    
                    table = doc.tables[table_idx]
                    
                    # 验证行列索引
                    if row_idx >= len(table.rows):
                        failed_fills.append({
                            'data': data_value,
                            'position': (table_idx, row_idx, col_idx),
                            'error': f'行索引超出范围: {row_idx}'
                        })
                        continue
                    
                    if col_idx >= len(table.rows[row_idx].cells):
                        failed_fills.append({
                            'data': data_value,
                            'position': (table_idx, row_idx, col_idx),
                            'error': f'列索引超出范围: {col_idx}'
                        })
                        continue
                    
                    # 执行填充
                    cell = table.cell(row_idx, col_idx)
                    original_text = cell.text.strip()
                    
                    # 检查是否已有内容
                    if original_text and original_text != data_value:
                        logger.warning(f"位置 ({table_idx}, {row_idx}, {col_idx}) 已有内容: '{original_text}'，将被覆盖为: '{data_value}'")
                    
                    # 填充数据
                    cell.text = str(data_value)
                    filled_count += 1
                    filled_positions.append({
                        'data': data_value,
                        'position': (table_idx, row_idx, col_idx),
                        'original_text': original_text
                    })
                    
                    logger.info(f"成功填充 '{data_value}' 到位置 ({table_idx}, {row_idx}, {col_idx})")
                    
                except Exception as e:
                    failed_fills.append({
                        'data': data_value,
                        'position': (table_idx, row_idx, col_idx),
                        'error': str(e)
                    })
                    logger.error(f"填充失败: '{data_value}' -> ({table_idx}, {row_idx}, {col_idx}): {e}")
            
            # 保存文档
            doc.save(file_path)
            
            return {
                'success': True,
                'filled_count': filled_count,
                'total_planned': len(fill_plan),
                'filled_positions': filled_positions,
                'failed_fills': failed_fills,
                'success_rate': filled_count / len(fill_plan) if fill_plan else 0
            }
            
        except Exception as e:
            logger.error(f"执行填充计划失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'filled_count': 0,
                'failed_fills': []
            }
    
    def _format_analysis_for_ai(self, analysis_result: Dict[str, Any]) -> Dict[str, Any]:
        """
        格式化分析结果为AI友好的格式
        
        Args:
            analysis_result: 原始分析结果
            
        Returns:
            AI友好的格式化结果
        """
        # 提取关键信息
        field_positions = analysis_result.get('field_positions', {})
        empty_positions = analysis_result.get('empty_positions', [])
        fill_rules = analysis_result.get('fill_rules', [])
        
        # 创建简化的字段映射（使用您要求的格式）
        simplified_field_map = {}
        for field_name, position in field_positions.items():
            simplified_field_map[field_name] = position
        
        # 创建空位映射
        empty_map = {}
        for i, position in enumerate(empty_positions):
            empty_map[f'empty_{i+1}'] = position
        
        # 创建填充规则摘要
        rules_summary = []
        for rule in fill_rules:
            rules_summary.append({
                'field': rule['field'],
                'fill_position': rule['fill_position'],
                'rule_type': rule['rule_type'],
                'confidence': rule['confidence']
            })
        
        return {
            'document_info': analysis_result.get('document_info', {}),
            'field_positions': simplified_field_map,
            'empty_positions': empty_map,
            'fill_rules': rules_summary,
            'ai_instructions': {
                'usage_notes': [
                    '1. 有字段内容的格子不可以填',
                    '2. 优先使用右侧填充规则',
                    '3. 字段名下方的格子也可能填充数据',
                    '4. 字段名右侧的格子通常填充该字段的数据',
                    '5. 返回格式: {"数据内容": (表格索引, 行索引, 列索引)}'
                ],
                'example_format': {
                    '张三': (1, 1, 3),
                    '2023001234': (1, 2, 3),
                    '计算机学院': (1, 1, 5)
                }
            },
            'total_tables': analysis_result.get('document_info', {}).get('total_tables', 0),
            'total_fields': len(field_positions),
            'total_empty_positions': len(empty_positions),
            'total_fill_rules': len(fill_rules)
        }
    
    def _format_result(self, analysis_result: Dict[str, Any], 
                      fill_plan: Dict[str, Tuple[int, int, int]], 
                      execution_result: Dict[str, Any]) -> str:
        """
        格式化最终结果
        
        Args:
            analysis_result: 分析结果
            fill_plan: 填充计划
            execution_result: 执行结果
            
        Returns:
            格式化的结果字符串
        """
        if not execution_result.get('success', False):
            return f"填充失败: {execution_result.get('error', '未知错误')}"
        
        filled_count = execution_result.get('filled_count', 0)
        total_planned = execution_result.get('total_planned', 0)
        success_rate = execution_result.get('success_rate', 0)
        
        result_parts = [
            f"✅ 智能表格填充完成",
            f"📊 填充统计: {filled_count}/{total_planned} 个字段成功填充 (成功率: {success_rate:.1%})"
        ]
        
        # 添加成功填充的详细信息
        if execution_result.get('filled_positions'):
            result_parts.append("\n🎯 成功填充的字段:")
            for item in execution_result['filled_positions']:
                data = item['data']
                pos = item['position']
                result_parts.append(f"   • '{data}' → 表格{pos[0]}, 行{pos[1]}, 列{pos[2]}")
        
        # 添加失败填充的信息
        if execution_result.get('failed_fills'):
            result_parts.append(f"\n❌ 填充失败的字段 ({len(execution_result['failed_fills'])}个):")
            for item in execution_result['failed_fills']:
                data = item['data']
                pos = item['position']
                error = item['error']
                result_parts.append(f"   • '{data}' → 位置{pos}: {error}")
        
        return "\n".join(result_parts)
    
    def _format_coordinate_result(self, coordinate_data: Dict[str, Tuple[int, int, int]], 
                                execution_result: Dict[str, Any]) -> str:
        """
        格式化坐标填充结果
        
        Args:
            coordinate_data: 坐标数据
            execution_result: 执行结果
            
        Returns:
            格式化的结果字符串
        """
        if not execution_result.get('success', False):
            return f"坐标填充失败: {execution_result.get('error', '未知错误')}"
        
        filled_count = execution_result.get('filled_count', 0)
        total_planned = len(coordinate_data)
        success_rate = execution_result.get('success_rate', 0)
        
        result_parts = [
            f"✅ 坐标表格填充完成",
            f"📊 填充统计: {filled_count}/{total_planned} 个位置成功填充 (成功率: {success_rate:.1%})"
        ]
        
        # 添加填充详情
        if execution_result.get('filled_positions'):
            result_parts.append("\n🎯 填充详情:")
            for item in execution_result['filled_positions']:
                data = item['data']
                pos = item['position']
                original = item.get('original_text', '')
                if original:
                    result_parts.append(f"   • '{data}' → 表格{pos[0]}, 行{pos[1]}, 列{pos[2]} (原内容: '{original}')")
                else:
                    result_parts.append(f"   • '{data}' → 表格{pos[0]}, 行{pos[1]}, 列{pos[2]}")
        
        return "\n".join(result_parts)
    
    def intelligent_fill(self, file_path: str, fill_data: Dict[str, Any]) -> str:
        """
        智能填充 - 基于字段名自动匹配填充
        
        Args:
            file_path: 文档路径
            fill_data: 填充数据字典
            
        Returns:
            填充结果信息
        """
        try:
            # 分析文档结构
            logger.info(f"智能填充文档: {file_path}")
            analysis_result = self.analyzer.analyze_document(file_path)
            
            if 'error' in analysis_result:
                return f"文档分析失败: {analysis_result['error']}"
            
            # 获取字段位置
            field_positions = analysis_result.get('field_positions', {})
            
            if not field_positions:
                return "未找到可填充字段"
            
            # 智能匹配字段
            matched_positions = {}
            unmatched_fields = []
            
            for data_key, data_value in fill_data.items():
                best_match = None
                best_confidence = 0.0
                
                # 尝试匹配字段名
                for field_name, position in field_positions.items():
                    confidence = self._calculate_field_similarity(data_key, field_name)
                    if confidence > best_confidence and confidence > 0.3:  # 最低匹配阈值
                        best_confidence = confidence
                        best_match = position
                
                if best_match:
                    matched_positions[data_value] = best_match
                else:
                    unmatched_fields.append(data_key)
            
            if not matched_positions:
                return f"无法匹配任何字段。未匹配字段: {unmatched_fields}"
            
            # 执行填充
            result = self._execute_fill_plan(file_path, matched_positions)
            
            # 构建返回信息
            filled_count = len(matched_positions)
            result_parts = [f"智能填充完成，共填充 {filled_count} 个字段"]
            
            if unmatched_fields:
                result_parts.append(f"未匹配字段: {', '.join(unmatched_fields)}")
            
            result_parts.append(f"\n填充详情:")
            for data, pos in matched_positions.items():
                result_parts.append(f"  • '{data}' → 表格{pos[0]}, 行{pos[1]}, 列{pos[2]}")
            
            return "\n".join(result_parts)
            
        except Exception as e:
            logger.error(f"智能填充失败: {e}")
            return f"智能填充失败: {str(e)}"
    
    def _calculate_field_similarity(self, field1: str, field2: str) -> float:
        """
        计算两个字段名的相似度
        
        Args:
            field1: 字段名1
            field2: 字段名2
            
        Returns:
            相似度分数 (0-1)
        """
        import difflib
        
        # 清理字段名
        clean_field1 = field1.replace(' ', '').replace('、', '').replace('：', '')
        clean_field2 = field2.replace(' ', '').replace('、', '').replace('：', '')
        
        # 计算相似度
        similarity = difflib.SequenceMatcher(None, clean_field1, clean_field2).ratio()
        
        # 检查是否包含关键词
        keywords = ['姓名', '学号', '学院', '专业', '实习', '时间', '单位']
        for keyword in keywords:
            if keyword in clean_field1 and keyword in clean_field2:
                similarity = max(similarity, 0.8)
        
        return similarity
