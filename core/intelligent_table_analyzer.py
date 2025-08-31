#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能表格分析器 - 通用表格内容提取和字段识别系统
支持自动识别表格中的字段名、空位和填充规则
"""

import json
import logging
import re
from typing import Dict, List, Optional, Any, Tuple, Set
from dataclasses import dataclass, asdict
from pathlib import Path
from docx import Document
from docx.table import Table
import difflib

logger = logging.getLogger(__name__)

@dataclass
class CellPosition:
    """单元格位置信息"""
    table_index: int  # 表格索引
    row_index: int    # 行索引
    col_index: int    # 列索引
    
    def __str__(self):
        return f"(表格{table_index}, 行{row_index}, 列{col_index})"
    
    def to_tuple(self) -> Tuple[int, int, int]:
        """转换为元组格式"""
        return (self.table_index, self.row_index, self.col_index)

@dataclass
class CellInfo:
    """单元格详细信息"""
    position: CellPosition
    text: str
    is_empty: bool = False
    is_field_name: bool = False
    is_fillable: bool = False
    field_type: Optional[str] = None  # 字段类型：姓名、学号、学院等
    confidence: float = 0.0  # 识别置信度
    adjacent_fields: List[str] = None  # 相邻字段信息
    
    def __post_init__(self):
        if self.adjacent_fields is None:
            self.adjacent_fields = []

@dataclass
class TableAnalysisResult:
    """表格分析结果"""
    table_index: int
    rows: int
    columns: int
    cells: List[List[CellInfo]]  # 二维数组，按行列组织
    field_positions: Dict[str, CellPosition]  # 字段名 -> 位置映射
    empty_positions: List[CellPosition]  # 空位列表
    fillable_positions: List[CellPosition]  # 可填充位置列表
    fill_rules: List[Dict[str, Any]]  # 填充规则
    analysis_summary: Dict[str, Any]  # 分析摘要

class IntelligentTableAnalyzer:
    """智能表格分析器"""
    
    def __init__(self):
        # 常见字段名模式
        self.field_patterns = {
            '姓名': ['姓名', '姓  名', '学生姓名', '学生名字', 'name'],
            '学号': ['学号', '学  号', '学生学号', 'student_id', '学籍号'],
            '学院': ['学院', '所在学院', '就读学院', '学院名称', 'department'],
            '专业': ['专业', '专业名称', '所学专业', '专业班别', '专业、班别', 'major'],
            '班级': ['班级', '班别', '班级名称', 'class'],
            '实习单位': ['实习单位', '实习公司', '实习企业', '单位名称', 'internship_company'],
            '实习时间': ['实习时间', '实习期间', '实习日期', '时间', 'internship_period'],
            '指导教师': ['指导教师', '指导老师', '导师', 'supervisor'],
            '联系方式': ['联系方式', '联系电话', '手机号码', 'phone'],
            '地址': ['地址', '家庭地址', '通讯地址', 'address'],
            '成绩': ['成绩', '分数', '得分', 'grade', 'score'],
            '评价': ['评价', '评语', '评价内容', 'evaluation'],
            '日期': ['日期', '时间', '年月日', 'date'],
            '签名': ['签名', '签字', 'signature'],
            '盖章': ['盖章', '印章', 'stamp']
        }
        
        # 填充规则模板
        self.fill_rule_templates = [
            {
                'name': '右侧填充',
                'description': '字段名右侧的单元格填充数据',
                'pattern': 'field_right',
                'condition': 'has_field_on_left'
            },
            {
                'name': '下方填充',
                'description': '字段名下方的单元格填充数据',
                'pattern': 'field_below',
                'condition': 'has_field_above'
            },
            {
                'name': '左侧填充',
                'description': '字段名左侧的单元格填充数据',
                'pattern': 'field_left',
                'condition': 'has_field_on_right'
            },
            {
                'name': '上方填充',
                'description': '字段名上方的单元格填充数据',
                'pattern': 'field_above',
                'condition': 'has_field_below'
            }
        ]
    
    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """
        分析整个文档的所有表格
        
        Args:
            file_path: 文档路径
            
        Returns:
            完整的分析结果，包含所有表格的结构和填充信息
        """
        try:
            doc = Document(file_path)
            results = {
                'file_path': file_path,
                'total_tables': len(doc.tables),
                'tables': {},
                'global_field_map': {},
                'global_empty_positions': [],
                'global_fill_rules': []
            }
            
            # 分析每个表格
            for table_idx in range(len(doc.tables)):
                table_result = self.analyze_table(doc, table_idx)
                results['tables'][f'table_{table_idx}'] = table_result
                
                # 合并全局信息
                for field_name, position in table_result.field_positions.items():
                    results['global_field_map'][field_name] = position.to_tuple()
                
                results['global_empty_positions'].extend([
                    pos.to_tuple() for pos in table_result.empty_positions
                ])
                results['global_fill_rules'].extend(table_result.fill_rules)
            
            # 生成AI友好的JSON格式
            ai_friendly_result = self._generate_ai_friendly_format(results)
            
            return ai_friendly_result
            
        except Exception as e:
            logger.error(f"文档分析失败: {e}")
            return {'error': str(e)}
    
    def analyze_table(self, doc: Document, table_index: int) -> TableAnalysisResult:
        """
        分析单个表格
        
        Args:
            doc: 文档对象
            table_index: 表格索引
            
        Returns:
            表格分析结果
        """
        try:
            table = doc.tables[table_index]
            rows = len(table.rows)
            columns = len(table.columns) if table.rows else 0
            
            # 提取所有单元格信息
            cells = []
            field_positions = {}
            empty_positions = []
            fillable_positions = []
            
            for row_idx in range(rows):
                row_cells = []
                for col_idx in range(columns):
                    position = CellPosition(table_index, row_idx, col_idx)
                    cell = table.cell(row_idx, col_idx)
                    cell_text = cell.text.strip()
                    
                    # 创建单元格信息
                    cell_info = CellInfo(
                        position=position,
                        text=cell_text,
                        is_empty=len(cell_text) == 0
                    )
                    
                    # 分析是否为字段名
                    field_type = self._identify_field_type(cell_text)
                    if field_type:
                        cell_info.is_field_name = True
                        cell_info.field_type = field_type
                        cell_info.confidence = self._calculate_field_confidence(cell_text, field_type)
                        field_positions[cell_text] = position
                    
                    # 分析是否可填充
                    cell_info.is_fillable = self._is_fillable_position(cell_info, cells, row_idx, col_idx)
                    
                    if cell_info.is_empty and cell_info.is_fillable:
                        empty_positions.append(position)
                        fillable_positions.append(position)
                    
                    row_cells.append(cell_info)
                cells.append(row_cells)
            
            # 生成填充规则
            fill_rules = self._generate_fill_rules(cells, field_positions)
            
            # 创建分析摘要
            analysis_summary = self._create_analysis_summary(
                rows, columns, field_positions, empty_positions, fill_rules
            )
            
            return TableAnalysisResult(
                table_index=table_index,
                rows=rows,
                columns=columns,
                cells=cells,
                field_positions=field_positions,
                empty_positions=empty_positions,
                fillable_positions=fillable_positions,
                fill_rules=fill_rules,
                analysis_summary=analysis_summary
            )
            
        except Exception as e:
            logger.error(f"表格{table_index}分析失败: {e}")
            raise
    
    def _identify_field_type(self, text: str) -> Optional[str]:
        """
        识别文本是否为字段名，并返回字段类型
        
        Args:
            text: 单元格文本
            
        Returns:
            字段类型，如果不是字段名则返回None
        """
        if not text or len(text.strip()) == 0:
            return None
        
        # 清理文本
        clean_text = text.strip().replace('\n', '').replace('\r', '')
        
        # 检查是否匹配已知字段模式
        for field_type, patterns in self.field_patterns.items():
            for pattern in patterns:
                # 精确匹配
                if clean_text == pattern:
                    return field_type
                
                # 包含匹配
                if pattern in clean_text or clean_text in pattern:
                    return field_type
        
        # 使用相似度匹配
        for field_type, patterns in self.field_patterns.items():
            for pattern in patterns:
                similarity = difflib.SequenceMatcher(None, clean_text, pattern).ratio()
                if similarity > 0.8:  # 80%相似度阈值
                    return field_type
        
        return None
    
    def _calculate_field_confidence(self, text: str, field_type: str) -> float:
        """
        计算字段识别置信度
        
        Args:
            text: 单元格文本
            field_type: 识别的字段类型
            
        Returns:
            置信度分数 (0.0-1.0)
        """
        clean_text = text.strip().replace('\n', '').replace('\r', '')
        patterns = self.field_patterns.get(field_type, [])
        
        max_confidence = 0.0
        for pattern in patterns:
            if clean_text == pattern:
                return 1.0  # 完全匹配
            elif pattern in clean_text or clean_text in pattern:
                return 0.9  # 包含匹配
            else:
                similarity = difflib.SequenceMatcher(None, clean_text, pattern).ratio()
                max_confidence = max(max_confidence, similarity)
        
        return max_confidence
    
    def _is_fillable_position(self, cell_info: CellInfo, cells: List[List[CellInfo]], 
                            row_idx: int, col_idx: int) -> bool:
        """
        判断位置是否可填充
        
        Args:
            cell_info: 当前单元格信息
            cells: 所有单元格信息 (二维数组: [行][列])
            row_idx: 行索引 (相对于当前表格)
            col_idx: 列索引 (相对于当前表格)
            
        Returns:
            是否可填充
        """
        # 如果单元格是字段名，则不可填充（让AI自己判断是否修改）
        if cell_info.is_field_name:
            return False
        
        # 如果单元格为空，则检查是否有相邻字段来决定是否可填充
        if cell_info.is_empty:
            # 检查相邻单元格是否有字段名
            has_adjacent_field = False
            directions = [
                (0, -1),  # 左 - 优先检查左侧字段
                (0, 1),   # 右
                (-1, 0),  # 上
                (1, 0),   # 下
            ]
            
            for dr, dc in directions:
                adj_row = row_idx + dr
                adj_col = col_idx + dc
                
                # 检查边界 - 使用二维坐标系统
                if (0 <= adj_row < len(cells) and 
                    0 <= adj_col < len(cells[adj_row])):
                    adj_cell = cells[adj_row][adj_col]
                    if adj_cell.is_field_name:
                        has_adjacent_field = True
                        cell_info.adjacent_fields.append(adj_cell.text)
            
            return has_adjacent_field
        
        # 如果单元格有内容但不是字段名，则允许填充（AI可以判断是否修改）
        return True
    
    def _generate_fill_rules(self, cells: List[List[CellInfo]], 
                           field_positions: Dict[str, CellPosition]) -> List[Dict[str, Any]]:
        """
        生成填充规则
        
        Args:
            cells: 所有单元格信息
            field_positions: 字段位置映射
            
        Returns:
            填充规则列表
        """
        rules = []
        used_positions = set()  # 记录已使用的填充位置，避免重复
        
        for field_name, position in field_positions.items():
            row_idx = position.row_index
            col_idx = position.col_index
            
            # 优先检查右侧是否可以填充
            if col_idx + 1 < len(cells[row_idx]):
                right_cell = cells[row_idx][col_idx + 1]
                right_pos = (position.table_index, row_idx, col_idx + 1)
                
                # 如果右侧单元格可填充且未被使用
                if right_cell.is_fillable and right_pos not in used_positions:
                    # 检查右侧单元格是否已有内容
                    has_content = not right_cell.is_empty
                    confidence = 0.9 if not has_content else 0.7  # 有内容时降低置信度
                    
                    rules.append({
                        'field_name': field_name,
                        'field_position': position.to_tuple(),
                        'fill_position': right_pos,
                        'rule_type': 'field_right',
                        'description': f'字段"{field_name}"右侧填充',
                        'confidence': confidence,
                        'has_existing_content': has_content,
                        'ai_guidance': f'右侧单元格{"已有内容" if has_content else "为空"}，{"AI需要判断是否覆盖现有内容" if has_content else "可以直接填充"}'
                    })
                    used_positions.add(right_pos)
                    continue  # 优先使用右侧填充，跳过下方填充
            
            # 如果右侧不可用，检查下方是否可以填充
            if row_idx + 1 < len(cells):
                below_cell = cells[row_idx + 1][col_idx]
                below_pos = (position.table_index, row_idx + 1, col_idx)
                
                if below_cell.is_fillable and below_pos not in used_positions:
                    # 检查下方单元格是否已有内容
                    has_content = not below_cell.is_empty
                    confidence = 0.8 if not has_content else 0.6  # 有内容时降低置信度
                    
                    rules.append({
                        'field_name': field_name,
                        'field_position': position.to_tuple(),
                        'fill_position': below_pos,
                        'rule_type': 'field_below',
                        'description': f'字段"{field_name}"下方填充',
                        'confidence': confidence,
                        'has_existing_content': has_content,
                        'ai_guidance': f'下方单元格{"已有内容" if has_content else "为空"}，{"AI需要判断是否覆盖现有内容" if has_content else "可以直接填充"}'
                    })
                    used_positions.add(below_pos)
        
        return rules
    
    def _create_analysis_summary(self, rows: int, columns: int, 
                               field_positions: Dict[str, CellPosition],
                               empty_positions: List[CellPosition],
                               fill_rules: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        创建分析摘要
        
        Args:
            rows: 行数
            columns: 列数
            field_positions: 字段位置
            empty_positions: 空位
            fill_rules: 填充规则
            
        Returns:
            分析摘要
        """
        return {
            'table_size': f'{rows}行 x {columns}列',
            'field_count': len(field_positions),
            'empty_count': len(empty_positions),
            'fillable_count': len([pos for pos in empty_positions]),
            'rule_count': len(fill_rules),
            'field_types': list(set(self._identify_field_type(field) for field in field_positions.keys())),
            'coverage_ratio': len(fill_rules) / max(len(field_positions), 1)
        }
    
    def _generate_ai_friendly_format(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """
        生成AI友好的JSON格式
        
        Args:
            results: 分析结果
            
        Returns:
            AI友好的格式
        """
        # 创建简化的字段映射
        field_map = {}
        for field_name, position_tuple in results['global_field_map'].items():
            field_map[field_name] = position_tuple
        
        # 创建空位列表
        empty_positions = results['global_empty_positions']
        
        # 创建填充规则摘要
        fill_rules_summary = []
        for table_key, table_data in results['tables'].items():
            for rule in table_data.fill_rules:
                fill_rules_summary.append({
                    'field': rule['field_name'],
                    'fill_position': rule['fill_position'],
                    'rule_type': rule['rule_type'],
                    'confidence': rule['confidence']
                })
        
        return {
            'document_info': {
                'file_path': results['file_path'],
                'total_tables': results['total_tables']
            },
            'field_positions': field_map,
            'empty_positions': empty_positions,
            'fill_rules': fill_rules_summary,
            'ai_instructions': {
                'usage_rules': [
                    '1. 有内容的格子不可以填',
                    '2. 字段名右侧的格子通常填充该字段的数据',
                    '3. 字段名下方的格子也可能填充数据',
                    '4. 优先使用右侧填充规则',
                    '5. 返回格式: {"数据内容": (表格索引, 行索引, 列索引)}'
                ],
                'example_format': {
                    '张三': (1, 1, 3),
                    '2023001234': (1, 2, 3),
                    '计算机学院': (1, 1, 5)
                }
            }
        }

    def create_coordinate_mapping(self, analysis_result: Dict[str, Any]) -> Dict[str, Any]:
        """
        创建坐标映射信息，供AI使用
        
        Args:
            analysis_result: 表格分析结果
            
        Returns:
            坐标映射信息，包含字段位置、空位和填充建议
        """
        field_positions = analysis_result.get('field_positions', {})
        empty_positions = analysis_result.get('empty_positions', [])
        fill_rules = analysis_result.get('fill_rules', [])
        
        # 创建字段到坐标的映射
        field_coordinates = {}
        for field_name, position in field_positions.items():
            field_coordinates[field_name] = position
        
        # 创建空位坐标列表
        empty_coordinates = []
        for i, empty_pos in enumerate(empty_positions):
            empty_coordinates.append({
                'index': i + 1,
                'position': empty_pos,
                'description': f'空位{i + 1}'
            })
        
        # 创建填充建议（增强版，包含AI判断指导）
        fill_suggestions = []
        for rule in fill_rules:
            suggestion = {
                'field_name': rule.get('field_name', rule.get('field', '未知字段')),
                'suggested_position': rule.get('fill_position', []),
                'rule_type': rule.get('rule_type', '未知'),
                'confidence': rule.get('confidence', 0.0)
            }
            
            # 添加AI判断指导信息
            if 'has_existing_content' in rule:
                suggestion['has_existing_content'] = rule['has_existing_content']
            if 'ai_guidance' in rule:
                suggestion['ai_guidance'] = rule['ai_guidance']
            
            fill_suggestions.append(suggestion)
        
        return {
            'field_coordinates': field_coordinates,
            'empty_positions': empty_coordinates,
            'fill_suggestions': fill_suggestions,
            'coordinate_format': '(表格索引, 行索引, 列索引)',
            'usage_instructions': [
                '1. 使用 field_coordinates 查看字段位置',
                '2. 使用 empty_positions 查看可填充的空位',
                '3. 使用 fill_suggestions 获取填充建议',
                '4. 坐标格式: (表格索引, 行索引, 列索引)',
                '5. 返回格式: {"数据内容": [表格索引, 行索引, 列索引]}',
                '6. AI判断指导: 查看 fill_suggestions 中的 ai_guidance 字段',
                '7. 字段保护: 如果单元格是字段名，AI需要谨慎判断是否修改',
                '8. 内容覆盖: 如果单元格有内容但不是字段，AI可以判断是否覆盖'
            ]
        }
    
    def _calculate_field_similarity(self, field1: str, field2: str) -> float:
        """
        计算两个字段名的相似度
        
        Args:
            field1: 第一个字段名
            field2: 第二个字段名
            
        Returns:
            相似度分数 (0.0-1.0)
        """
        if not field1 or not field2:
            return 0.0
        
        # 清理字段名
        f1 = field1.replace('\n', '').replace('\r', '').strip()
        f2 = field2.replace('\n', '').replace('\r', '').strip()
        
        # 完全匹配
        if f1 == f2:
            return 1.0
        
        # 包含匹配
        if f1 in f2 or f2 in f1:
            return 0.9
        
        # 使用difflib计算相似度
        return difflib.SequenceMatcher(None, f1, f2).ratio()
