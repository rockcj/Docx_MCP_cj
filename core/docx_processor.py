import docx
from typing import List, Dict, Any, Union
from io import BytesIO
from docx.document import Document
from docx.table import _Cell, Table
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from .models import DocumentPatch, TableOperation, TableOperationType, CellAlignment, BorderStyle

class DocxProcessor:
    """
    负责处理DOCX文件的核心类。
    功能包括：解析文档结构、为元素生成ID、提取内容与格式、以及应用修改。
    所有操作现在都在内存中进行。
    """

    @staticmethod
    def extract_structure_with_ids(file_stream: BytesIO) -> Dict[str, Any]:
        """
        从内存中的文件流解析DOCX文件，提取其结构，并为每个可编辑元素生成唯一ID。

        :param file_stream: 包含.docx文件内容的BytesIO流。
        :return: 一个代表文档结构的字典。
        """
        document = docx.Document(file_stream)
        structure = {"elements": []}
        element_counter = 0

        # 遍历文档的所有顶级元素（段落和表格）
        for element in document.element.body:
            if element.tag.endswith('p'):  # 如果是段落
                p = docx.text.paragraph.Paragraph(element, document)
                paragraph_id = f"p_{element_counter}"
                structure["elements"].append({
                    "id": paragraph_id,
                    "type": "paragraph",
                    "text": p.text,
                    "style": p.style.name,
                })
                element_counter += 1
            elif element.tag.endswith('tbl'):  # 如果是表格
                table = docx.table.Table(element, document)
                table_id = f"tbl_{element_counter}"
                table_data = {
                    "id": table_id,
                    "type": "table",
                    "rows": [],
                    "row_count": len(table.rows),
                    "column_count": len(table.columns) if table.rows else 0,
                    "style": table.style.name if table.style else None
                }
                for i, row in enumerate(table.rows):
                    row_data = {"cells": []}
                    for j, cell in enumerate(row.cells):
                        cell_id = f"{table_id}_r{i}c{j}"
                        row_data["cells"].append({
                            "id": cell_id,
                            "text": cell.text,
                        })
                    table_data["rows"].append(row_data)
                structure["elements"].append(table_data)
                element_counter += 1
        
        return structure

    @staticmethod
    def apply_patches(original_stream: BytesIO, new_stream: BytesIO, patches: List[DocumentPatch]):
        """
        将一系列修改（补丁）应用到内存中的原始DOCX文件流，并将结果写入新的流。

        :param original_stream: 包含原始.docx文件内容的BytesIO流。
        :param new_stream: 用于写入修改后文件内容的BytesIO流。
        :param patches: 一个包含修改指令的列表。
        """
        document = docx.Document(original_stream)
        patches_dict = {p.element_id: p.new_content for p in patches}
        table_operations = [p for p in patches if p.table_operation is not None]

        # 首先处理表格操作（这些操作可能会改变表格结构）
        for patch in table_operations:
            DocxProcessor._apply_table_operation(document, patch.table_operation)

        element_counter = 0
        
        # 然后处理常规的内容修改
        for element in document.element.body:
            if element.tag.endswith('p'):
                current_id = f"p_{element_counter}"
                if current_id in patches_dict:
                    p = docx.text.paragraph.Paragraph(element, document)
                    DocxProcessor._replace_paragraph_text(p, patches_dict[current_id])
                element_counter += 1
            elif element.tag.endswith('tbl'):
                table = docx.table.Table(element, document)
                table_id = f"tbl_{element_counter}"
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        cell_id = f"{table_id}_r{i}c{j}"
                        if cell_id in patches_dict:
                            # 清空单元格并填充新内容
                            cell.text = str(patches_dict[cell_id])
                element_counter += 1

        document.save(new_stream)

    @staticmethod
    def _replace_paragraph_text(p: docx.text.paragraph.Paragraph, new_text: str):
        """
        替换段落的文本，同时尽量保留原始格式。
        
        该方法采用以下策略来保留格式：
        1. 如果段落只有一个run，直接替换文本内容，完全保留格式
        2. 如果段落有多个runs，尝试保留格式分布
        3. 如果新文本较短，使用第一个run的格式
        4. 如果新文本较长，按比例分配格式

        :param p: 要修改的段落对象。
        :param new_text: 新的文本内容。
        """
        if not p.runs:
            # 如果段落没有runs，创建一个新的run
            p.add_run(new_text)
            return
        
        # 保存原始runs的格式信息
        original_runs_info = []
        original_text = ""
        
        for run in p.runs:
            run_info = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': None,
                'style': None
            }
            
            # 保存字体颜色（如果有）
            if run.font.color.rgb:
                run_info['font_color'] = run.font.color.rgb
            
            original_runs_info.append(run_info)
            original_text += run.text
        
        # 清空现有的runs
        for run in p.runs[:]:
            r = run._r
            r.getparent().remove(r)
        
        # 策略1: 如果原段落只有一个run，直接替换并保留所有格式
        if len(original_runs_info) == 1:
            new_run = p.add_run(new_text)
            DocxProcessor._apply_run_formatting(new_run, original_runs_info[0])
            return
        
        # 策略2: 如果新文本为空，不添加任何内容
        if not new_text.strip():
            return
        
        # 策略3: 如果新文本很短或原文本很短，使用第一个run的格式
        if len(new_text) <= 50 or len(original_text) <= 10:
            new_run = p.add_run(new_text)
            DocxProcessor._apply_run_formatting(new_run, original_runs_info[0])
            return
        
        # 策略4: 尝试按比例分配格式
        DocxProcessor._distribute_formatting_proportionally(p, new_text, original_runs_info, original_text)

    @staticmethod
    def _apply_run_formatting(run, format_info: Dict[str, Any]):
        """
        将格式信息应用到run对象上。
        
        :param run: docx run对象
        :param format_info: 包含格式信息的字典
        """
        try:
            if format_info.get('bold') is not None:
                run.bold = format_info['bold']
            if format_info.get('italic') is not None:
                run.italic = format_info['italic']
            if format_info.get('underline') is not None:
                run.underline = format_info['underline']
            if format_info.get('font_name'):
                run.font.name = format_info['font_name']
            if format_info.get('font_size'):
                run.font.size = format_info['font_size']
            if format_info.get('font_color'):
                run.font.color.rgb = format_info['font_color']
        except Exception:
            # 如果应用格式失败，忽略错误继续执行
            pass

    @staticmethod
    def _distribute_formatting_proportionally(p, new_text: str, original_runs_info: List[Dict], original_text: str):
        """
        按比例分配格式到新文本中。
        
        :param p: 段落对象
        :param new_text: 新文本
        :param original_runs_info: 原始runs的格式信息
        :param original_text: 原始文本
        """
        if not original_text:
            # 如果原文本为空，使用第一个run的格式
            new_run = p.add_run(new_text)
            if original_runs_info:
                DocxProcessor._apply_run_formatting(new_run, original_runs_info[0])
            return
        
        # 计算每个原始run在文本中的位置比例
        current_pos = 0
        run_positions = []
        
        for run_info in original_runs_info:
            run_text = run_info['text']
            start_ratio = current_pos / len(original_text) if len(original_text) > 0 else 0
            end_ratio = (current_pos + len(run_text)) / len(original_text) if len(original_text) > 0 else 1
            
            run_positions.append({
                'start_ratio': start_ratio,
                'end_ratio': end_ratio,
                'format_info': run_info
            })
            current_pos += len(run_text)
        
        # 根据比例将新文本分段并应用相应格式
        new_text_len = len(new_text)
        last_end = 0
        
        for i, pos_info in enumerate(run_positions):
            start_pos = int(pos_info['start_ratio'] * new_text_len)
            end_pos = int(pos_info['end_ratio'] * new_text_len)
            
            # 确保不会超出边界
            start_pos = max(start_pos, last_end)
            end_pos = min(end_pos, new_text_len)
            
            # 如果是最后一个run，确保包含所有剩余文本
            if i == len(run_positions) - 1:
                end_pos = new_text_len
            
            if start_pos < end_pos:
                segment_text = new_text[start_pos:end_pos]
                if segment_text:  # 只有当片段不为空时才创建run
                    new_run = p.add_run(segment_text)
                    DocxProcessor._apply_run_formatting(new_run, pos_info['format_info'])
                    last_end = end_pos
        
        # 如果还有剩余文本，使用最后一个run的格式
        if last_end < new_text_len:
            remaining_text = new_text[last_end:]
            if remaining_text:
                new_run = p.add_run(remaining_text)
                if original_runs_info:
                    DocxProcessor._apply_run_formatting(new_run, original_runs_info[-1])

    @staticmethod
    def _apply_table_operation(document: Document, operation: TableOperation):
        """
        应用表格操作到文档中。

        :param document: docx文档对象
        :param operation: 表格操作指令
        """
        try:
            # 找到目标表格
            table = DocxProcessor._find_table_by_id(document, operation.table_id)
            if table is None:
                raise ValueError(f"找不到表格: {operation.table_id}")

            # 根据操作类型执行相应操作
            if operation.operation_type == TableOperationType.ADD_ROW:
                DocxProcessor._add_table_row(table, operation.row_index, operation.cell_data)
            elif operation.operation_type == TableOperationType.ADD_COLUMN:
                DocxProcessor._add_table_column(table, operation.column_index, operation.cell_data)
            elif operation.operation_type == TableOperationType.DELETE_ROW:
                DocxProcessor._delete_table_row(table, operation.row_index)
            elif operation.operation_type == TableOperationType.DELETE_COLUMN:
                DocxProcessor._delete_table_column(table, operation.column_index)
            elif operation.operation_type == TableOperationType.FORMAT_TABLE:
                DocxProcessor._format_table(table, operation.table_format)
            elif operation.operation_type == TableOperationType.FORMAT_CELL:
                DocxProcessor._format_cell(table, operation.cell_id, operation.cell_format)
            elif operation.operation_type == TableOperationType.MODIFY_CELL:
                DocxProcessor._modify_cell_content(table, operation.cell_id, operation.new_content)
                
        except Exception as e:
            # 记录错误并抛出异常以便调试
            error_msg = f"表格操作失败: {operation.operation_type}, 错误: {str(e)}"
            print(error_msg)
            raise Exception(error_msg)

    @staticmethod
    def _find_table_by_id(document: Document, table_id: str) -> Table:
        """
        根据表格ID查找表格对象。

        :param document: docx文档对象
        :param table_id: 表格ID（如 "tbl_2"）
        :return: 表格对象或None
        """
        try:
            # 从table_id中提取元素索引（包括段落和表格的总计数）
            element_index = int(table_id.split('_')[1])
            
            # 遍历文档元素，计数所有元素（段落和表格）
            element_counter = 0
            for element in document.element.body:
                if element_counter == element_index:
                    if element.tag.endswith('tbl'):
                        return docx.table.Table(element, document)
                    else:
                        return None  # 索引指向的不是表格
                element_counter += 1
            return None
        except (ValueError, IndexError):
            return None

    @staticmethod
    def _add_table_row(table: Table, row_index: int = None, cell_data: List[str] = None):
        """
        向表格添加新行。

        :param table: 表格对象
        :param row_index: 插入位置（None表示在末尾添加）
        :param cell_data: 新行的单元格数据列表
        """
        if row_index is None:
            row_index = len(table.rows)
        
        # 添加新行
        if row_index is None or row_index >= len(table.rows):
            # 在末尾添加
            new_row = table.add_row()
            # 填充数据
            if cell_data:
                for i, data in enumerate(cell_data):
                    if i < len(new_row.cells):
                        new_row.cells[i].text = str(data)
        else:
            # 在指定位置插入（这个逻辑比较复杂，暂时先实现末尾添加）
            new_row = table.add_row()
            # 填充数据
            if cell_data:
                for i, data in enumerate(cell_data):
                    if i < len(new_row.cells):
                        new_row.cells[i].text = str(data)

    @staticmethod
    def _add_table_column(table: Table, column_index: int = None, cell_data: List[str] = None):
        """
        向表格添加新列。

        :param table: 表格对象
        :param column_index: 插入位置（None表示在末尾添加）
        :param cell_data: 新列的单元格数据列表
        """
        if column_index is None:
            column_index = len(table.columns)
        
        # 为每一行添加新单元格
        for row_idx, row in enumerate(table.rows):
            # 创建新单元格
            new_cell_element = OxmlElement('w:tc')
            
            # 如果有模板单元格，复制其属性
            if len(row.cells) > 0:
                template_cell = row.cells[0]._tc
                for child in template_cell:
                    if child.tag.endswith('tcPr'):
                        new_cell_element.append(child)
                        break
            
            # 添加空段落
            p_element = OxmlElement('w:p')
            new_cell_element.append(p_element)
            
            # 插入到指定位置
            if column_index >= len(row.cells):
                row._tr.append(new_cell_element)
            else:
                row._tr.insert(column_index, new_cell_element)
            
            # 填充数据
            if cell_data and row_idx < len(cell_data):
                new_cell = row.cells[column_index]
                new_cell.text = str(cell_data[row_idx])

    @staticmethod
    def _delete_table_row(table: Table, row_index: int):
        """
        删除表格行。

        :param table: 表格对象
        :param row_index: 要删除的行索引
        """
        if 0 <= row_index < len(table.rows):
            row_element = table.rows[row_index]._tr
            row_element.getparent().remove(row_element)

    @staticmethod
    def _delete_table_column(table: Table, column_index: int):
        """
        删除表格列。

        :param table: 表格对象
        :param column_index: 要删除的列索引
        """
        for row in table.rows:
            if 0 <= column_index < len(row.cells):
                cell_element = row.cells[column_index]._tc
                cell_element.getparent().remove(cell_element)

    @staticmethod
    def _format_table(table: Table, table_format):
        """
        格式化表格样式。

        :param table: 表格对象
        :param table_format: 表格格式配置
        """
        if not table_format:
            return
        
        try:
            # 设置表格宽度
            if table_format.width:
                if table_format.width.endswith('%'):
                    # 百分比宽度
                    width_pct = int(table_format.width.rstrip('%'))
                    table.autofit = False
                    # 这里需要更复杂的XML操作来设置百分比宽度
                
            # 设置列宽
            if table_format.column_widths:
                for i, width_str in enumerate(table_format.column_widths):
                    if i < len(table.columns):
                        try:
                            if width_str.endswith('cm'):
                                width_cm = float(width_str.rstrip('cm'))
                                table.columns[i].width = Inches(width_cm / 2.54)
                            elif width_str.endswith('in'):
                                width_in = float(width_str.rstrip('in'))
                                table.columns[i].width = Inches(width_in)
                        except ValueError:
                            continue
            
            # 设置边框样式
            if table_format.border_style and table_format.border_style != BorderStyle.NONE:
                DocxProcessor._set_table_borders(table, table_format.border_style, table_format.border_color)
                
        except Exception as e:
            print(f"表格格式化失败: {str(e)}")

    @staticmethod
    def _format_cell(table: Table, cell_id: str, cell_format):
        """
        格式化单元格样式。

        :param table: 表格对象
        :param cell_id: 单元格ID
        :param cell_format: 单元格格式配置
        """
        if not cell_format:
            return
        
        try:
            # 解析单元格ID获取行列索引
            parts = cell_id.split('_')
            if len(parts) >= 3:
                row_col = parts[2]  # 如 "r0c1"
                row_idx = int(row_col.split('c')[0][1:])
                col_idx = int(row_col.split('c')[1])
                
                if 0 <= row_idx < len(table.rows) and 0 <= col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    DocxProcessor._apply_cell_formatting(cell, cell_format)
                    
        except (ValueError, IndexError) as e:
            print(f"单元格格式化失败: {str(e)}")

    @staticmethod
    def _modify_cell_content(table: Table, cell_id: str, new_content):
        """
        修改单元格内容。

        :param table: 表格对象
        :param cell_id: 单元格ID
        :param new_content: 新内容
        """
        try:
            # 解析单元格ID获取行列索引
            parts = cell_id.split('_')
            if len(parts) >= 3:
                row_col = parts[2]  # 如 "r0c1"
                row_idx = int(row_col.split('c')[0][1:])
                col_idx = int(row_col.split('c')[1])
                
                if 0 <= row_idx < len(table.rows) and 0 <= col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = str(new_content)
                    
        except (ValueError, IndexError) as e:
            print(f"单元格内容修改失败: {str(e)}")

    @staticmethod
    def _apply_cell_formatting(cell, cell_format):
        """
        应用单元格格式。

        :param cell: 单元格对象
        :param cell_format: 单元格格式配置
        """
        try:
            # 获取单元格的第一个段落
            if len(cell.paragraphs) > 0:
                paragraph = cell.paragraphs[0]
                
                # 设置对齐方式
                if cell_format.alignment:
                    if cell_format.alignment == CellAlignment.LEFT:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif cell_format.alignment == CellAlignment.CENTER:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif cell_format.alignment == CellAlignment.RIGHT:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    elif cell_format.alignment == CellAlignment.JUSTIFY:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                
                # 设置字体格式
                if len(paragraph.runs) > 0:
                    run = paragraph.runs[0]
                else:
                    run = paragraph.add_run()
                
                if cell_format.bold is not None:
                    run.bold = cell_format.bold
                if cell_format.italic is not None:
                    run.italic = cell_format.italic
                if cell_format.font_size:
                    run.font.size = Pt(cell_format.font_size)
                if cell_format.font_name:
                    run.font.name = cell_format.font_name
                if cell_format.text_color:
                    color_hex = cell_format.text_color.lstrip('#')
                    if len(color_hex) == 6:
                        r = int(color_hex[0:2], 16)
                        g = int(color_hex[2:4], 16)
                        b = int(color_hex[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
            
            # 设置背景色
            if cell_format.background_color:
                DocxProcessor._set_cell_background_color(cell, cell_format.background_color)
                
        except Exception as e:
            print(f"单元格格式应用失败: {str(e)}")

    @staticmethod
    def _set_cell_background_color(cell, color_hex: str):
        """
        设置单元格背景色。

        :param cell: 单元格对象
        :param color_hex: 十六进制颜色代码
        """
        try:
            color_hex = color_hex.lstrip('#')
            if len(color_hex) == 6:
                # 获取或创建单元格属性
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                
                # 设置背景色
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), color_hex.upper())
                tcPr.append(shd)
                
        except Exception as e:
            print(f"设置单元格背景色失败: {str(e)}")

    @staticmethod
    def _set_table_borders(table: Table, border_style: BorderStyle, border_color: str = None):
        """
        设置表格边框。

        :param table: 表格对象
        :param border_style: 边框样式
        :param border_color: 边框颜色
        """
        try:
            # 边框样式映射
            style_map = {
                BorderStyle.SINGLE: 'single',
                BorderStyle.DOUBLE: 'double',
                BorderStyle.THICK: 'thick',
                BorderStyle.THIN: 'single'
            }
            
            border_xml_style = style_map.get(border_style, 'single')
            color = border_color.lstrip('#') if border_color else '000000'
            
            # 创建边框XML
            borders_xml = f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{border_xml_style}" w:sz="4" w:space="0" w:color="{color}"/>
                <w:left w:val="{border_xml_style}" w:sz="4" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{border_xml_style}" w:sz="4" w:space="0" w:color="{color}"/>
                <w:right w:val="{border_xml_style}" w:sz="4" w:space="0" w:color="{color}"/>
                <w:insideH w:val="{border_xml_style}" w:sz="4" w:space="0" w:color="{color}"/>
                <w:insideV w:val="{border_xml_style}" w:sz="4" w:space="0" w:color="{color}"/>
            </w:tblBorders>
            '''
            
            # 获取或创建表格属性
            tbl = table._tbl
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # 移除现有边框设置
            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)
            
            # 添加新边框设置
            borders_element = parse_xml(borders_xml)
            tblPr.append(borders_element)
            
        except Exception as e:
            print(f"设置表格边框失败: {str(e)}") 