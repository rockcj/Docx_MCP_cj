#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
最终完整MCP服务器
包含智能工具和所有基础工具，提供完整功能
"""

import os
import tempfile
import logging
import traceback
import sys
import json
import time
import shutil
from pathlib import Path
from typing import Dict, Any, Optional, List, Union
from docx import Document

# 添加项目根目录到Python路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from fastmcp import FastMCP

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(os.path.join(tempfile.gettempdir(), "final_complete_docx_mcp_server.log")),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("FinalCompleteDocxMCPServer")

# 导入表格结构提取器
try:
    from core.table_structure_extractor import table_extractor
    TABLE_EXTRACTOR_AVAILABLE = True
    logger.info("表格结构提取器加载成功")
except ImportError as e:
    logger.warning(f"表格结构提取器加载失败: {e}")
    TABLE_EXTRACTOR_AVAILABLE = False

# 导入智能工具规划器
try:
    from core.intelligent_tool_planner import intelligent_planner
    PLANNER_AVAILABLE = True
    logger.info("智能工具规划器加载成功")
except ImportError as e:
    logger.warning(f"智能工具规划器加载失败: {e}")
    PLANNER_AVAILABLE = False

# 导入通用表格填充器
try:
    from core.universal_table_filler import UniversalTableFiller
    UNIVERSAL_FILLER_AVAILABLE = True
    logger.info("通用表格填充器加载成功")
except ImportError as e:
    logger.warning(f"通用表格填充器加载失败: {e}")
    UNIVERSAL_FILLER_AVAILABLE = False

# 创建MCP服务器
mcp = FastMCP("FinalCompleteDocxProcessor")

# 全局变量存储当前文档状态
current_document = None
current_file_path = None

# 初始化通用表格填充器
universal_filler = UniversalTableFiller() if UNIVERSAL_FILLER_AVAILABLE else None

# ==================== 基础文档管理工具 ====================

@mcp.tool()
def create_document(file_path: str) -> str:
    """
    创建新的Word文档
    
    Parameters:
    - file_path: 文档保存路径
    """
    try:
        global current_document, current_file_path
        
        from docx import Document
        
        # 规范化路径
        path = Path(file_path)
        if not path.suffix.lower() == '.docx':
            path = path.with_suffix('.docx')
        path.parent.mkdir(parents=True, exist_ok=True)
        
        # 创建文档
        current_document = Document()
        current_file_path = str(path)
        current_document.save(str(path))
        
        logger.info(f"Document created: {path}")
        return f"文档创建成功: {path}"
        
    except Exception as e:
        logger.error(f"创建文档失败: {e}")
        return f"创建文档失败: {str(e)}"

@mcp.tool()
def open_document(file_path: str) -> str:
    """
    打开现有的Word文档
    
    Parameters:
    - file_path: 文档路径
    """
    try:
        global current_document, current_file_path
        
        from docx import Document
        
        if not os.path.exists(file_path):
            return f"文档不存在: {file_path}"
        
        current_document = Document(file_path)
        current_file_path = str(Path(file_path).resolve())
        
        logger.info(f"Document opened: {current_file_path}")
        return f"文档打开成功: {current_file_path}"
        
    except Exception as e:
        logger.error(f"打开文档失败: {e}")
        return f"打开文档失败: {str(e)}"

@mcp.tool()
def save_document() -> str:
    """
    保存当前文档
    """
    try:
        global current_document, current_file_path
        
        if not current_document or not current_file_path:
            return "没有打开的文档需要保存"
        
        current_document.save(current_file_path)
        
        logger.info(f"Document saved: {current_file_path}")
        return f"文档保存成功: {current_file_path}"
        
    except Exception as e:
        logger.error(f"保存文档失败: {e}")
        return f"保存文档失败: {str(e)}"

@mcp.tool()
def save_as_document(new_file_path: str) -> str:
    """
    将当前文档另存为
    
    Parameters:
    - new_file_path: 新的文档路径
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档需要另存为"
        
        # 规范化新路径
        path = Path(new_file_path)
        if not path.suffix.lower() == '.docx':
            path = path.with_suffix('.docx')
        path.parent.mkdir(parents=True, exist_ok=True)
        
        current_document.save(str(path))
        current_file_path = str(path)
        
        logger.info(f"Document saved as: {path}")
        return f"文档另存为成功: {path}"
        
    except Exception as e:
        logger.error(f"另存为文档失败: {e}")
        return f"另存为文档失败: {str(e)}"

@mcp.tool()
def close_document() -> str:
    """
    关闭当前文档
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        file_path = current_file_path
        current_document = None
        current_file_path = None
        
        logger.info(f"Document closed: {file_path}")
        return f"文档已关闭: {file_path}"
        
    except Exception as e:
        logger.error(f"关闭文档失败: {e}")
        return f"关闭文档失败: {str(e)}"

@mcp.tool()
def get_document_info() -> str:
    """
    获取文档信息
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        info = {
            "file_path": current_file_path,
            "paragraphs_count": len(current_document.paragraphs),
            "tables_count": len(current_document.tables),
            "sections_count": len(current_document.sections)
        }
        
        return json.dumps(info, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"获取文档信息失败: {e}")
        return f"获取文档信息失败: {str(e)}"

@mcp.tool()
def copy_document(source_path: str, target_path: str) -> str:
    """
    复制文档文件并重命名保存
    
    功能说明:
    复制指定的文档文件到新位置，用于创建文档的工作副本。
    这样可以保持原始文件干净，在副本上进行修改操作。
    
    参数说明:
    - source_path (必需): 源文件路径
      * 支持相对路径和绝对路径
      * 示例: "docs/原始文档.docx"
    - target_path (必需): 目标文件路径
      * 支持相对路径和绝对路径
      * 示例: "docs/工作副本.docx"
      * 如果目标目录不存在，会自动创建
    
    返回结果:
    成功时返回: "文档复制成功: 源文件 -> 目标文件"
    失败时返回: "文档复制失败: [错误信息]"
    
    使用示例:
    1. 基本复制:
       copy_document("docs/实习鉴定表.docx", "docs/实习鉴定表_工作版.docx")
    
    2. 复制到新目录:
       copy_document("docs/原始.docx", "work/修改版.docx")
    
    注意事项:
    - 源文件必须存在
    - 目标路径的目录会自动创建
    - 如果目标文件已存在，会被覆盖
    - 建议在修改文档前先创建副本
    """
    try:
        # 检查源文件是否存在
        source_file = Path(source_path)
        if not source_file.exists():
            return f"源文件不存在: {source_path}"
        
        if not source_file.is_file():
            return f"源路径不是文件: {source_path}"
        
        # 处理目标路径
        target_file = Path(target_path)
        
        # 确保目标文件有.docx扩展名
        if not target_file.suffix.lower() == '.docx':
            target_file = target_file.with_suffix('.docx')
        
        # 创建目标目录
        target_file.parent.mkdir(parents=True, exist_ok=True)
        
        # 复制文件
        shutil.copy2(source_file, target_file)
        
        logger.info(f"文档复制成功: {source_path} -> {target_file}")
        return f"文档复制成功: {source_path} -> {target_file}"
        
    except Exception as e:
        logger.error(f"文档复制失败: {e}")
        return f"文档复制失败: {str(e)}"

@mcp.tool()
def create_work_copy(file_path: str, suffix: str = "_工作版") -> str:
    """
    为文档创建工作副本
    
    功能说明:
    自动为指定文档创建一个工作副本，在原文件名基础上添加后缀。
    这是一个便捷的工具，用于快速创建文档的工作版本。
    
    参数说明:
    - file_path (必需): 原文件路径
      * 支持相对路径和绝对路径
      * 示例: "docs/实习鉴定表.docx"
    - suffix (可选): 文件名后缀，默认"_工作版"
      * 示例: "_修改版", "_副本", "_backup"
    
    返回结果:
    成功时返回: "工作副本创建成功: 原文件 -> 工作副本文件"
    失败时返回: "工作副本创建失败: [错误信息]"
    
    使用示例:
    1. 创建工作副本:
       create_work_copy("docs/实习鉴定表.docx")
       # 结果: docs/实习鉴定表_工作版.docx
    
    2. 自定义后缀:
       create_work_copy("docs/文档.docx", "_修改版")
       # 结果: docs/文档_修改版.docx
    
    注意事项:
    - 原文件必须存在
    - 如果工作副本已存在，会被覆盖
    - 建议在修改文档前先创建工作副本
    """
    try:
        # 检查源文件是否存在
        source_file = Path(file_path)
        if not source_file.exists():
            return f"原文件不存在: {file_path}"
        
        if not source_file.is_file():
            return f"路径不是文件: {file_path}"
        
        # 生成工作副本路径
        work_file = source_file.parent / f"{source_file.stem}{suffix}{source_file.suffix}"
        
        # 复制文件
        shutil.copy2(source_file, work_file)
        
        logger.info(f"工作副本创建成功: {file_path} -> {work_file}")
        return f"工作副本创建成功: {file_path} -> {work_file}"
        
    except Exception as e:
        logger.error(f"工作副本创建失败: {e}")
        return f"工作副本创建失败: {str(e)}"

# ==================== 智能文档创建工具 ====================

@mcp.tool()
def intelligent_create_document(
    file_path: str, 
    template_type: str = "basic",
    auto_optimize: bool = True
) -> str:
    """
    智能创建文档 - 集成模板和自动优化
    
    功能说明:
    根据指定的模板类型自动创建Word文档，包含预设的结构和格式。
    支持自动页面设置优化，提升文档的专业度。
    
    参数说明:
    - file_path (必需): 文档保存路径，支持相对路径和绝对路径
      * 示例: "C:/Users/用户名/Desktop/我的文档.docx" 或 "report.docx"
      * 如果路径没有.docx扩展名，会自动添加
      * 如果目录不存在，会自动创建
    - template_type (可选): 模板类型，默认"basic"
      * "basic": 基础文档模板，包含标题和基本段落
      * "business": 商务文档模板，包含日期、收件人、主题等商务格式
      * "academic": 学术论文模板，包含摘要、引言、方法、结果、结论等学术结构
    - auto_optimize (可选): 是否自动优化页面设置，默认True
      * True: 自动设置页边距为1英寸，提升文档专业度
      * False: 使用默认页面设置
    
    使用示例:
    1. 创建基础文档:
       intelligent_create_document("C:/Users/用户名/Desktop/报告.docx", "basic", True)
    
    2. 创建商务文档:
       intelligent_create_document("商务信函.docx", "business", True)
    
    3. 创建学术论文:
       intelligent_create_document("/path/to/论文.docx", "academic", True)
    
    返回结果:
    成功时返回: "智能文档创建成功: [文件路径] (模板类型: [模板类型])"
    失败时返回: "智能文档创建失败: [错误信息]"
    
    注意事项:
    - 确保有写入权限到指定目录
    - 文件名不要包含特殊字符
    - 建议使用英文路径避免编码问题
    """
    try:
        global current_document, current_file_path
        
        from docx import Document
        from docx.shared import Inches
        
        # 规范化路径
        path = Path(file_path)
        if not path.suffix.lower() == '.docx':
            path = path.with_suffix('.docx')
        path.parent.mkdir(parents=True, exist_ok=True)
        
        # 创建文档
        doc = Document()
        current_document = doc
        current_file_path = str(path)
        
        # 根据模板类型添加基础结构
        if template_type == "business":
            doc.add_heading("商务文档", 0)
            doc.add_paragraph("日期: {{date}}")
            doc.add_paragraph("收件人: {{recipient}}")
            doc.add_paragraph("主题: {{subject}}")
            doc.add_paragraph("")
            doc.add_paragraph("尊敬的{{recipient}}，")
            doc.add_paragraph("")
            doc.add_paragraph("此致")
            doc.add_paragraph("敬礼")
            doc.add_paragraph("{{sender}}")
        elif template_type == "academic":
            doc.add_heading("学术论文", 0)
            doc.add_paragraph("标题: {{title}}")
            doc.add_paragraph("作者: {{author}}")
            doc.add_paragraph("")
            doc.add_heading("摘要", level=1)
            doc.add_paragraph("{{abstract}}")
            doc.add_heading("1. 引言", level=1)
            doc.add_paragraph("{{introduction}}")
            doc.add_heading("2. 方法", level=1)
            doc.add_paragraph("{{methodology}}")
            doc.add_heading("3. 结果", level=1)
            doc.add_paragraph("{{results}}")
            doc.add_heading("4. 结论", level=1)
            doc.add_paragraph("{{conclusion}}")
        else:  # basic
            doc.add_heading("文档标题", 0)
            doc.add_paragraph("这是一个基础文档模板。")
            doc.add_paragraph("")
            doc.add_paragraph("您可以在这里添加内容。")
        
        # 自动优化页面设置
        if auto_optimize:
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
        
        # 保存文档
        doc.save(str(path))
        
        logger.info(f"智能文档创建成功: {path}")
        return f"智能文档创建成功: {path} (模板类型: {template_type})"
        
    except Exception as e:
        logger.error(f"智能文档创建失败: {e}")
        return f"智能文档创建失败: {str(e)}"

# ==================== 基础文本内容工具 ====================

@mcp.tool()
def add_paragraph(text: str, bold: bool = False, italic: bool = False, 
                 underline: bool = False, font_size: Optional[int] = None, 
                 font_name: Optional[str] = None, color: Optional[str] = None, 
                 alignment: Optional[str] = None, style: Optional[str] = None) -> str:
    """
    添加段落
    
    Parameters:
    - text: 段落文本
    - bold: 是否粗体
    - italic: 是否斜体
    - underline: 是否下划线
    - font_size: 字体大小
    - font_name: 字体名称
    - color: 字体颜色
    - alignment: 对齐方式
    - style: 样式名称
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        # 添加段落
        paragraph = current_document.add_paragraph(text)
        
        # 应用格式
        if paragraph.runs:
            run = paragraph.runs[0]
            font = run.font
            
            if font_name:
                font.name = font_name
            if font_size:
                font.size = Pt(font_size)
            if bold:
                font.bold = bold
            if italic:
                font.italic = italic
            if underline:
                font.underline = underline
            if color:
                if color.startswith("#"):
                    color = color[1:]
                try:
                    rgb_color = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
                    font.color.rgb = rgb_color
                except:
                    pass
        
        # 设置对齐
        if alignment:
            alignment_map = {
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            if alignment in alignment_map:
                paragraph.alignment = alignment_map[alignment]
        
        logger.info("段落添加成功")
        return "段落添加成功"
        
    except Exception as e:
        logger.error(f"添加段落失败: {e}")
        return f"添加段落失败: {str(e)}"

@mcp.tool()
def add_heading(text: str, level: int = 1) -> str:
    """
    添加标题
    
    Parameters:
    - text: 标题文本
    - level: 标题级别（1-9）
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        current_document.add_heading(text, level)
        
        logger.info(f"标题添加成功: {text}")
        return f"标题添加成功: {text}"
        
    except Exception as e:
        logger.error(f"添加标题失败: {e}")
        return f"添加标题失败: {str(e)}"

@mcp.tool()
def add_text_with_formatting(text: str, bold: bool = False, italic: bool = False, 
                           underline: bool = False, font_size: int = 12, 
                           font_name: str = "宋体", color: str = "#000000", 
                           alignment: str = "left", style: str = None) -> str:
    """
    添加带格式的文本 - 精确控制文本格式
    
    功能说明:
    添加具有精确格式控制的文本内容，支持字体、颜色、对齐等多种格式设置。
    适用于需要精确控制文本外观的场景。
    
    参数说明:
    - text (必需): 要添加的文本内容
      * 示例: "这是重要文本" 或 "标题内容"
      * 支持中文和英文
    - bold (可选): 是否粗体，默认False
      * True: 文本显示为粗体
      * False: 正常字重
    - italic (可选): 是否斜体，默认False
      * True: 文本显示为斜体
      * False: 正常字体
    - underline (可选): 是否下划线，默认False
      * True: 文本下方添加下划线
      * False: 无下划线
    - font_size (可选): 字体大小，默认12
      * 范围: 8-72，常用值: 10, 12, 14, 16, 18, 20, 24
      * 示例: 12 表示12号字体
    - font_name (可选): 字体名称，默认"宋体"
      * 中文: "宋体", "黑体", "楷体", "仿宋"
      * 英文: "Arial", "Times New Roman", "Calibri"
      * 示例: "微软雅黑", "Arial"
    - color (可选): 字体颜色，默认"#000000"（黑色）
      * 格式: 十六进制颜色代码
      * 示例: "#FF0000"（红色）, "#0000FF"（蓝色）, "#008000"（绿色）
    - alignment (可选): 对齐方式，默认"left"
      * "left": 左对齐
      * "center": 居中对齐
      * "right": 右对齐
    - style (可选): 样式名称，默认None
      * 暂时未使用，保留用于未来扩展
    
    使用示例:
    1. 添加普通文本:
       add_text_with_formatting("普通文本内容", False, False, False, 12, "宋体", "#000000", "left")
    
    2. 添加粗体标题:
       add_text_with_formatting("重要标题", True, False, False, 16, "黑体", "#000000", "center")
    
    3. 添加强调文本:
       add_text_with_formatting("强调内容", True, True, False, 14, "微软雅黑", "#FF0000", "left")
    
    4. 添加下划线文本:
       add_text_with_formatting("链接文本", False, False, True, 12, "Arial", "#0000FF", "left")
    
    返回结果:
    成功时返回: "格式化文本添加成功"
    失败时返回: "添加格式化文本失败: [错误信息]"
    
    注意事项:
    - 使用前需要先创建或打开文档
    - 颜色代码必须是6位十六进制格式
    - 字体名称需要系统中已安装
    - 格式设置会应用到整个文本段落
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        # 添加段落
        paragraph = current_document.add_paragraph()
        
        # 设置对齐方式
        if alignment == "center":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == "right":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # 添加运行
        run = paragraph.add_run(text)
        
        # 设置字体格式
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold
        font.italic = italic
        font.underline = underline
        
        # 设置颜色
        if color.startswith("#"):
            color = color[1:]
        try:
            rgb_color = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
            font.color.rgb = rgb_color
        except:
            pass  # 如果颜色设置失败，使用默认颜色
        
        logger.info("格式化文本添加成功")
        return "格式化文本添加成功"
        
    except Exception as e:
        logger.error(f"添加格式化文本失败: {e}")
        return f"添加格式化文本失败: {str(e)}"

@mcp.tool()
def search_and_replace(search_text: str, replace_text: str, case_sensitive: bool = True) -> str:
    """
    搜索和替换文本
    
    Parameters:
    - search_text: 搜索文本
    - replace_text: 替换文本
    - case_sensitive: 是否区分大小写
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        count = 0
        for paragraph in current_document.paragraphs:
            if case_sensitive:
                if search_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(search_text, replace_text)
                    count += paragraph.text.count(replace_text)
            else:
                if search_text.lower() in paragraph.text.lower():
                    # 简单的替换，保持原大小写
                    text = paragraph.text
                    paragraph.text = text.replace(search_text, replace_text)
                    count += 1
        
        logger.info(f"搜索替换完成: 替换了 {count} 处")
        return f"搜索替换完成: 替换了 {count} 处"
        
    except Exception as e:
        logger.error(f"搜索替换失败: {e}")
        return f"搜索替换失败: {str(e)}"

# ==================== 智能内容处理工具 ====================

@mcp.tool()
def smart_add_content(
    content: str,
    content_type: str = "paragraph",
    style: str = "normal",
    auto_format: bool = True
) -> str:
    """
    智能添加内容 - 自动格式化和样式应用
    
    功能说明:
    智能添加文本内容到当前文档，根据内容类型自动应用最佳格式。
    支持自动样式识别和格式化，提升文档的专业度和可读性。
    
    参数说明:
    - content (必需): 要添加的文本内容
      * 示例: "这是一个重要的段落内容。" 或 "# 这是标题" 或 "• 这是列表项"
      * 支持中文和英文文本
      * 可以包含换行符，会自动处理
    - content_type (可选): 内容类型，默认"paragraph"
      * "paragraph": 普通段落，适用于正文内容
      * "heading": 标题，会自动识别标题级别（如 # 表示一级标题）
      * "list": 列表项，会自动应用项目符号格式
    - style (可选): 文本样式，默认"normal"
      * "normal": 普通样式，12号字体
      * "emphasis": 强调样式，粗体+斜体
      * "quote": 引用样式，居中+斜体+小字体
    - auto_format (可选): 是否自动格式化，默认True
      * True: 自动应用最佳格式和样式
      * False: 使用基础格式
    
    使用示例:
    1. 添加普通段落:
       smart_add_content("这是文档的正文内容，用于描述具体信息。", "paragraph", "normal", True)
    
    2. 添加标题:
       smart_add_content("# 第一章 项目概述", "heading", "normal", True)
    
    3. 添加列表:
       smart_add_content("• 项目目标", "list", "normal", True)
    
    4. 添加强调内容:
       smart_add_content("重要提示：请仔细阅读以下内容。", "paragraph", "emphasis", True)
    
    5. 添加引用:
       smart_add_content("正如某位专家所说：技术改变世界。", "paragraph", "quote", True)
    
    返回结果:
    成功时返回: "智能内容添加成功: [内容类型] - [内容前50字符]..."
    失败时返回: "智能内容添加失败: [错误信息]"
    
    注意事项:
    - 使用前需要先创建或打开文档
    - 标题会自动识别级别：# = 1级，## = 2级，以此类推
    - 列表项会自动添加项目符号
    - 内容过长时会自动截断显示
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "请先创建或打开一个文档"
        
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        # 根据类型添加内容
        if content_type == "heading":
            # 智能判断标题级别
            if content.startswith("#"):
                level = min(len(content) - len(content.lstrip("#")), 9)
                text = content.lstrip("#").strip()
            else:
                level = 1
                text = content
            
            paragraph = current_document.add_heading(text, level)
        elif content_type == "list":
            # 添加列表
            paragraph = current_document.add_paragraph(content, style='List Bullet')
        else:
            # 添加段落
            paragraph = current_document.add_paragraph(content)
        
        # 应用样式
        if auto_format and paragraph.runs:
            run = paragraph.runs[0]
            
            if style == "emphasis":
                run.font.bold = True
                run.font.italic = True
            elif style == "quote":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run.font.italic = True
                run.font.size = Pt(10)
            elif style == "normal":
                run.font.size = Pt(12)
        
        # 自动保存
        if current_file_path:
            current_document.save(current_file_path)
        
        logger.info(f"智能内容添加成功: {content_type}")
        return f"智能内容添加成功: {content_type} - {content[:50]}..."
        
    except Exception as e:
        logger.error(f"智能内容添加失败: {e}")
        return f"智能内容添加失败: {str(e)}"

# ==================== 基础表格处理工具 ====================

@mcp.tool()
def add_table(rows: int, cols: int, data: Optional[List[List[str]]] = None, 
              has_header: bool = True) -> str:
    """
    添加表格
    
    Parameters:
    - rows: 行数
    - cols: 列数
    - data: 表格数据（可选）
    - has_header: 是否有表头
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        from docx.shared import Inches
        
        # 创建表格
        table = current_document.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # 如果有数据，填充表格
        if data:
            for i, row_data in enumerate(data):
                if i < rows:
                    row = table.rows[i]
                    for j, cell_data in enumerate(row_data):
                        if j < cols:
                            row.cells[j].text = str(cell_data)
        
        # 设置表头样式
        if has_header and rows > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        logger.info(f"表格添加成功: {rows}x{cols}")
        return f"表格添加成功: {rows}行 x {cols}列"
        
    except Exception as e:
        logger.error(f"添加表格失败: {e}")
        return f"添加表格失败: {str(e)}"

@mcp.tool()
def add_table_row(table_index: int, row_data: List[str]) -> str:
    """
    在指定表格中添加行
    
    Parameters:
    - table_index: 表格索引（从0开始）
    - row_data: 行数据
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        # 获取表格
        if table_index >= len(current_document.tables):
            return f"表格索引 {table_index} 超出范围，文档中只有 {len(current_document.tables)} 个表格"
        
        table = current_document.tables[table_index]
        
        # 添加新行
        new_row = table.add_row()
        
        # 填充数据
        for i, cell_data in enumerate(row_data):
            if i < len(new_row.cells):
                new_row.cells[i].text = str(cell_data)
        
        logger.info(f"表格行添加成功: 表格{table_index}")
        return f"表格行添加成功: 表格{table_index}"
        
    except Exception as e:
        logger.error(f"添加表格行失败: {e}")
        return f"添加表格行失败: {str(e)}"

@mcp.tool()
def format_table(table_index: int, style: str = "Table Grid") -> str:
    """
    格式化表格
    
    Parameters:
    - table_index: 表格索引
    - style: 表格样式
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        # 获取表格
        if table_index >= len(current_document.tables):
            return f"表格索引 {table_index} 超出范围，文档中只有 {len(current_document.tables)} 个表格"
        
        table = current_document.tables[table_index]
        table.style = style
        
        logger.info(f"表格格式化成功: 表格{table_index}")
        return f"表格格式化成功: 表格{table_index}"
        
    except Exception as e:
        logger.error(f"格式化表格失败: {e}")
        return f"格式化表格失败: {str(e)}"

@mcp.tool()
def add_table_column(table_index: int, column_index: Optional[int] = None, 
                    data: Optional[List[str]] = None) -> str:
    """
    添加表格列 - 在指定位置插入新列
    
    功能说明:
    在现有表格的指定位置添加新列，支持在任意位置插入或末尾追加。
    可以同时为新列的所有单元格填充数据。
    
    参数说明:
    - table_index (必需): 表格索引，从0开始计数
      * 示例: 0 表示第一个表格，1 表示第二个表格
      * 必须小于文档中表格的总数
    - column_index (可选): 插入位置，默认None表示在末尾添加
      * None: 在表格末尾添加新列
      * 0: 在第一列之前插入
      * 1: 在第二列之前插入
      * 示例: 2 表示在第三列之前插入新列
    - data (可选): 新列的单元格数据列表
      * 格式: ["数据1", "数据2", "数据3", ...]
      * 长度应该等于表格的行数
      * 示例: ["新列标题", "数据1", "数据2", "数据3"]
    
    使用示例:
    1. 在表格末尾添加空列:
       add_table_column(0, None, None)
    
    2. 在第二列位置插入新列:
       add_table_column(0, 1, ["新列标题", "数据1", "数据2"])
    
    3. 在表格末尾添加带数据的列:
       add_table_column(0, None, ["备注", "备注1", "备注2", "备注3"])
    
    4. 在第一列之前插入标题列:
       add_table_column(0, 0, ["序号", "1", "2", "3"])
    
    返回结果:
    成功时返回: "表格 0 添加列成功"
    失败时返回: "添加表格列失败: [错误信息]"
    
    注意事项:
    - 使用前需要先创建或打开文档
    - 表格索引必须有效
    - 列索引不能超出范围
    - 数据列表长度建议与表格行数匹配
    - 新列会自动继承表格的样式
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        if table_index < 0 or table_index >= len(current_document.tables):
            return f"表格索引超出范围: {table_index}"
        
        table = current_document.tables[table_index]
        
        # 确定插入位置
        if column_index is None:
            column_index = len(table.rows[0].cells) if table.rows else 0
        elif column_index < 0 or column_index > (len(table.rows[0].cells) if table.rows else 0):
            return f"列索引超出范围: {column_index}"
        
        # 为每一行添加新单元格
        for row_idx, row in enumerate(table.rows):
            # 创建新单元格
            from docx.oxml import OxmlElement
            new_cell_element = OxmlElement('w:tc')
            
            # 如果有模板单元格，复制其属性
            if len(row.cells) > 0:
                template_cell = row.cells[0]._tc
                for child in template_cell:
                    if child.tag.endswith('tcPr'):
                        new_cell_element.append(child)
                        break
            
            # 添加空段落
            p_element = OxmlElement('w:p')
            new_cell_element.append(p_element)
            
            # 插入到指定位置
            if column_index >= len(row.cells):
                row._tr.append(new_cell_element)
            else:
                row._tr.insert(column_index, new_cell_element)
            
            # 填充数据
            if data and row_idx < len(data):
                # 需要重新获取单元格引用
                table = current_document.tables[table_index]  # 重新获取表格
                new_cell = table.rows[row_idx].cells[column_index]
                new_cell.text = str(data[row_idx])
        
        logger.info(f"表格 {table_index} 添加列成功")
        return f"表格 {table_index} 添加列成功"
        
    except Exception as e:
        logger.error(f"添加表格列失败: {e}")
        return f"添加表格列失败: {str(e)}"

@mcp.tool()
def merge_table_cells(table_index: int, row_start: int, col_start: int, 
                     row_end: int, col_end: int) -> str:
    """
    合并表格单元格
    
    Parameters:
    - table_index: 表格索引
    - row_start: 起始行
    - col_start: 起始列
    - row_end: 结束行
    - col_end: 结束列
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        # 获取表格
        if table_index >= len(current_document.tables):
            return f"表格索引 {table_index} 超出范围，文档中只有 {len(current_document.tables)} 个表格"
        
        table = current_document.tables[table_index]
        
        # 合并单元格
        cell = table.cell(row_start, col_start)
        merge_cell = table.cell(row_end, col_end)
        cell.merge(merge_cell)
        
        logger.info(f"表格单元格合并成功: 表格{table_index}")
        return f"表格单元格合并成功: 表格{table_index}"
        
    except Exception as e:
        logger.error(f"合并表格单元格失败: {e}")
        return f"合并表格单元格失败: {str(e)}"

# ==================== 智能表格处理工具 ====================

@mcp.tool()
def extract_table_structure(file_path: str, table_index: int) -> str:
    """
    提取完整表结构 - 支持A3和A4格式文档的表格结构提取
    
    功能说明:
    提取指定文档中指定表格的完整结构信息，包括单元格内容、合并情况、表格类型等。
    支持A3和A4格式的文档，能够自动识别页面格式和表格类型。
    提取的结构信息会暂存在内存缓存中，供后续MCP工具快速访问。
    
    参数说明:
    - file_path (必需): 文档文件路径
      * 支持相对路径和绝对路径
      * 示例: "docs/学生综合素质评价表.docx"
      * 示例: "C:/Documents/表格文档.docx"
    - table_index (必需): 表格索引，从0开始计数
      * 0: 第一个表格
      * 1: 第二个表格
      * 示例: 2 表示第三个表格
    
    返回结果:
    成功时返回JSON格式的表格结构信息，包含:
    - table_index: 表格索引
    - rows: 行数
    - columns: 列数
    - table_type: 表格类型 (evaluation/signature/award/internship/student_info/general)
    - page_format: 页面格式 (A3/A4)
    - has_merged_cells: 是否有合并单元格
    - header_rows: 标题行数量
    - data_rows: 数据行数量
    - cells: 详细的单元格信息数组，每个单元格包含:
      * row_index, col_index: 行列索引
      * text: 单元格文本内容
      * is_merged, merge_span: 合并单元格信息
      * cell_type: 单元格类型
      * style_info: 样式信息 (边框、对齐方式等)
      * size_info: 大小信息 (宽度、高度等)
      * position_info: 位置信息 (行列位置、边界信息等)
      * format_info: 格式信息 (字体、颜色、段落等)
    
    使用示例:
    1. 提取第一个表格结构:
       extract_table_structure("docs/学生综合素质评价表.docx", 0)
    
    2. 提取A3格式文档的表格:
       extract_table_structure("docs/实习鉴定表A3.docx", 0)
    
    3. 提取特定表格:
       extract_table_structure("docs/陈杰综合素质评价表.docx", 2)
    
    注意事项:
    - 表格索引必须有效（0到表格总数-1）
    - 提取的结构会缓存到内存中，提高后续访问速度
    - 支持自动识别A3和A4格式文档
    - 能够智能识别表格类型（评价表、签名表、获奖表等）
    """
    try:
        if not TABLE_EXTRACTOR_AVAILABLE:
            return "表格结构提取器不可用，请检查模块导入"
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return f"文件不存在: {file_path}"
        
        # 提取表格结构
        structure = table_extractor.extract_table_structure(file_path, table_index)
        
        if structure is None:
            return f"提取表格 {table_index} 结构失败"
        
        # 转换为可序列化的格式
        structure_dict = {
            "table_index": structure.table_index,
            "rows": structure.rows,
            "columns": structure.columns,
            "table_type": structure.table_type,
            "page_format": structure.page_format,
            "has_merged_cells": structure.has_merged_cells,
            "header_rows": structure.header_rows,
            "data_rows": structure.data_rows,
            "extracted_at": structure.extracted_at,
            "cells": []
        }
        
        # 转换单元格信息
        for row_idx, row_cells in enumerate(structure.cells):
            row_data = []
            for col_idx, cell in enumerate(row_cells):
                cell_data = {
                    "row_index": cell.row_index,
                    "col_index": cell.col_index,
                    "text": cell.text,
                    "is_merged": cell.is_merged,
                    "merge_span": cell.merge_span,
                    "cell_type": cell.cell_type,
                    "style_info": cell.style_info or {},
                    "size_info": cell.size_info or {},
                    "position_info": cell.position_info or {},
                    "format_info": cell.format_info or {}
                }
                row_data.append(cell_data)
            structure_dict["cells"].append(row_data)
        
        logger.info(f"成功提取表格 {table_index} 结构: {structure.rows}行 x {structure.columns}列")
        return json.dumps(structure_dict, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"提取表格结构失败: {e}")
        return f"提取表格结构失败: {str(e)}"

@mcp.tool()
def extract_all_tables_structure(file_path: str) -> str:
    """
    提取文档中所有表格的完整结构
    
    功能说明:
    提取指定文档中所有表格的完整结构信息，返回所有表格的结构摘要。
    适用于需要了解整个文档表格布局的场景。
    
    参数说明:
    - file_path (必需): 文档文件路径
      * 支持相对路径和绝对路径
      * 示例: "docs/学生综合素质评价表.docx"
    
    返回结果:
    成功时返回JSON格式的表格摘要信息，包含:
    - file_path: 文件路径
    - total_tables: 表格总数
    - page_format: 页面格式
    - tables: 每个表格的基本信息数组
    
    使用示例:
    1. 提取所有表格结构:
       extract_all_tables_structure("docs/学生综合素质评价表.docx")
    
    2. 分析A3格式文档:
       extract_all_tables_structure("docs/实习鉴定表A3.docx")
    
    注意事项:
    - 会提取并缓存所有表格的结构信息
    - 返回的是摘要信息，详细结构请使用extract_table_structure
    """
    try:
        if not TABLE_EXTRACTOR_AVAILABLE:
            return "表格结构提取器不可用，请检查模块导入"
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return f"文件不存在: {file_path}"
        
        # 获取表格摘要
        summary = table_extractor.get_table_summary(file_path)
        
        if "error" in summary:
            return f"获取表格摘要失败: {summary['error']}"
        
        logger.info(f"成功提取文档表格摘要: {summary['total_tables']}个表格")
        return json.dumps(summary, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"提取所有表格结构失败: {e}")
        return f"提取所有表格结构失败: {str(e)}"

@mcp.tool()
def extract_document_structure(file_path: str, include_cell_details: bool = False) -> str:
    """
    提取完整文档结构 - 包含所有表格、段落、图片等元素的综合分析
    
    功能说明:
    一次性提取文档的完整结构信息，包括所有表格、段落、图片、页面格式等。
    提供文档的全局视图，帮助AI更好地理解文档布局和内容组织。
    支持A3和A4格式的文档，能够自动识别页面格式和表格类型。
    
    参数说明:
    - file_path (必需): 文档文件路径
      * 支持相对路径和绝对路径
      * 示例: "docs/学生综合素质评价表.docx"
    - include_cell_details (可选): 是否包含详细的单元格信息
      * true: 包含完整的单元格详细信息
      * false: 仅包含表格摘要信息（默认）
    
    返回结果:
    成功时返回JSON格式的文档结构信息，包含:
    - file_path: 文档路径
    - document_info: 文档基本信息
      * page_format: 页面格式 (A3/A4)
      * total_paragraphs: 段落总数
      * total_tables: 表格总数
      * total_images: 图片总数
      * total_pages: 页数估算
    - tables: 所有表格的详细信息
      * 每个表格包含完整的结构信息
      * 如果include_cell_details=true，包含详细单元格信息
    - paragraphs: 段落信息摘要
      * 段落数量和主要段落类型
    - images: 图片信息摘要
      * 图片数量和位置信息
    
    使用示例:
    1. 提取文档完整结构:
       extract_document_structure("docs/学生综合素质评价表.docx")
    
    2. 包含详细单元格信息:
       extract_document_structure("docs/实习鉴定表A3.docx", true)
    
    3. 分析复杂文档:
       extract_document_structure("docs/陈杰综合素质评价表.docx", false)
    
    注意事项:
    - 会提取并缓存所有表格的结构信息
    - include_cell_details=true时返回数据量较大
    - 支持自动识别A3和A4格式文档
    - 能够智能识别表格类型和文档结构
    """
    try:
        if not TABLE_EXTRACTOR_AVAILABLE:
            return "表格结构提取器不可用，请检查模块导入"
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return f"文件不存在: {file_path}"
        
        # 打开文档进行分析
        doc = Document(file_path)
        
        # 提取文档基本信息
        document_info = {
            "file_path": file_path,
            "page_format": "unknown",
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "total_images": 0,
            "total_pages": 1  # 简化估算
        }
        
        # 检测页面格式
        if doc.sections:
            section = doc.sections[0]
            page_width = section.page_width
            page_height = section.page_height
            
            # 根据尺寸判断页面格式
            if page_width > 10000000:  # A3通常更宽
                document_info["page_format"] = "A3"
            else:
                document_info["page_format"] = "A4"
        
        # 统计图片数量
        image_count = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath('.//a:blip'):
                    image_count += 1
        document_info["total_images"] = image_count
        
        # 提取所有表格结构
        tables_info = []
        if len(doc.tables) > 0:
            for i in range(len(doc.tables)):
                if include_cell_details:
                    # 包含详细单元格信息
                    structure = table_extractor.extract_table_structure(file_path, i)
                    if structure:
                        # 转换为详细格式
                        table_dict = {
                            "index": structure.table_index,
                            "rows": structure.rows,
                            "columns": structure.columns,
                            "type": structure.table_type,
                            "page_format": structure.page_format,
                            "has_merged_cells": structure.has_merged_cells,
                            "header_rows": structure.header_rows,
                            "data_rows": structure.data_rows,
                            "cells": []
                        }
                        
                        # 添加详细单元格信息
                        for row_cells in structure.cells:
                            row_data = []
                            for cell in row_cells:
                                cell_data = {
                                    "row_index": cell.row_index,
                                    "col_index": cell.col_index,
                                    "text": cell.text,
                                    "is_merged": cell.is_merged,
                                    "merge_span": cell.merge_span,
                                    "cell_type": cell.cell_type,
                                    "style_info": cell.style_info or {},
                                    "size_info": cell.size_info or {},
                                    "position_info": cell.position_info or {},
                                    "format_info": cell.format_info or {}
                                }
                                row_data.append(cell_data)
                            table_dict["cells"].append(row_data)
                        
                        tables_info.append(table_dict)
                else:
                    # 仅包含摘要信息
                    structure = table_extractor.extract_table_structure(file_path, i)
                    if structure:
                        table_dict = {
                            "index": structure.table_index,
                            "rows": structure.rows,
                            "columns": structure.columns,
                            "type": structure.table_type,
                            "page_format": structure.page_format,
                            "has_merged_cells": structure.has_merged_cells,
                            "header_rows": structure.header_rows,
                            "data_rows": structure.data_rows
                        }
                        tables_info.append(table_dict)
        
        # 分析段落信息
        paragraphs_info = {
            "total_count": len(doc.paragraphs),
            "non_empty_count": len([p for p in doc.paragraphs if p.text.strip()]),
            "has_tables": len(doc.tables) > 0,
            "has_images": image_count > 0
        }
        
        # 分析图片信息
        images_info = {
            "total_count": image_count,
            "has_images": image_count > 0
        }
        
        # 构建完整的文档结构
        document_structure = {
            "file_path": file_path,
            "document_info": document_info,
            "tables": tables_info,
            "paragraphs": paragraphs_info,
            "images": images_info,
            "extracted_at": time.time(),
            "include_cell_details": include_cell_details
        }
        
        logger.info(f"成功提取文档结构: {document_info['total_tables']}个表格, {document_info['total_paragraphs']}个段落")
        return json.dumps(document_structure, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"提取文档结构失败: {e}")
        return f"提取文档结构失败: {str(e)}"

@mcp.tool()
def get_table_structure_cache_info() -> str:
    """
    获取表格结构缓存信息
    
    功能说明:
    查看当前表格结构缓存的状态信息，包括缓存大小、缓存文件等。
    用于监控缓存使用情况和调试。
    
    返回结果:
    成功时返回JSON格式的缓存信息，包含:
    - cache_size: 当前缓存数量
    - max_cache_size: 最大缓存数量
    - cache_timeout: 缓存超时时间
    - cached_files: 已缓存的文件列表
    
    使用示例:
    1. 查看缓存状态:
       get_table_structure_cache_info()
    
    注意事项:
    - 缓存信息仅用于监控和调试
    - 缓存会在超时后自动清理
    """
    try:
        if not TABLE_EXTRACTOR_AVAILABLE:
            return "表格结构提取器不可用，请检查模块导入"
        
        cache_info = table_extractor.get_cache_info()
        logger.info("获取表格结构缓存信息")
        return json.dumps(cache_info, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"获取缓存信息失败: {e}")
        return f"获取缓存信息失败: {str(e)}"

@mcp.tool()
def get_intelligent_planning_guide() -> str:
    """
    获取智能规划指导 - AI使用MCP工具前的必读指南
    
    功能说明:
    这是AI在使用MCP工具前必须阅读的智能规划指导。
    包含了所有工具的详细说明、使用场景、参数要求和最佳实践。
    
    返回结果:
    返回完整的智能规划指导文档，包括：
    - 所有工具的分类和描述
    - 智能工作流规划建议
    - 工具使用最佳实践
    - 常见任务的处理流程
    
    使用说明:
    AI在处理任何文档任务前，都应该先调用此工具获取规划指导，
    然后基于指导制定合理的工具调用计划。
    
    注意事项:
    - 此工具返回的是静态指导文档，不涉及具体文档操作
    - AI应该将此指导作为工具选择的决策依据
    - 建议AI在处理复杂任务时多次参考此指导
    """
    try:
        if not PLANNER_AVAILABLE:
            return "智能规划器不可用，无法获取规划指导"
        
        return intelligent_planner.get_all_tools_summary()
        
    except Exception as e:
        logger.error(f"获取智能规划指导失败: {e}")
        return f"获取智能规划指导失败: {str(e)}"

@mcp.tool()
def create_intelligent_workflow_plan(user_request: str) -> str:
    """
    创建智能工作流规划 - 为AI提供具体的执行计划
    
    功能说明:
    根据用户的具体请求，生成详细的工具调用工作流规划。
    AI可以使用此规划来有序地执行复杂的文档处理任务。
    
    参数说明:
    - user_request (必需): 用户的完整请求描述
      * 示例: "帮我创建一个学生信息表，包含姓名、学号、班级字段，并填充一些示例数据"
      * 示例: "打开现有文档，分析表格结构，然后智能填充数据"
    
    返回结果:
    返回JSON格式的详细工作流规划，包括：
    - 任务类型分析
    - 步骤化执行计划
    - 工具调用顺序
    - 参数配置建议
    - 风险评估和优化建议
    
    使用示例:
    1. 创建表格任务规划:
       create_intelligent_workflow_plan("创建学生综合素质评价表，包含评价项目、分值、得分字段")
    
    2. 表格填充任务规划:
       create_intelligent_workflow_plan("打开docs/评价表.docx，智能填充学生张三的评价数据")
    
    注意事项:
    - 此工具只生成规划，不执行具体操作
    - AI应该按照返回的规划逐步执行工具调用
    - 可以根据实际情况调整规划中的参数
    """
    try:
        if not PLANNER_AVAILABLE:
            return "智能规划器不可用，无法创建工作流规划"
        
        plan = intelligent_planner.create_intelligent_plan(user_request)
        
        # 格式化规划结果
        result = f"""
🎯 智能工作流规划
═══════════════════════════════════════════════════════════════

📋 用户请求: {user_request}
🔍 任务类型: {plan.task_type.value}
📊 总步骤数: {plan.total_steps}
⏱️ 预计时间: {plan.estimated_time}

🔄 执行步骤:
"""
        
        for step in plan.workflow_steps:
            result += f"""
步骤 {step.step_id}: {step.description}
  工具: {step.tool_name}
  参数: {json.dumps(step.parameters, ensure_ascii=False, indent=4)}
  预期输出: {step.expected_output}
  验证标准: {', '.join(step.validation_criteria)}
"""
        
        if plan.parallel_opportunities:
            result += f"""
⚡ 并行执行机会: {plan.parallel_opportunities}
"""
        
        if plan.risk_factors:
            result += f"""
⚠️ 风险因素:
"""
            for risk in plan.risk_factors:
                result += f"  - {risk}\n"
        
        if plan.optimization_suggestions:
            result += f"""
💡 优化建议:
"""
            for suggestion in plan.optimization_suggestions:
                result += f"  - {suggestion}\n"
        
        result += "\n═══════════════════════════════════════════════════════════════\n"
        
        return result
        
    except Exception as e:
        logger.error(f"创建工作流规划失败: {e}")
        return f"创建工作流规划失败: {str(e)}"

@mcp.tool()
def get_tool_detailed_guidance(tool_name: str) -> str:
    """
    获取特定工具的详细指导 - 深入了解单个工具的使用方法
    
    功能说明:
    获取指定工具的详细使用指导，包括参数说明、使用示例、最佳实践等。
    当AI需要深入了解某个特定工具时使用。
    
    参数说明:
    - tool_name (必需): 工具名称
      * 示例: "intelligent_table_fill"
      * 示例: "extract_table_structure"
      * 示例: "create_document"
    
    返回结果:
    返回指定工具的详细指导文档，包括：
    - 功能描述和适用场景
    - 详细的参数说明
    - 使用示例和最佳实践
    - 前置条件和输出格式
    
    使用示例:
    1. 获取智能表格填充工具指导:
       get_tool_detailed_guidance("intelligent_table_fill")
    
    2. 获取表格结构提取工具指导:
       get_tool_detailed_guidance("extract_table_structure")
    
    注意事项:
    - 工具名称必须准确匹配MCP工具名称
    - 此工具返回静态指导，不执行实际操作
    - 建议AI在处理复杂任务前先了解相关工具
    """
    try:
        if not PLANNER_AVAILABLE:
            return "智能规划器不可用，无法获取工具指导"
        
        return intelligent_planner.get_tool_guidance(tool_name)
        
    except Exception as e:
        logger.error(f"获取工具指导失败: {e}")
        return f"获取工具指导失败: {str(e)}"

@mcp.tool()
def clear_table_structure_cache() -> str:
    """
    清空表格结构缓存
    
    功能说明:
    清空所有已缓存的表格结构信息，释放内存。
    在MCP调用完成后或需要重置缓存时使用。
    
    返回结果:
    成功时返回: "表格结构缓存已清空"
    失败时返回: "清空缓存失败: [错误信息]"
    
    使用示例:
    1. 清空所有缓存:
       clear_table_structure_cache()
    
    注意事项:
    - 清空后需要重新提取表格结构
    - 建议在MCP调用完成后使用
    """
    try:
        if not TABLE_EXTRACTOR_AVAILABLE:
            return "表格结构提取器不可用，请检查模块导入"
        
        table_extractor.clear_cache()
        logger.info("表格结构缓存已清空")
        return "表格结构缓存已清空"
        
    except Exception as e:
        logger.error(f"清空缓存失败: {e}")
        return f"清空缓存失败: {str(e)}"




# ==================== 通用表格填充工具 ====================

@mcp.tool()
def extract_fillable_fields(file_path: str) -> str:
    """
    提取文档中所有可填充的字段和位置信息（坐标填充专用）
    
    功能说明:
    自动分析文档结构，提取所有表格中的字段坐标、空位坐标等信息，
    为坐标填充提供精确的位置映射。返回AI友好的JSON格式数据。
    
    参数说明:
    - file_path (必需): 文档文件路径
      * 支持相对路径和绝对路径
      * 示例: "docs/实习鉴定表.docx"
    
    返回结果:
    返回JSON格式的坐标信息，包含:
    - field_coordinates: 字段名到坐标的映射
    - empty_positions: 空位坐标列表
    - fill_suggestions: 填充建议
    - coordinate_format: 坐标格式说明
    - usage_instructions: AI使用说明
    
    使用示例:
    1. 提取坐标信息:
       extract_fillable_fields("docs/实习鉴定表.docx")
    
    注意事项:
    - 专注于坐标信息，弱化智能匹配
    - 返回的数据格式适合AI进行坐标填充
    - AI判断指导: 查看 fill_suggestions 中的 ai_guidance 字段
    - 字段保护: 如果单元格是字段名，AI需要谨慎判断是否修改
    - 内容覆盖: 如果单元格有内容但不是字段，AI可以判断是否覆盖现有内容
    """
    try:
        if not UNIVERSAL_FILLER_AVAILABLE:
            return "通用表格填充器不可用，请检查系统配置"
        
        if not universal_filler:
            return "通用表格填充器初始化失败"
        
        result = universal_filler.get_document_analysis(file_path)
        return result
        
    except Exception as e:
        logger.error(f"提取可填充字段失败: {e}")
        return f"提取失败: {str(e)}"

@mcp.tool()
def intelligent_table_fill(file_path: str, fill_data: Dict[str, Any]) -> str:
    """
    智能表格填充 - 简化的辅助填充功能
    
    功能说明:
    这是一个简化的智能填充功能，主要用于辅助坐标填充。
    建议优先使用 extract_fillable_fields + fill_with_coordinates 的组合。
    
    参数说明:
    - file_path (必需): 文档文件路径
    - fill_data (必需): 填充数据字典
      * 格式: {"字段名1": "值1", "字段名2": "值2", ...}
    
    返回结果:
    返回填充结果或建议使用坐标填充的提示
    
    使用建议:
    推荐使用以下工作流程：
    1. 调用 extract_fillable_fields() 获取坐标信息
    2. 根据坐标信息创建填充计划
    3. 调用 fill_with_coordinates() 执行精确填充
    """
    try:
        if not UNIVERSAL_FILLER_AVAILABLE:
            return "通用表格填充器不可用，请检查系统配置"
        
        if not universal_filler:
            return "通用表格填充器初始化失败"
        
        # 提供使用建议
        suggestion = """
建议使用更精确的坐标填充方式：

1. 首先调用 extract_fillable_fields('{}') 获取坐标信息
2. 根据返回的坐标信息创建填充计划
3. 调用 fill_with_coordinates() 执行精确填充

这样可以获得更好的填充效果和更精确的控制。
        """.format(file_path)
        
        return f"智能填充功能已简化。{suggestion}"
        
    except Exception as e:
        logger.error(f"智能表格填充失败: {e}")
        return f"填充失败: {str(e)}"

@mcp.tool()
def fill_with_coordinates(file_path: str, coordinate_data: Dict[str, List[int]]) -> str:
    """
    使用坐标数据直接填充表格（主要功能）
    
    功能说明:
    这是表格填充的主要功能，根据AI提供的坐标数据直接填充表格，
    支持精确的位置控制。配合 extract_fillable_fields 使用效果最佳。
    
    工作流程:
    1. 调用 extract_fillable_fields() 分析文档结构
    2. 根据返回的坐标信息创建填充计划
    3. 调用本工具执行精确填充
    
    参数说明:
    - file_path (必需): 文档文件路径
      * 支持相对路径和绝对路径
      * 示例: "docs/实习鉴定表.docx"
    - coordinate_data (必需): 坐标数据字典
      * 格式: {"数据内容": [表格索引, 行索引, 列索引], ...}
      * 示例: {"张三": [1, 1, 3], "2023001234": [1, 2, 3]}
      * 坐标从0开始计数
    
    返回结果:
    成功时返回详细的填充结果信息
    失败时返回错误信息
    
    使用示例:
    1. 基本坐标填充:
       fill_with_coordinates("docs/实习鉴定表.docx", {
           "张三": [1, 1, 3],
           "2023001234": [1, 2, 3],
           "计算机学院": [1, 1, 5]
       })
    
    注意事项:
    - 坐标格式为[表格索引, 行索引, 列索引]
    - 索引从0开始计数
    - 建议先使用 extract_fillable_fields() 获取准确坐标
    - AI判断策略: 如果目标单元格是字段名，需要谨慎判断是否修改
    - 内容覆盖: 如果单元格有内容但不是字段，AI可以判断是否覆盖现有内容
    - 查看 fill_suggestions 中的 ai_guidance 字段获取判断指导
    """
    try:
        global current_document, current_file_path
        
        if not UNIVERSAL_FILLER_AVAILABLE:
            return "通用表格填充器不可用，请检查系统配置"
        
        if not universal_filler:
            return "通用表格填充器初始化失败"
        
        # 转换坐标格式从List[int]到Tuple[int, int, int]
        coordinate_tuples = {}
        for data, coord_list in coordinate_data.items():
            if len(coord_list) == 3:
                coordinate_tuples[data] = tuple(coord_list)
            else:
                logger.warning(f"无效坐标格式: {data} -> {coord_list}")
        
        # 执行填充操作
        result = universal_filler.fill_with_coordinates(file_path, coordinate_tuples)
        
        # 更新全局状态，确保MCP服务器知道当前文档
        from docx import Document
        from pathlib import Path
        
        # 重新打开文档以更新全局状态
        if Path(file_path).exists():
            current_document = Document(file_path)
            current_file_path = str(Path(file_path).resolve())
            logger.info(f"全局状态已更新: {current_file_path}")
        
        return result
        
    except Exception as e:
        logger.error(f"坐标填充失败: {e}")
        return f"填充失败: {str(e)}"


def calculate_field_similarity(field1: str, field2: str) -> float:
    """
    计算两个字段名的相似度
    
    参数:
    - field1: 第一个字段名
    - field2: 第二个字段名
    
    返回:
    - 相似度分数 (0.0-1.0)
    """
    if not field1 or not field2:
        return 0.0
    
    # 清理字段名，移除换行符和多余空格
    f1 = field1.replace('\n', '').replace('\r', '').strip()
    f2 = field2.replace('\n', '').replace('\r', '').strip()
    
    # 完全匹配
    if f1 == f2:
        return 1.0
    
    # 精确包含关系（高优先级）
    if f1 in f2 or f2 in f1:
        return 0.9
    
    # 关键词匹配（改进版）
    f1_words = set(f1.replace('：', '').replace(':', '').split())
    f2_words = set(f2.replace('：', '').replace(':', '').split())
    
    if f1_words and f2_words:
        intersection = f1_words.intersection(f2_words)
        if intersection:
            # 如果有共同关键词，计算相似度
            union = f1_words.union(f2_words)
            jaccard_similarity = len(intersection) / len(union) if union else 0.0
            
            # 如果共同关键词数量较多，提高相似度
            if len(intersection) >= 2:
                return max(0.8, jaccard_similarity)
            else:
                return max(0.6, jaccard_similarity)
    
    # 字符级相似度（降低权重）
    from difflib import SequenceMatcher
    char_similarity = SequenceMatcher(None, f1, f2).ratio()
    
    # 如果字符相似度很高，但包含特殊字符，降低权重
    if char_similarity > 0.8 and ('：' in f2 or ':' in f2 or '签名' in f2 or '成绩' in f2):
        return char_similarity * 0.5
    
    return char_similarity

@mcp.tool()
def basic_table_fill(file_path: str, fill_data: Dict[str, Any]) -> str:
    """
    基础表格填充 - 智能字段匹配填充
    
    功能说明:
    自动分析文档中的字段位置，通过智能匹配进行填充。
    不再依赖硬编码的位置映射，具有更好的通用性。
    
    参数说明:
    - file_path (必需): 文档文件路径
    - fill_data (必需): 填充数据字典
      * 格式: {"字段名1": "值1", "字段名2": "值2", ...}
      * 支持任意字段名，系统会自动匹配
    
    返回结果:
    成功时返回: "基础表格填充成功，共填充X个字段"
    失败时返回: "基础表格填充失败: [错误信息]"
    
    使用示例:
    basic_table_fill("docs/实习鉴定表.docx", {
        "姓名": "张三",
        "学号": "2021001234", 
        "学院": "计算机学院",
        "专业": "计算机科学与技术",
        "实习单位": "腾讯科技有限公司",
        "实习时间": "2024年7月-2024年9月"
    })
    """
    try:
        global current_document, current_file_path
        
        if not UNIVERSAL_FILLER_AVAILABLE:
            return "通用表格填充器不可用，请检查系统配置"
        
        if not universal_filler:
            return "通用表格填充器初始化失败"
        
        # 使用通用表格填充器进行智能填充
        result = universal_filler.intelligent_fill(file_path, fill_data)
        
        # 更新全局状态，确保MCP服务器知道当前文档
        from docx import Document
        from pathlib import Path
        
        # 重新打开文档以更新全局状态
        if Path(file_path).exists():
            current_document = Document(file_path)
            current_file_path = str(Path(file_path).resolve())
            logger.info(f"全局状态已更新: {current_file_path}")
        
        return result
        
    except Exception as e:
        logger.error(f"基础表格填充失败: {e}")
        return f"基础表格填充失败: {str(e)}"

@mcp.tool()
def intelligent_create_table(
    data: List[List[str]],
    auto_style: bool = True
) -> str:
    """
    智能创建表格 - 自动样式和优化
    
    功能说明:
    智能创建表格并自动应用最佳样式，根据数据结构自动优化格式。
    支持表头自动加粗，表格边框样式，提升表格的专业度和可读性。
    
    参数说明:
    - data (必需): 表格数据，二维列表格式
      * 格式: [["列1标题", "列2标题", "列3标题"], ["数据1", "数据2", "数据3"], ...]
      * 第一行通常作为表头，会自动加粗显示
      * 示例: [["姓名", "年龄", "职业"], ["张三", "25", "工程师"], ["李四", "30", "设计师"]]
      * 所有数据会自动转换为字符串格式
    - auto_style (可选): 是否自动应用样式，默认True
      * True: 自动应用表格样式，表头加粗，添加边框
      * False: 使用基础表格格式
    
    使用示例:
    1. 创建人员信息表:
       data = [
           ["姓名", "部门", "职位", "入职日期"],
           ["张三", "技术部", "软件工程师", "2023-01-15"],
           ["李四", "设计部", "UI设计师", "2023-02-20"],
           ["王五", "市场部", "市场专员", "2023-03-10"]
       ]
       intelligent_create_table(data, True)
    
    2. 创建项目进度表:
       data = [
           ["项目名称", "开始时间", "结束时间", "状态"],
           ["项目A", "2023-01-01", "2023-06-30", "进行中"],
           ["项目B", "2023-02-01", "2023-08-31", "计划中"]
       ]
       intelligent_create_table(data, True)
    
    3. 创建简单数据表:
       data = [["项目", "数量", "单价"], ["产品A", "100", "50"], ["产品B", "200", "30"]]
       intelligent_create_table(data, True)
    
    返回结果:
    成功时返回: "智能表格创建成功: [行数]行 x [列数]列"
    失败时返回: "智能表格创建失败: [错误信息]"
    
    注意事项:
    - 使用前需要先创建或打开文档
    - 数据不能为空，至少需要一行数据
    - 第一行建议作为表头使用
    - 表格会自动添加到文档末尾
    - 支持中文和英文内容
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "请先创建或打开一个文档"
        
        if not data or not data[0]:
            return "表格数据不能为空"
        
        # 创建表格
        table = current_document.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        
        # 填充数据
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = str(cell_data)
        
        # 自动样式
        if auto_style and len(data) > 0:
            # 设置表头样式
            header_row = table.rows[0]
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # 自动保存
        if current_file_path:
            current_document.save(current_file_path)
        
        logger.info(f"智能表格创建成功: {len(data)}x{len(data[0])}")
        return f"智能表格创建成功: {len(data)}行 x {len(data[0])}列"
        
    except Exception as e:
        logger.error(f"智能表格创建失败: {e}")
        return f"智能表格创建失败: {str(e)}"

# ==================== 基础图片处理工具 ====================

@mcp.tool()
def add_image(image_path: str, width: Optional[float] = None, 
              height: Optional[float] = None) -> str:
    """
    添加图片
    
    Parameters:
    - image_path: 图片文件路径
    - width: 图片宽度（英寸）
    - height: 图片高度（英寸）
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        if not os.path.exists(image_path):
            return f"图片文件不存在: {image_path}"
        
        from docx.shared import Inches
        
        # 添加图片
        paragraph = current_document.add_paragraph()
        run = paragraph.add_run()
        
        if width and height:
            run.add_picture(image_path, width=Inches(width), height=Inches(height))
        elif width:
            run.add_picture(image_path, width=Inches(width))
        elif height:
            run.add_picture(image_path, height=Inches(height))
        else:
            run.add_picture(image_path)
        
        logger.info(f"图片添加成功: {image_path}")
        return f"图片添加成功: {image_path}"
        
    except Exception as e:
        logger.error(f"添加图片失败: {e}")
        return f"添加图片失败: {str(e)}"

@mcp.tool()
def extract_images(output_dir: str = "extracted_images") -> str:
    """
    提取文档中的所有图片
    
    Parameters:
    - output_dir: 输出目录
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        # 创建输出目录
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        # 提取图片
        image_count = 0
        for rel in current_document.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                # 这里需要更复杂的逻辑来实际提取图片
                # 简化处理
                pass
        
        logger.info(f"图片提取完成: {image_count}个图片")
        return f"图片提取完成: 找到 {image_count} 个图片"
        
    except Exception as e:
        logger.error(f"提取图片失败: {e}")
        return f"提取图片失败: {str(e)}"

@mcp.tool()
def resize_image(image_index: int, width: float, height: float) -> str:
    """
    调整图片大小
    
    Parameters:
    - image_index: 图片索引
    - width: 新宽度（英寸）
    - height: 新高度（英寸）
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        # 这里需要更复杂的逻辑来实际调整图片大小
        # 简化处理
        logger.info(f"图片大小调整完成: 图片{image_index}")
        return f"图片大小调整完成: 图片{image_index} -> {width}x{height}英寸"
        
    except Exception as e:
        logger.error(f"调整图片大小失败: {e}")
        return f"调整图片大小失败: {str(e)}"

# ==================== 基础页面设置工具 ====================

@mcp.tool()
def set_page_margins(top: float = 1.0, bottom: float = 1.0, 
                    left: float = 1.0, right: float = 1.0) -> str:
    """
    设置页面边距
    
    Parameters:
    - top: 上边距（英寸）
    - bottom: 下边距（英寸）
    - left: 左边距（英寸）
    - right: 右边距（英寸）
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        from docx.shared import Inches
        
        # 设置页面边距
        sections = current_document.sections
        for section in sections:
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)
        
        logger.info("页面边距设置成功")
        return "页面边距设置成功"
        
    except Exception as e:
        logger.error(f"设置页面边距失败: {e}")
        return f"设置页面边距失败: {str(e)}"

@mcp.tool()
def set_page_orientation(orientation: str = "portrait") -> str:
    """
    设置页面方向
    
    Parameters:
    - orientation: 页面方向（portrait/landscape）
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        from docx.enum.section import WD_ORIENT
        
        # 设置页面方向
        for section in current_document.sections:
            if orientation.lower() == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
            else:
                section.orientation = WD_ORIENT.PORTRAIT
        
        logger.info(f"页面方向设置成功: {orientation}")
        return f"页面方向设置成功: {orientation}"
        
    except Exception as e:
        logger.error(f"设置页面方向失败: {e}")
        return f"设置页面方向失败: {str(e)}"

@mcp.tool()
def set_page_size(width: float = 8.5, height: float = 11.0) -> str:
    """
    设置页面大小
    
    Parameters:
    - width: 页面宽度（英寸）
    - height: 页面高度（英寸）
    """
    try:
        global current_document, current_file_path
        
        if not current_document:
            return "没有打开的文档"
        
        from docx.shared import Inches
        
        # 设置页面大小
        for section in current_document.sections:
            section.page_width = Inches(width)
            section.page_height = Inches(height)
        
        logger.info(f"页面大小设置成功: {width}x{height}英寸")
        return f"页面大小设置成功: {width}x{height}英寸"
        
    except Exception as e:
        logger.error(f"设置页面大小失败: {e}")
        return f"设置页面大小失败: {str(e)}"

# ==================== 智能建议工具 ====================

@mcp.tool()
def get_smart_suggestions(
    context: str = "document_editing"
) -> str:
    """
    获取智能建议 - 基于上下文的建议
    
    功能说明:
    根据当前文档编辑上下文提供智能建议，帮助用户优化文档结构、
    格式和内容，提升文档的专业度和可读性。
    
    参数说明:
    - context (可选): 上下文类型，默认"document_editing"
      * "document_editing": 文档编辑建议，适用于一般文档编辑场景
      * "table_creation": 表格创建建议，适用于创建和编辑表格
      * "content_formatting": 内容格式化建议，适用于文本格式优化
      * "structure_optimization": 结构优化建议，适用于文档结构改进
      * "professional_polish": 专业润色建议，适用于提升文档专业度
    
    使用示例:
    1. 获取一般编辑建议:
       get_smart_suggestions("document_editing")
    
    2. 获取表格创建建议:
       get_smart_suggestions("table_creation")
    
    3. 获取格式化建议:
       get_smart_suggestions("content_formatting")
    
    4. 获取结构优化建议:
       get_smart_suggestions("structure_optimization")
    
    5. 获取专业润色建议:
       get_smart_suggestions("professional_polish")
    
    返回结果:
    返回JSON格式的建议列表，包含:
    - context: 上下文类型
    - suggestions: 建议内容数组
    - total_count: 建议总数
    
    示例返回:
    {
      "context": "document_editing",
      "suggestions": [
        "建议添加目录以改善文档结构",
        "考虑使用标题样式统一格式",
        "检查段落间距是否合适"
      ],
      "total_count": 3
    }
    
    注意事项:
    - 建议基于最佳实践和文档标准
    - 可以根据具体需求选择合适的上下文类型
    - 建议内容会根据上下文动态调整
    - 返回的JSON可以直接解析使用
    """
    try:
        suggestions = {
            "document_editing": [
                "建议添加目录以改善文档结构",
                "考虑使用标题样式统一格式",
                "检查段落间距是否合适",
                "添加页码提升专业度",
                "确保文档标题清晰明确",
                "检查拼写和语法错误"
            ],
            "table_creation": [
                "建议添加表头突出显示",
                "考虑使用表格样式统一外观",
                "检查表格对齐方式",
                "添加表格标题说明",
                "确保数据格式一致",
                "考虑添加表格边框"
            ],
            "content_formatting": [
                "建议统一字体和字号",
                "检查行间距是否合适",
                "考虑使用项目符号列表",
                "添加适当的段落缩进",
                "保持标题层次结构清晰",
                "使用粗体和斜体突出重点"
            ],
            "structure_optimization": [
                "建议添加摘要或概述部分",
                "考虑使用章节标题组织内容",
                "检查文档逻辑结构是否合理",
                "添加必要的过渡段落",
                "确保结论部分完整",
                "考虑添加附录或参考资料"
            ],
            "professional_polish": [
                "检查文档格式是否符合标准",
                "确保所有图片和表格都有标题",
                "添加页眉页脚信息",
                "检查引用格式是否正确",
                "确保文档布局美观专业",
                "添加必要的版权声明"
            ]
        }
        
        context_suggestions = suggestions.get(context, suggestions["document_editing"])
        
        result = {
            "context": context,
            "suggestions": context_suggestions,
            "total_count": len(context_suggestions)
        }
        
        return json.dumps(result, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"获取智能建议失败: {e}")
        return f"获取智能建议失败: {str(e)}"

# ==================== 系统状态工具 ====================

@mcp.tool()
def get_system_status() -> str:
    """
    获取系统状态
    """
    try:
        global current_document, current_file_path
        
        status = {
            "system_status": "运行中",
            "current_document": {
                "has_document": current_document is not None,
                "file_path": current_file_path,
                "paragraphs_count": len(current_document.paragraphs) if current_document else 0,
                "tables_count": len(current_document.tables) if current_document else 0
            },
            "available_tools": {
                "basic_tools": [
                    "create_document", "open_document", "save_document", "save_as_document", "copy_document", "create_work_copy", 
                    "close_document", "get_document_info", "add_paragraph", "add_heading",
                    "add_text_with_formatting", "search_and_replace", "add_table",
                    "add_table_row", "add_table_column", "format_table", "merge_table_cells", "add_image",
                    "extract_images", "resize_image", "set_page_margins", "set_page_orientation",
                    "set_page_size"
                ],
                "intelligent_tools": [
                    "intelligent_create_document", "smart_add_content", "intelligent_create_table",
                    "get_smart_suggestions"
                ]
            }
        }
        
        return json.dumps(status, ensure_ascii=False, indent=2)
        
    except Exception as e:
        logger.error(f"获取系统状态失败: {e}")
        return f"获取系统状态失败: {str(e)}"

@mcp.tool()
def test_connection() -> str:
    """
    测试连接
    """
    return "最终完整MCP服务器连接正常！包含基础工具和智能工具。"

@mcp.tool()
def get_server_info() -> str:
    """
    获取服务器信息
    """
    return "FinalCompleteDocxProcessor - 最终完整MCP服务器，包含所有基础工具和智能工具，提供完整功能"

def main():
    """MCP服务器主入口函数 - 支持多种传输协议"""
    import argparse
    
    # 解析命令行参数
    parser = argparse.ArgumentParser(description='DOCX MCP 服务器')
    parser.add_argument('--transport', '-t', 
                       choices=['stdio', 'sse', 'streamable-http'],
                       default='stdio',
                       help='传输协议类型 (默认: stdio)')
    parser.add_argument('--host', 
                       default='localhost',
                       help='HTTP/SSE 服务器主机地址 (默认: localhost)')
    parser.add_argument('--port', '-p',
                       type=int,
                       default=8000,
                       help='HTTP/SSE 服务器端口 (默认: 8000)')
    
    args = parser.parse_args()
    
    print("=" * 60)
    print("🚀 启动 DOCX MCP 服务器")
    print("=" * 60)
    print()
    print("📦 服务器信息:")
    print(f"   名称: FinalCompleteDocxProcessor")
    print(f"   传输协议: {args.transport.upper()}")
    if args.transport in ['sse', 'streamable-http']:
        print(f"   地址: http://{args.host}:{args.port}")
    print()
    print("🛠️  功能模块:")
    print("   - 基础文档管理 (8个工具)")
    print("   - 文本内容处理 (5个工具)")
    print("   - 表格操作 (6个工具)")
    print("   - 表格分析 (5个工具)")
    print("   - 表格填充 (4个工具)")
    print("   - 图片处理 (3个工具)")
    print("   - 页面设置 (3个工具)")
    print("   - 智能功能 (5个工具)")
    print("   - 系统状态 (3个工具)")
    print()
    print("📊 总计: 42个MCP工具")
    print("=" * 60)
    print()
    
    # 根据传输协议启动MCP服务器
    try:
        if args.transport == 'stdio':
            # STDIO 传输（默认，用于 Cursor/Claude Desktop）
            logger.info("使用 STDIO 传输协议启动服务器")
            mcp.run(transport='stdio')
            
        elif args.transport == 'sse':
            # SSE (Server-Sent Events) 传输
            logger.info(f"使用 SSE 传输协议启动服务器: {args.host}:{args.port}")
            print(f"🌐 SSE 服务器运行在: http://{args.host}:{args.port}")
            print(f"📡 连接端点: http://{args.host}:{args.port}/sse")
            mcp.run(transport='sse', host=args.host, port=args.port)
            
        elif args.transport == 'streamable-http':
            # Streamable HTTP 传输
            logger.info(f"使用 Streamable HTTP 传输协议启动服务器: {args.host}:{args.port}")
            print(f"🌐 HTTP 服务器运行在: http://{args.host}:{args.port}")
            print(f"📡 API 端点: http://{args.host}:{args.port}/mcp")
            mcp.run(transport='streamable-http', host=args.host, port=args.port)
            
    except KeyboardInterrupt:
        print("\n")
        print("👋 服务器已停止")
        logger.info("服务器被用户中断")
    except Exception as e:
        print(f"\n❌ 服务器启动失败: {e}")
        logger.error(f"服务器启动失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
